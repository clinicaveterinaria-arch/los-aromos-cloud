from datetime import datetime, date, timedelta
from zoneinfo import ZoneInfo
import os
import uuid
import mimetypes
import re
from typing import Optional
from io import BytesIO, StringIO
from openpyxl import load_workbook, Workbook
import csv
import base64
import json
import zipfile
import urllib.request
import urllib.error
from urllib.parse import urlencode
import unicodedata
from difflib import SequenceMatcher
from pypdf import PdfReader
from fastapi import FastAPI, Request, Form, Depends, HTTPException, UploadFile, File
from fastapi.responses import RedirectResponse, HTMLResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from sqlalchemy import or_, and_, text
from passlib.context import CryptContext
from starlette.middleware.sessions import SessionMiddleware
from supabase import create_client

SUPABASE_URL = os.getenv("SUPABASE_URL", "")
SUPABASE_KEY = os.getenv("SUPABASE_KEY", "")

print("SUPABASE_URL =", SUPABASE_URL)
print("SUPABASE_KEY cargada =", bool(SUPABASE_KEY))
supabase = create_client(SUPABASE_URL, SUPABASE_KEY) if SUPABASE_URL and SUPABASE_KEY else None
from .database import Base, engine, get_db
from .models import User, Owner, Patient, ClinicalEvent, EventAttachment, Appointment, Product, Sale, SaleItem, SalePayment, WaitingListEntry, Hospitalization, HospitalizationMedication, HospitalizationFluid, LaboratoryResult
from .vademecum_parser import read_rows_from_upload, parse_vademecum_rows
from .vademecum_importer import import_vademecum
from .senasa_sync import update_from_senasa
Base.metadata.create_all(bind=engine)
# =====================================
# Zona horaria Argentina
# =====================================

ARG_TZ = ZoneInfo("America/Argentina/Buenos_Aires")

def argentina_now():
    return datetime.now(ARG_TZ).replace(tzinfo=None)
def safe_storage_filename(filename: str) -> str:
    """
    Convierte un nombre de archivo en una key segura para Supabase Storage.
    El nombre original puede seguir guardándose en la base de datos.
    """
    original = os.path.basename(filename or "archivo")

    stem, extension = os.path.splitext(original)

    # Elimina tildes/diacríticos: á -> a, ó -> o, etc.
    stem = unicodedata.normalize("NFKD", stem)
    stem = "".join(
        char for char in stem
        if not unicodedata.combining(char)
    )

    # Deja únicamente caracteres seguros
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem)

    # Evita varios guiones bajos consecutivos
    stem = re.sub(r"_+", "_", stem)

    # Limpia extremos
    stem = stem.strip("._-")

    if not stem:
        stem = "archivo"

    # También sanitizamos la extensión
    extension = re.sub(r"[^A-Za-z0-9.]", "", extension.lower())

    return f"{stem}{extension}"    
# ==========================================================
# IA CLÍNICA
# ==========================================================

def ai_clinical_summary(event):
    import json
    import urllib.request
    import urllib.error

    api_key = os.getenv("OPENAI_API_KEY", "").strip()

    fallback = {
        "summary": "Asistente IA no disponible todavía.",
        "owner_explanation": "",
        "priority": "Normal",
        "problems": [],
        "differentials": [],
        "recommended_tests": [],
        "treatment": [],
        "alerts": []
    }

    patient = getattr(event, "patient", None)
    previous_events_text = ""

    try:
        previous_events = (
            event.patient.events
            if getattr(event, "patient", None) and getattr(event.patient, "events", None)
            else []
        )

        previous_events = sorted(
            previous_events,
            key=lambda e: e.event_date or datetime.min,
            reverse=True
        )

        previous_events = [
            e for e in previous_events
            if e.id != event.id
        ][:8]

        previous_events_lines = []

        for previous in previous_events:
            previous_events_lines.append(
                f"""
Fecha: {previous.event_date.strftime('%d/%m/%Y') if previous.event_date else ''}
Tipo: {previous.event_type or ''}
Título: {previous.title or ''}
Descripción: {previous.description or ''}
Anamnesis: {previous.anamnesis or ''}
Examen físico: {previous.physical_exam or ''}
Diagnóstico: {previous.diagnosis or ''}
Tratamiento: {previous.treatment or ''}
"""
            )

        previous_events_text = "\n---\n".join(previous_events_lines)

    except Exception:
        previous_events_text = ""
    age_text = ""

    try:
        if patient and getattr(patient, "birth_date", None):
            today = argentina_now().date()
            birth = patient.birth_date
            years = today.year - birth.year - ((today.month, today.day) < (birth.month, birth.day))
            age_text = f"{years} años"
    except Exception:
        age_text = ""
    clinical_text = f"""
Paciente: {getattr(patient, "name", "") if patient else ""}
Especie: {getattr(patient, "species", "") if patient else ""}
Raza: {getattr(patient, "breed", "") if patient else ""}
Sexo: {getattr(patient, "sex", "") if patient else ""}
Edad: {age_text}
Castrado: {getattr(patient, "neutered", "") if patient else ""}
Fecha del evento: {event.event_date.strftime('%d/%m/%Y %H:%M') if event.event_date else ""}
Peso paciente: {getattr(patient, "weight", "") if patient else ""}
Peso evento: {getattr(event, "weight", "") or ""}

Tipo de evento: {getattr(event, "event_type", "") or ""}
Título: {getattr(event, "title", "") or ""}
Descripción: {getattr(event, "description", "") or ""}
Anamnesis: {getattr(event, "anamnesis", "") or ""}
Examen físico: {getattr(event, "physical_exam", "") or ""}
Diagnóstico cargado: {getattr(event, "diagnosis", "") or ""}
Tratamiento cargado: {getattr(event, "treatment", "") or ""}

Constantes:
Temperatura: {getattr(event, "temperature", "") or ""}
FC: {getattr(event, "heart_rate", "") or ""}
FR: {getattr(event, "respiratory_rate", "") or ""}
Mucosas: {getattr(event, "mucous_membranes", "") or ""}
TLLC/CRT: {getattr(event, "crt", "") or ""}
Hidratación: {getattr(event, "hydration", "") or ""}

ECG:
FC ECG: {getattr(event, "ecg_hr", "") or ""}
Ritmo: {getattr(event, "ecg_rhythm", "") or ""}
PR: {getattr(event, "ecg_pr", "") or ""}
QRS: {getattr(event, "ecg_qrs", "") or ""}
QT: {getattr(event, "ecg_qt", "") or ""}
Eje: {getattr(event, "ecg_axis", "") or ""}
Interpretación ECG: {getattr(event, "ecg_interpretation", "") or ""}

Eco:
AI/Ao: {getattr(event, "eco_aiao", "") or ""}
FS: {getattr(event, "eco_fs", "") or ""}
ACVIM: {getattr(event, "eco_acvim", "") or ""}
Diagnóstico eco: {getattr(event, "eco_diagnosis", "") or ""}
Tratamiento eco: {getattr(event, "eco_treatment", "") or ""}

RX:
VHS: {getattr(event, "rx_vhs", "") or ""}
VLAS: {getattr(event, "rx_vlas", "") or ""}
Patrón pulmonar: {getattr(event, "rx_lung_pattern", "") or ""}
Edema: {getattr(event, "rx_edema", "") or ""}
Congestión: {getattr(event, "rx_congestion", "") or ""}
Observaciones RX: {getattr(event, "rx_observations", "") or ""}
Historia clínica previa relevante:
{previous_events_text}
"""

    if not api_key:
        fallback["summary"] = "Falta configurar OPENAI_API_KEY en Render."
        fallback["alerts"] = ["No se encontró OPENAI_API_KEY."]
        return fallback
    event_type_text = (getattr(event, "event_type", "") or "").lower()

    study_instruction = ""

    if "radiografía" in event_type_text or "rx" in event_type_text:
        study_instruction = """
Este evento corresponde a RADIOGRAFÍA.
Si hay imágenes adjuntas, analizalas como radiografías veterinarias.
Describí hallazgos radiológicos con prudencia.
No afirmes lesiones si la calidad o proyección no lo permite.
Correlacioná siempre con la clínica.
"""
    elif "ecg" in event_type_text:
        study_instruction = """
Este evento corresponde a ELECTROCARDIOGRAMA.
Priorizá ritmo, frecuencia, intervalos, eje eléctrico, alteraciones de conducción y signos compatibles con agrandamiento de cámaras.
Si hay imagen adjunta, interpretala junto con los valores cargados.
"""
    elif "ecocardiografía" in event_type_text:
        study_instruction = """
Este evento corresponde a ECOCARDIOGRAFÍA.
Priorizá AI/Ao, FS, FE, EPSS, ACVIM, válvulas, cámaras cardíacas, función sistólica y signos de congestión o hipertensión pulmonar.
"""
    elif "laboratorio" in event_type_text:
        study_instruction = """
Este evento corresponde a LABORATORIO.
Interpretá alteraciones hematológicas y bioquímicas en contexto clínico.
Indicá patrones compatibles, diferenciales y estudios complementarios si faltan datos.
"""
    image_count = 0

    content_blocks = []

    try:
        for attachment in getattr(event, "attachments", []) or []:
            file_url = getattr(attachment, "file_path", "") or ""
            file_name = (getattr(attachment, "filename", "") or "").lower()

            if (
                file_url.startswith("http")
                and (
                    file_name.endswith(".jpg")
                    or file_name.endswith(".jpeg")
                    or file_name.endswith(".png")
                    or file_name.endswith(".webp")
                )
            ):
                content_blocks.append({
                    "type": "input_image",
                    "image_url": file_url
                })

                image_count += 1

                if image_count >= 4:
                    break

    except Exception:
        pass

    prompt = f"""
Sos un veterinario clínico experto en pequeños animales.

{study_instruction}

Analizá TODA la información clínica disponible.
Si hay imágenes adjuntas, utilizalas junto con los datos clínicos.
Nunca inventes hallazgos.

Respondé únicamente en JSON con esta estructura:

{{
  "summary":"",
  "owner_explanation":"",
  "priority":"Normal",
  "problems":[],
  "differentials":[],
  "recommended_tests":[],
  "treatment":[],
  "alerts":[]
}}

Información clínica:

{clinical_text}

Cantidad de imágenes adjuntas enviadas para análisis:
{image_count}
"""

    content_blocks.insert(0, {
        "type": "input_text",
        "text": prompt
    })
    payload = {
        "model": os.getenv("OPENAI_MODEL", "gpt-5.4-mini"),
        "input": [
            {
                "role": "user",
                "content": content_blocks
            }
        ],
        "max_output_tokens": 1400
    }
    try:
        req = urllib.request.Request(
            "https://api.openai.com/v1/responses",
            data=json.dumps(payload).encode("utf-8"),
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            method="POST"
        )

        with urllib.request.urlopen(req, timeout=35) as response:
            data = json.loads(response.read().decode("utf-8"))

        output_text = data.get("output_text", "")

        if not output_text:
            parts = []
            for item in data.get("output", []):
                for content in item.get("content", []):
                    if content.get("type") in ["output_text", "text"]:
                        parts.append(content.get("text", ""))
            output_text = "\n".join(parts).strip()

        cleaned = output_text.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.replace("```json", "").replace("```", "").strip()

        result = json.loads(cleaned)

        for key in fallback:
            if key not in result:
                result[key] = fallback[key]

        if result.get("priority") not in ["Normal", "Alta", "Crítica"]:
            result["priority"] = "Normal"

        return result

    except Exception as e:
        print("ERROR IA CLINICA:", str(e))
        fallback["summary"] = "No se pudo generar el análisis IA en este momento."
        fallback["alerts"] = [f"Error IA: {str(e)}"]
        return fallback

app = FastAPI(title='Los Aromos Cloud')
app.add_middleware(SessionMiddleware, secret_key=os.getenv('SECRET_KEY', 'dev-secret-change-me'))
app.mount('/static', StaticFiles(directory='app/static'), name='static')
@app.get('/service-worker.js')
def service_worker():
    service_worker_path = 'app/static/service-worker.js'

    if not os.path.exists(service_worker_path):
        raise HTTPException(
            status_code=404,
            detail='Service worker no encontrado'
        )

    with open(service_worker_path, 'r', encoding='utf-8') as file:
        content = file.read()

    return Response(
        content=content,
        media_type='application/javascript',
        headers={
            'Service-Worker-Allowed': '/',
            'Cache-Control': 'no-cache'
        }
    )    
os.makedirs("app/uploads", exist_ok=True)
os.makedirs('app/uploads', exist_ok=True)
app.mount('/uploads', StaticFiles(directory='app/uploads'), name='uploads')
templates = Jinja2Templates(directory='app/templates')
# ==========================================================
# ACTUALIZADOR DE PRODUCTOS DESDE FACTURA PDF
# ==========================================================

INVOICE_PREVIEW_DIR = "app/uploads/invoice_previews"
os.makedirs(INVOICE_PREVIEW_DIR, exist_ok=True)

INVOICE_STOPWORDS = {
    "X", "DE", "DEL", "CON", "AL", "EL", "LA", "LOS", "LAS", "PARA",
    "JM", "LABYES", "COMP", "COMPRIMIDO", "COMPRIMIDOS", "INY", "INYECTABLE",
    "ML", "CC", "GR", "GRS", "G", "MG", "KG", "DOSIS", "DS"
}


def invoice_money(value: str) -> float:
    value = (value or "").strip().replace("$", "").replace(" ", "")
    if not value:
        return 0.0
    if "," in value and "." in value:
        value = value.replace(",", "")
    elif "," in value:
        value = value.replace(".", "").replace(",", ".")
    try:
        return float(value)
    except ValueError:
        return 0.0


def invoice_normalize(text_value: str) -> str:
    text_value = unicodedata.normalize("NFKD", text_value or "")
    text_value = "".join(ch for ch in text_value if not unicodedata.combining(ch))
    text_value = text_value.upper().replace("REPEXXIN", "REPEXXXIN")
    text_value = re.sub(r"\bC\s*/\s*", " ", text_value)
    text_value = re.sub(r"[^A-Z0-9,.]+", " ", text_value)
    text_value = re.sub(r"\s+", " ", text_value).strip()
    return text_value


def invoice_tokens(text_value: str) -> set[str]:
    normalized = invoice_normalize(text_value)
    return {
        token for token in normalized.split()
        if token and token not in INVOICE_STOPWORDS
    }


def invoice_numbers(text_value: str, unit: str) -> list[float]:
    normalized = invoice_normalize(text_value).replace(",", ".")
    pattern = rf"(?<!\d)(\d+(?:\.\d+)?)\s*{unit}\b"
    return [float(v) for v in re.findall(pattern, normalized)]


def invoice_weight_ranges(text_value: str) -> list[tuple[float, float]]:
    normalized = invoice_normalize(text_value).replace(",", ".")
    return [
        (float(a), float(b))
        for a, b in re.findall(r"(\d+(?:\.\d+)?)\s*A\s*(\d+(?:\.\d+)?)\s*KG", normalized)
    ]


def invoice_has_numeric_conflict(source: str, target: str) -> bool:
    for unit in ("MG", "ML", "GR", "KG"):
        source_values = invoice_numbers(source, unit)
        target_values = invoice_numbers(target, unit)
        if source_values and target_values and not set(source_values).intersection(target_values):
            return True

    source_ranges = invoice_weight_ranges(source)
    target_ranges = invoice_weight_ranges(target)
    if source_ranges and target_ranges and not set(source_ranges).intersection(target_ranges):
        return True

    return False


def invoice_explicit_product(source_name: str, products: list[Product]):
    source = invoice_normalize(source_name)

    def find_name(*needles):
        for product in products:
            name = invoice_normalize(product.name or "")
            if all(invoice_normalize(n) in name for n in needles):
                return product
        return None

    if "MELTRA GOLD" in source and re.search(r"10\s*A\s*20\s*KG", source):
        return find_name("MELTRA GOLD", "GRANDE")

    if "MELTRA GOLD" in source and re.search(r"\b10\s*KG\b", source) and " A " not in source:
        return find_name("MELTRA GOLD", "CHICO")

    if "NEXGARD COMBO L" in source or ("NEXGARD COMBO" in source and "0.9 ML" in source.replace(",", ".")):
        return find_name("NEXGARD", "CAT", "GRANDE")

    if "REPEXXXIN" in source and re.search(r"4\s*A\s*25\s*KG", source):
        return find_name("REPEXXXIN", "4", "25")

    if "REPEXXXIN" in source and re.search(r"25\s*A\s*50\s*KG", source):
        product = find_name("REPEXXXIN", "25", "60")
        return product or find_name("REPEXXXIN", "25", "50")

    return None


def invoice_match_score(source_name: str, product_name: str) -> float:
    source_norm = invoice_normalize(source_name)
    product_norm = invoice_normalize(product_name)

    if invoice_has_numeric_conflict(source_norm, product_norm):
        return 0.0

    source_tokens = invoice_tokens(source_norm)
    product_tokens = invoice_tokens(product_norm)
    union = source_tokens | product_tokens
    token_score = len(source_tokens & product_tokens) / len(union) if union else 0.0
    sequence_score = SequenceMatcher(None, source_norm, product_norm).ratio()

    return round((token_score * 0.65) + (sequence_score * 0.35), 4)


def invoice_find_match(db: Session, provider_cuit: str, description: str, products: list[Product]):
    alias_row = db.execute(
        text("""
            SELECT product_id
            FROM product_invoice_aliases
            WHERE provider_cuit = :provider_cuit
              AND source_normalized = :source_normalized
            LIMIT 1
        """),
        {
            "provider_cuit": provider_cuit or "",
            "source_normalized": invoice_normalize(description),
        }
    ).mappings().first()

    if alias_row:
        product = db.get(Product, alias_row["product_id"])
        if product and product.active:
            return product, 1.0, "alias"

    explicit = invoice_explicit_product(description, products)
    if explicit:
        return explicit, 1.0, "regla"

    scored = sorted(
        ((invoice_match_score(description, p.name or ""), p) for p in products),
        key=lambda item: item[0],
        reverse=True
    )

    best_score, best_product = scored[0] if scored else (0.0, None)
    second_score = scored[1][0] if len(scored) > 1 else 0.0

    if best_product and best_score >= 0.78 and (best_score - second_score) >= 0.08:
        return best_product, best_score, "texto"

    return None, best_score, "revisar"


def invoice_parse_pdf(content: bytes) -> dict:
    reader = PdfReader(BytesIO(content))
    page_texts = [(page.extract_text(extraction_mode="layout") or "") for page in reader.pages]
    metadata_pages = [(page.extract_text() or "") for page in reader.pages]

    selected_pages = []
    for page_text in page_texts:
        upper = page_text.upper()
        if "DUPLICADO" in upper and any("ORIGINAL" in p.upper() for p in page_texts):
            continue
        selected_pages.append(page_text)

    full_text = "\n".join(selected_pages)
    metadata_text = "\n".join(metadata_pages)
    invoice_number_match = re.search(r"N[º°]?\s*:?\s*([0-9]{4}-[0-9]{8})", metadata_text, re.I)
    date_match = re.search(r"\b(\d{2}/\d{2}/\d{4})\b", metadata_text)
    vat_match = re.search(r"IVA\s*(\d+(?:[.,]\d+)?)\s*%", metadata_text, re.I)
    cuit_matches = re.findall(r"\b\d{2}-\d{8}-\d\b", metadata_text)

    provider_name = ""
    for candidate in ("LA RED COMERCIAL SRL",):
        if candidate in metadata_text.upper():
            provider_name = candidate.title()
            break

    provider_cuit = ""
    if cuit_matches:
        provider_cuit = "30-71657247-8" if "30-71657247-8" in cuit_matches else cuit_matches[-1]

    vat_percent = invoice_money(vat_match.group(1)) if vat_match else 21.0

    line_pattern = re.compile(
        r"^\s*(\d+(?:[.,]\d+)?)\s+(.+?)\s+"
        r"(?:(\d+(?:[.,]\d+)?)%\s+)?"
        r"(?:(\d+(?:[.,]\d+)?)%\s+)?"
        r"([\d.,]+)\s+([\d.,]+)\s*$"
    )

    items = []
    seen = set()
    in_items = False

    for raw_line in full_text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        upper = line.upper()

        if "CANTIDAD" in upper and "DESCRIP" in upper and "PRECIO" in upper:
            in_items = True
            continue
        if upper in {"ORIGINAL", "DUPLICADO"} or upper.startswith("SUBTOTAL"):
            in_items = False
        if not in_items:
            continue

        match = line_pattern.match(line)
        if not match:
            continue

        quantity = invoice_money(match.group(1))
        description = match.group(2).strip()
        discount_1 = invoice_money(match.group(3)) if match.group(3) else 0.0
        discount_2 = invoice_money(match.group(4)) if match.group(4) else 0.0
        unit_net = invoice_money(match.group(5))
        line_total = invoice_money(match.group(6))

        if quantity <= 0 or unit_net <= 0 or not re.search(r"[A-ZÁÉÍÓÚÑ]", description.upper()):
            continue

        key = (quantity, invoice_normalize(description), unit_net, line_total)
        if key in seen:
            continue
        seen.add(key)

        calculated_total = quantity * unit_net
        if line_total and abs(calculated_total - line_total) > max(1.0, line_total * 0.015):
            unit_net = line_total / quantity

        items.append({
            "quantity": quantity,
            "description": description,
            "discount_percent": max(discount_1, discount_2),
            "unit_net": round(unit_net, 4),
            "unit_cost_vat": round(unit_net * (1 + vat_percent / 100), 4),
            "line_total": round(line_total, 4),
        })

    if not items:
        raise ValueError("No se pudieron reconocer renglones de productos en el PDF.")

    return {
        "provider_name": provider_name,
        "provider_cuit": provider_cuit,
        "invoice_number": invoice_number_match.group(1) if invoice_number_match else "",
        "invoice_date": date_match.group(1) if date_match else "",
        "vat_percent": vat_percent,
        "items": items,
    }


def invoice_preview_path(token: str) -> str:
    safe_token = re.sub(r"[^a-zA-Z0-9_-]", "", token or "")
    return os.path.join(INVOICE_PREVIEW_DIR, f"{safe_token}.json")


def invoice_save_preview(payload: dict) -> str:
    token = uuid.uuid4().hex
    with open(invoice_preview_path(token), "w", encoding="utf-8") as file:
        json.dump(payload, file, ensure_ascii=False, indent=2)
    return token


def invoice_load_preview(token: str) -> dict:
    path = invoice_preview_path(token)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="La vista previa venció o no existe.")
    with open(path, "r", encoding="utf-8") as file:
        return json.load(file)
def get_pending_count():
    return 0

def get_waiting_count():
    return 0

templates.env.globals['get_waiting_count'] = get_waiting_count
templates.env.globals['get_pending_count'] = get_pending_count
pwd_context = CryptContext(schemes=['bcrypt'], deprecated='auto')

EVENT_TYPES = ['Consulta clínica','Control','Vacuna','Desparasitación','Radiografía','ECG','Ecocardiografía','Ecografía','Laboratorio','Cirugía','Anestesia','Internación','Alta','Otro procedimiento']
DUE_ACTIVE_MARKER = '[DEUDA_ACTIVA]'
DUE_CLOSED_MARKER = '[DEUDA_CERRADA]'


def pending_action_matches_event(pending_event, real_event):
    pending_type = (pending_event.event_type or '').lower()
    real_type = (real_event.event_type or '').lower()

    pending_text = f'{pending_event.title or ""} {pending_event.description or ""}'.lower()
    real_text = f'{real_event.title or ""} {real_event.description or ""} {real_event.vaccine_name or ""} {real_event.dewormer_product or ""}'.lower()

    if pending_type == real_type:
        return True

    if 'vacun' in pending_text and real_type == 'vacuna':
        return True

    if (
        'despar' in pending_text
        or 'antiparas' in pending_text
        or 'aprax' in pending_text
        or 'pipeta' in pending_text
    ) and real_type == 'desparasitación':
        return True

    if 'control' in pending_text and real_type in ['control', 'consulta clínica']:
        return True

    return False


def close_pending_actions_after_real_event(db, patient, real_event):
    pending_actions = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .filter(ClinicalEvent.description.ilike(f'%{DUE_ACTIVE_MARKER}%'))
        .all()
    )

    for pending_event in pending_actions:
        if pending_action_matches_event(pending_event, real_event):
            pending_event.description = (pending_event.description or '').replace(
                DUE_ACTIVE_MARKER,
                DUE_CLOSED_MARKER
            )
            pending_event.reminder_date = None
def init_db():
    # Base.metadata.create_all(bind=engine)
    db = next(get_db())

    with engine.begin() as conn:
        conn.execute(text("ALTER TABLE patients ADD COLUMN IF NOT EXISTS photo_url TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS weight FLOAT"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS temperature FLOAT"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS heart_rate INTEGER"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS respiratory_rate INTEGER"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS mucous_membranes TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS crt TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS hydration TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS anamnesis TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS physical_exam TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS vaccine_name TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS vaccine_lot TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS vaccine_expiration TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS next_vaccine_date TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS dewormer_product TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS dewormer_drug TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS dewormer_dose TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS next_deworming_date TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_hr TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_rhythm TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_p TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_pr TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_qrs TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_st TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_t TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_qt TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_axis TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_interpretation TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_aiao TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_fs TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_acvim TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_diagnosis TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_treatment TEXT DEFAULT ''"))
        # ECG ampliado
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_p_mv TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_qrs_mv TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_t_mv TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_qtc TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_polarity TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_arrhythmia TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_conduction TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS ecg_notes TEXT DEFAULT ''"))

        # Ecocardiografía ampliada
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_epss TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_lvidd TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_lvids TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_ivsd TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_ivss TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_lvpwd TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_lvpws TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_fe TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_la_size TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_lv_size TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_rv_size TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_ra_size TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_mitral TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_tricuspid TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_aortic TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_pulmonary TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_pulmonary_htn TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_pericardium TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_doppler TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS eco_observations TEXT DEFAULT ''"))

        # Radiografía cardiológica
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_vhs TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_vlas TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_heart_size TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_left_atrium TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_left_heart TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_right_heart TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_pulmonary_vessels TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_lung_pattern TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_edema TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_congestion TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_trachea TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE clinical_events ADD COLUMN IF NOT EXISTS rx_observations TEXT DEFAULT ''"))
        conn.execute(text("ALTER TABLE sales ADD COLUMN IF NOT EXISTS patient_id INTEGER REFERENCES patients(id)"))
        conn.execute(text("ALTER TABLE sales ADD COLUMN IF NOT EXISTS owner_id INTEGER REFERENCES owners(id)"))
        conn.execute(text("ALTER TABLE sales ADD COLUMN IF NOT EXISTS discount_percent FLOAT DEFAULT 0"))
        conn.execute(text("ALTER TABLE sales ADD COLUMN IF NOT EXISTS discount_amount FLOAT DEFAULT 0"))
        conn.execute(text("ALTER TABLE sales ADD COLUMN IF NOT EXISTS credit_surcharge_percent FLOAT DEFAULT 0"))
        conn.execute(text("ALTER TABLE sales ADD COLUMN IF NOT EXISTS credit_surcharge_amount FLOAT DEFAULT 0"))
        conn.execute(text("ALTER TABLE sales ADD COLUMN IF NOT EXISTS cost_total FLOAT DEFAULT 0"))
        conn.execute(text("ALTER TABLE sales ADD COLUMN IF NOT EXISTS profit_amount FLOAT DEFAULT 0"))
        conn.execute(text("ALTER TABLE sales ADD COLUMN IF NOT EXISTS margin_percent FLOAT DEFAULT 0"))
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS vademecum_active_ingredients (
                id SERIAL PRIMARY KEY,
                name VARCHAR(200) DEFAULT '',
                category VARCHAR(120) DEFAULT '',
                species VARCHAR(120) DEFAULT '',
                dog_dose TEXT DEFAULT '',
                cat_dose TEXT DEFAULT '',
                route VARCHAR(120) DEFAULT '',
                frequency VARCHAR(120) DEFAULT '',
                indications TEXT DEFAULT '',
                contraindications TEXT DEFAULT '',
                interactions TEXT DEFAULT '',
                warnings TEXT DEFAULT '',
                observations TEXT DEFAULT '',
                active BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """))

        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS vademecum_brands (
                id SERIAL PRIMARY KEY,
                active_ingredient_id INTEGER REFERENCES vademecum_active_ingredients(id),
                brand_name VARCHAR(200) DEFAULT '',
                laboratory VARCHAR(180) DEFAULT '',
                presentation VARCHAR(200) DEFAULT '',
                concentration VARCHAR(120) DEFAULT '',
                active BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """))
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS vademecum_drugs (
                id SERIAL PRIMARY KEY,
                commercial_name VARCHAR(200) DEFAULT '',
                active_ingredient VARCHAR(200) DEFAULT '',
                category VARCHAR(120) DEFAULT '',
                species VARCHAR(120) DEFAULT '',
                dose TEXT DEFAULT '',
                route VARCHAR(120) DEFAULT '',
                frequency VARCHAR(120) DEFAULT '',
                indications TEXT DEFAULT '',
                contraindications TEXT DEFAULT '',
                observations TEXT DEFAULT '',
                active BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """))       
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS event_attachments (
                id SERIAL PRIMARY KEY,
                event_id INTEGER REFERENCES clinical_events(id),
                filename VARCHAR(255) NOT NULL,
                file_path VARCHAR(500) NOT NULL
            )
        """))
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS products (
                id SERIAL PRIMARY KEY,
                rubro VARCHAR(120) DEFAULT '',
                tipo VARCHAR(120) DEFAULT '',
                name VARCHAR(200) DEFAULT '',
                code VARCHAR(100) DEFAULT '',
                barcode VARCHAR(120) DEFAULT '',
                cost_price FLOAT,
                sale_price FLOAT,
                margin_percent FLOAT,
                stock FLOAT,
                min_stock FLOAT,
                expiration_date DATE,
                provider VARCHAR(180) DEFAULT '',
                manufacturer VARCHAR(180) DEFAULT '',
                active BOOLEAN DEFAULT TRUE,
                notes TEXT DEFAULT '',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """))
        conn.execute(text("""
        CREATE TABLE IF NOT EXISTS sale_payments (
            id SERIAL PRIMARY KEY,
            sale_id INTEGER REFERENCES sales(id),
            method VARCHAR(50),
            amount FLOAT DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """))
        conn.execute(text("""
        CREATE TABLE IF NOT EXISTS hospitalizations (
            id SERIAL PRIMARY KEY,
            patient_id INTEGER REFERENCES patients(id),
            clinical_event_id INTEGER REFERENCES clinical_events(id),
            admission_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            discharge_date TIMESTAMP NULL,
            status VARCHAR(40) DEFAULT 'Internado',
            cage VARCHAR(80) DEFAULT '',
            responsible_vet VARCHAR(120) DEFAULT '',
            reason TEXT DEFAULT '',
            diagnosis TEXT DEFAULT '',
            treatment_plan TEXT DEFAULT '',
            notes TEXT DEFAULT '',
            initial_weight FLOAT,
            initial_temperature FLOAT,
            initial_heart_rate INTEGER,
            initial_respiratory_rate INTEGER,
            initial_mucous_membranes VARCHAR(100) DEFAULT '',
            initial_crt VARCHAR(50) DEFAULT '',
            initial_hydration VARCHAR(100) DEFAULT '',
            discharge_summary TEXT DEFAULT '',
            discharge_indications TEXT DEFAULT '',
            created_by VARCHAR(100) DEFAULT 'admin',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """)) 
        conn.execute(text("""
        CREATE TABLE IF NOT EXISTS hospitalization_medications (
            id SERIAL PRIMARY KEY,
            hospitalization_id INTEGER REFERENCES hospitalizations(id),
            medication_name VARCHAR(200) DEFAULT '',
            dose VARCHAR(120) DEFAULT '',
            route VARCHAR(80) DEFAULT '',
            frequency VARCHAR(120) DEFAULT '',
            scheduled_time VARCHAR(50) DEFAULT '',
            status VARCHAR(40) DEFAULT 'Pendiente',
            notes TEXT DEFAULT '',
            created_by VARCHAR(100) DEFAULT 'admin',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """))
        conn.execute(text("ALTER TABLE hospitalization_medications ADD COLUMN IF NOT EXISTS applied_at TIMESTAMP NULL"))
        conn.execute(text("ALTER TABLE hospitalization_fluids ADD COLUMN IF NOT EXISTS finished_at TIMESTAMP NULL"))
        conn.execute(text("ALTER TABLE hospitalization_fluids ADD COLUMN IF NOT EXISTS finished_by VARCHAR(100) DEFAULT ''"))
        conn.execute(text("ALTER TABLE hospitalization_medications ADD COLUMN IF NOT EXISTS applied_by VARCHAR(100) DEFAULT ''"))
        conn.execute(text("""
        CREATE TABLE IF NOT EXISTS hospitalization_fluids (
            id SERIAL PRIMARY KEY,
            hospitalization_id INTEGER REFERENCES hospitalizations(id),
            fluid_type VARCHAR(200) DEFAULT '',
            fluid_rate VARCHAR(80) DEFAULT '',
            ml_kg_h VARCHAR(80) DEFAULT '',
            drip_set VARCHAR(80) DEFAULT '',
            notes TEXT DEFAULT '',
            status VARCHAR(40) DEFAULT 'Activo',
            created_by VARCHAR(100) DEFAULT 'admin',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """))
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS product_invoice_aliases (
                id SERIAL PRIMARY KEY,
                provider_cuit VARCHAR(30) DEFAULT '',
                provider_name VARCHAR(180) DEFAULT '',
                source_description TEXT DEFAULT '',
                source_normalized TEXT NOT NULL,
                product_id INTEGER NOT NULL REFERENCES products(id),
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(provider_cuit, source_normalized)
            )
        """))
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS purchase_invoices (
                id SERIAL PRIMARY KEY,
                provider_cuit VARCHAR(30) DEFAULT '',
                provider_name VARCHAR(180) DEFAULT '',
                invoice_number VARCHAR(50) DEFAULT '',
                invoice_date VARCHAR(20) DEFAULT '',
                vat_percent FLOAT DEFAULT 21,
                items_count INTEGER DEFAULT 0,
                updated_products INTEGER DEFAULT 0,
                created_products INTEGER DEFAULT 0,
                skipped_items INTEGER DEFAULT 0,
                created_by VARCHAR(100) DEFAULT '',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(provider_cuit, invoice_number)
            )
        """))
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS product_cost_history (
                id SERIAL PRIMARY KEY,
                product_id INTEGER NOT NULL REFERENCES products(id),
                provider_cuit VARCHAR(30) DEFAULT '',
                provider_name VARCHAR(180) DEFAULT '',
                invoice_number VARCHAR(50) DEFAULT '',
                invoice_date VARCHAR(20) DEFAULT '',
                old_cost FLOAT DEFAULT 0,
                new_cost FLOAT DEFAULT 0,
                old_sale FLOAT DEFAULT 0,
                new_sale FLOAT DEFAULT 0,
                margin_percent FLOAT DEFAULT 0,
                quantity_added FLOAT DEFAULT 0,
                created_by VARCHAR(100) DEFAULT '',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """))
    try:
        admin = db.query(User).filter(User.username == 'admin').first()
        if not admin:
            admin = User(
                username='admin',
                full_name='Carolina',
                password_hash=pwd_context.hash('losaromos')
            )
            db.add(admin)
            db.commit()
    finally:
        db.close()

@app.on_event("startup")
def on_startup():
    init_db()

def current_user(request: Request, db: Session = Depends(get_db)):
    username = request.session.get('user')
    if not username:
        return None
    return db.query(User).filter(User.username == username, User.is_active == True).first()

def require_user(request: Request, db: Session = Depends(get_db)):
    user = current_user(request, db)
    if not user:
        raise HTTPException(status_code=303, headers={'Location': '/login'})
    return user

@app.get('/login', response_class=HTMLResponse)
def login_form(request: Request):
    return templates.TemplateResponse('login.html', {'request': request, 'error': None})

@app.post('/login')
def login(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == username).first()
    if not user or not pwd_context.verify(password, user.password_hash):
        return templates.TemplateResponse('login.html', {'request': request, 'error': 'Usuario o clave incorrectos'})
    request.session['user'] = user.username
    return RedirectResponse('/', status_code=303)

@app.get('/logout')
def logout(request: Request):
    request.session.clear()
    return RedirectResponse('/login', status_code=303)

@app.get('/', response_class=HTMLResponse)
def home(request: Request, db: Session = Depends(get_db), user: User = Depends(require_user)):
    today = argentina_now().date()
    day_start = datetime.combine(today, datetime.min.time())
    day_end = datetime.combine(today, datetime.max.time())

    stats = {
        'owners': db.query(Owner).count(),
        'patients': db.query(Patient).count(),
        'events': db.query(ClinicalEvent).count(),
        'pending': db.query(ClinicalEvent).filter(
            ClinicalEvent.reminder_date != None
        ).count(),
        'overdue': db.query(ClinicalEvent).filter(
            ClinicalEvent.reminder_date != None,
            ClinicalEvent.reminder_date < today
        ).count(),
        'today': db.query(ClinicalEvent).filter(
            ClinicalEvent.reminder_date == today
        ).count(),
        'vaccines': db.query(ClinicalEvent).filter(
            ClinicalEvent.event_type == 'Vacuna'
        ).count(),
        'rx': db.query(ClinicalEvent).filter(
            ClinicalEvent.event_type == 'Radiografía'
        ).count(),
        'ecg': db.query(ClinicalEvent).filter(
            ClinicalEvent.event_type == 'ECG'
        ).count(),
    }

    latest_events = (
        db.query(ClinicalEvent)
        .order_by(ClinicalEvent.event_date.desc())
        .limit(8)
        .all()
    )

    from calendar import monthrange
    
    last_day = monthrange(today.year, today.month)[1]
    end_of_month = today.replace(day=last_day)
    
    upcoming = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.reminder_date != None)
        .filter(ClinicalEvent.reminder_date >= today)
        .filter(ClinicalEvent.reminder_date <= end_of_month)
        .order_by(ClinicalEvent.reminder_date.asc())
        .limit(8)
        .all()
    )

    overdue_events = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.reminder_date != None,
            ClinicalEvent.reminder_date < today
        )
        .order_by(ClinicalEvent.reminder_date.asc())
        .limit(5)
        .all()
    )

    today_events = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.reminder_date == today)
        .order_by(ClinicalEvent.reminder_date.asc())
        .limit(5)
        .all()
    )

    upcoming_appointments = (
        db.query(Appointment)
        .filter(Appointment.appointment_date >= day_start)
        .filter(Appointment.status.in_(['Pendiente', 'Confirmado']))
        .order_by(Appointment.appointment_date.asc(), Appointment.start_time.asc())
        .limit(8)
        .all()
    )

    today_appointments = (
        db.query(Appointment)
        .filter(Appointment.appointment_date >= day_start)
        .filter(Appointment.appointment_date <= day_end)
        .filter(Appointment.status != 'Cancelado')
        .order_by(Appointment.start_time.asc())
        .all()
    )
    
    stats['today'] = len(today_appointments)
    stats['surgery_today'] = len([
        a for a in today_appointments
        if 'cirug' in (a.service or '').lower()
    ])
    stats['ecg_today'] = len([
        a for a in today_appointments
        if 'ecg' in (a.service or '').lower() or 'ecg' in (a.title or '').lower()
    ])
    stats['rx_today'] = len([
        a for a in today_appointments
        if 'rx' in (a.service or '').lower() or 'radiograf' in (a.service or '').lower() or 'rx' in (a.title or '').lower()
    ])

    recent_patients = (
        db.query(Patient)
        .order_by(Patient.id.desc())
        .limit(8)
        .all()
    )



    today_sales = (
        db.query(Sale)
        .filter(Sale.date >= day_start)
        .filter(Sale.date <= day_end)
        .filter(Sale.status != 'cancelled')
        .filter(Sale.status != 'quote')
        .all()
    )

    today_sales_total = sum(s.total or 0 for s in today_sales)
    today_sales_count = len(today_sales)

    critical_stock_count = (
        db.query(Product)
        .filter(Product.active == True)
        .filter(Product.stock != None)
        .filter(Product.min_stock != None)
        .filter(Product.min_stock > 0)
        .filter(Product.stock <= Product.min_stock)
        .count()
    )



    dashboard_alerts = []

    if stats['overdue'] > 0:
        dashboard_alerts.append(f"Hay {stats['overdue']} recordatorios vencidos.")

    if critical_stock_count > 0:
        dashboard_alerts.append(f"Hay {critical_stock_count} productos con stock crítico.")



    if not dashboard_alerts:
        dashboard_alerts.append("Todo se ve ordenado para hoy.")

    return templates.TemplateResponse(
        'home.html',
        {
            'request': request,
            'user': user,
            'stats': stats,
            'latest_events': latest_events,
            'upcoming': upcoming,
            'overdue_events': overdue_events,
            'today_events': today_events,
            'upcoming_appointments': upcoming_appointments,
            'today_appointments': today_appointments,
            'recent_patients': recent_patients, 
            'today_sales_total': today_sales_total,
            'today_sales_count': today_sales_count,
            'critical_stock_count': critical_stock_count,
            'dashboard_alerts': dashboard_alerts,
            'today': today
        }
    )
@app.get('/quick/rx')
def quick_rx():
    return RedirectResponse('/search?quick=rx', status_code=303)
@app.get('/quick/ecg')
def quick_ecg():
    return RedirectResponse('/search?quick=ecg', status_code=303)


@app.get('/quick/eco')
def quick_eco():
    return RedirectResponse('/search?quick=eco', status_code=303)    
@app.get('/owners', response_class=HTMLResponse)
def owners(request: Request, q: str = '', db: Session = Depends(get_db), user: User = Depends(require_user)):
    query = db.query(Owner)
    if q:
        like = f'%{q}%'
        query = query.filter(or_(Owner.name.ilike(like), Owner.phone.ilike(like), Owner.whatsapp.ilike(like)))
    return templates.TemplateResponse('owners.html', {'request': request, 'owners': query.order_by(Owner.name).limit(200).all(), 'q': q})

@app.post('/owners')
def owner_create(name: str = Form(...), phone: str = Form(''), whatsapp: str = Form(''), email: str = Form(''), address: str = Form(''), notes: str = Form(''), db: Session = Depends(get_db), user: User = Depends(require_user)):
    owner = Owner(name=name, phone=phone, whatsapp=whatsapp, email=email, address=address, notes=notes)
    db.add(owner); db.commit()
    return RedirectResponse('/owners', status_code=303)
@app.get('/owners/{owner_id}/edit', response_class=HTMLResponse)
def owner_edit_form(request: Request, owner_id: int, db: Session = Depends(get_db), user: User = Depends(require_user)):
    owner = db.get(Owner, owner_id)
    if not owner:
        raise HTTPException(status_code=404)

    return templates.TemplateResponse(
        'owner_edit.html',
        {'request': request, 'owner': owner}
    )


@app.post('/owners/{owner_id}/edit')
def owner_edit_save(
    owner_id: int,
    name: str = Form(...),
    phone: str = Form(''),
    whatsapp: str = Form(''),
    email: str = Form(''),
    address: str = Form(''),
    notes: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    owner = db.get(Owner, owner_id)
    if not owner:
        raise HTTPException(status_code=404)

    owner.name = name
    owner.phone = phone
    owner.whatsapp = whatsapp
    owner.email = email
    owner.address = address
    owner.notes = notes

    db.commit()
    return RedirectResponse('/owners', status_code=303)
@app.get('/patients/new', response_class=HTMLResponse)
def patient_new(
    request: Request,
    owner_id: Optional[int] = None,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    owners = (
        db.query(Owner)
        .order_by(Owner.name)
        .all()
    )

    return templates.TemplateResponse(
        'patient_new.html',
        {
            'request': request,
            'owners': owners,
            'owner_id': owner_id
        }
    )

@app.post('/patients')
def patient_create(
    name: str = Form(...),
    owner_id: str = Form(""),
    owner_name: str = Form(""),
    owner_phone: str = Form(""),
    owner_whatsapp: str = Form(""),
    owner_address: str = Form(""),
    owner_email: str = Form(""),
    species: str = Form(""),
    breed: str = Form(""),
    sex: str = Form(""),
    birth_date: str = Form(""),
    weight: str = Form(""),
    alerts: str = Form(""),
    notes: str = Form(""),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    owner_id_value = None
    
    if owner_id and owner_id.strip():
        try:
            owner_id_value = int(owner_id)
        except ValueError:
            owner_id_value = None
    
    if owner_id_value is None:
        if owner_name.strip():
            owner = Owner(
                name=owner_name.strip(),
                phone=owner_phone.strip(),
                whatsapp=owner_whatsapp.strip(),
                address=owner_address.strip(),
                email=owner_email.strip()
            )
    
            db.add(owner)
            db.commit()
            db.refresh(owner)
    
            owner_id_value = owner.id
        else:
            raise HTTPException(
                status_code=400,
                detail="Debe seleccionar un propietario existente o crear uno nuevo."
            )

    w = float(weight.replace(',', '.')) if weight.strip() else None

    bd = None
    if birth_date and birth_date.strip():
        try:
            bd = datetime.strptime(birth_date.strip(), '%Y-%m-%d').date()
        except ValueError:
            bd = None

    p = Patient(
        name=name,
        owner_id=owner_id_value,
        species=species,
        breed=breed,
        sex=sex,
        birth_date=bd,
        weight=w,
        alerts=alerts,
        notes=notes
    )

    db.add(p)
    db.commit()

    return RedirectResponse(f'/patients/{p.id}', status_code=303)
@app.get('/agenda', response_class=HTMLResponse)
def agenda(
    request: Request,
    date: str = '',
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    if date:
        selected_day = datetime.strptime(date, '%Y-%m-%d').date()
    else:
        selected_day = argentina_now().date()

    start_dt = datetime.combine(selected_day, datetime.min.time())
    end_dt = start_dt + timedelta(days=1)
    import calendar

    cal = calendar.Calendar(firstweekday=0)
    
    month_days = cal.monthdatescalendar(
        selected_day.year,
        selected_day.month
    )
    
    prev_month = (selected_day.replace(day=1) - timedelta(days=1)).replace(day=1)
    
    if selected_day.month == 12:
        next_month = selected_day.replace(year=selected_day.year + 1, month=1, day=1)
    else:
        next_month = selected_day.replace(month=selected_day.month + 1, day=1)
    appointments = (
        db.query(Appointment)
        .filter(Appointment.appointment_date >= start_dt)
        .filter(Appointment.appointment_date < end_dt)
        .order_by(Appointment.start_time)
        .all()
    )
    month_start = selected_day.replace(day=1)
    month_end = next_month

    month_appointments = (
        db.query(Appointment)
        .filter(Appointment.appointment_date >= month_start)
        .filter(Appointment.appointment_date < month_end)
        .all()
    )

    appointments_count_by_day = {}
    appointments_tooltip_by_day = {}
    for appointment in month_appointments:
        day_key = appointment.appointment_date.day

        appointments_count_by_day[day_key] = (
            appointments_count_by_day.get(day_key, 0) + 1
        )

        patient_name = (
            appointment.patient.name
            if appointment.patient
            else "Sin paciente"
        )

        title = appointment.title or ""

        text = f"• {patient_name}"
        if title:
            text += f" - {title}"

        appointments_tooltip_by_day.setdefault(day_key, []).append(text)
    owners = (
        db.query(Owner)
        .order_by(Owner.name)
        .all()
    )
    
    patients = (
        db.query(Patient)
        .order_by(Patient.name)
        .all()
    )

    return templates.TemplateResponse(
        'agenda.html',
        {
            'request': request,
            'selected_day': selected_day,
            'appointments': appointments,
            'owners': owners,
            'patients': patients,
            'month_days': month_days,
            'prev_month': prev_month,
            'next_month': next_month,
            'appointments_count_by_day': appointments_count_by_day,
            'appointments_tooltip_by_day': appointments_tooltip_by_day,
        }
    )


@app.post('/agenda')
def agenda_create(
    service: str = Form(...),
    title: str = Form(''),
    appointment_date: str = Form(...),
    start_time: str = Form(''),
    end_time: str = Form(''),
    owner_id: str = Form(''),
    patient_id: str = Form(''),
    notes: str = Form(''),
    reminder_24h: str = Form(''),
    reminder_12h: str = Form(''),
    contact_whatsapp: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    selected_day = datetime.strptime(appointment_date, '%Y-%m-%d').date()
    appointment_dt = datetime.combine(selected_day, datetime.min.time())

    owner_value = int(owner_id) if owner_id else None
    patient_value = int(patient_id) if patient_id else None

    appointment = Appointment(
        service=service,
        title=title,
        appointment_date=appointment_dt,
        start_time=start_time,
        end_time=end_time,
        owner_id=owner_value,
        patient_id=patient_value,
        reminder_12h=True if reminder_12h else False,
        contact_whatsapp=contact_whatsapp,
        notes=notes,
        reminder_24h=True if reminder_24h else False,
        status='Pendiente'
    )

    db.add(appointment)
    db.commit()

    return RedirectResponse(f'/agenda?date={appointment_date}', status_code=303)
@app.post('/agenda/{appointment_id}/status')
def agenda_update_status(
    appointment_id: int,
    status: str = Form(...),
    date: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    appointment = db.get(Appointment, appointment_id)

    if not appointment:
        raise HTTPException(status_code=404, detail='Turno no encontrado')

    appointment.status = status
    db.commit()

    if date:
        return RedirectResponse(f'/agenda?date={date}', status_code=303)

    return RedirectResponse('/agenda', status_code=303)
@app.post('/agenda/{appointment_id}/arrived')
def agenda_patient_arrived(
    appointment_id: int,
    date: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    appointment = db.get(Appointment, appointment_id)

    if not appointment:
        raise HTTPException(status_code=404, detail='Turno no encontrado')

    existing = (
        db.query(WaitingListEntry)
        .filter(WaitingListEntry.appointment_id == appointment.id)
        .filter(WaitingListEntry.status.in_(['Esperando', 'En consulta']))
        .first()
    )

    if not existing:
        entry = WaitingListEntry(
            appointment_id=appointment.id,
            owner_id=appointment.owner_id,
            patient_id=appointment.patient_id,
            reason=appointment.service or appointment.title or 'Turno agendado',
            notes=appointment.notes or '',
            priority='Normal',
            status='Esperando',
            arrival_time=argentina_now(),
            created_by=user.username
        )

        db.add(entry)

    appointment.status = 'Confirmado'
    db.commit()

    return RedirectResponse('/waitlist', status_code=303)

@app.get('/api/waitlist/search')
def waitlist_search_api(
    q: str = '',
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    q = (q or '').strip()

    if len(q) < 2:
        return {
            'owners': [],
            'patients': []
        }

    like = f'%{q}%'

    owners = (
        db.query(Owner)
        .filter(
            or_(
                Owner.name.ilike(like),
                Owner.phone.ilike(like),
                Owner.whatsapp.ilike(like)
            )
        )
        .order_by(Owner.name)
        .limit(20)
        .all()
    )

    patients = (
        db.query(Patient)
        .join(Owner, Patient.owner_id == Owner.id, isouter=True)
        .filter(
            or_(
                Patient.name.ilike(like),
                Owner.name.ilike(like),
                Owner.phone.ilike(like),
                Owner.whatsapp.ilike(like)
            )
        )
        .order_by(Patient.name)
        .limit(20)
        .all()
    )

    return {
        'owners': [
            {
                'id': o.id,
                'name': o.name or '',
                'phone': o.whatsapp or o.phone or ''
            }
            for o in owners
        ],
        'patients': [
            {
                'id': p.id,
                'name': p.name or '',
                'species': p.species or '',
                'breed': p.breed or '',
                'owner_id': p.owner.id if p.owner else '',
                'owner_name': p.owner.name if p.owner else '',
                'owner_phone': (p.owner.whatsapp or p.owner.phone or '') if p.owner else ''
            }
            for p in patients
        ]
    }


@app.get('/api/waitlist/owner/{owner_id}/patients')
def waitlist_owner_patients_api(
    owner_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patients = (
        db.query(Patient)
        .filter(Patient.owner_id == owner_id)
        .order_by(Patient.name)
        .all()
    )

    return {
        'patients': [
            {
                'id': p.id,
                'name': p.name or '',
                'species': p.species or '',
                'breed': p.breed or ''
            }
            for p in patients
        ]
    }
@app.get('/waitlist', response_class=HTMLResponse)
def waitlist_page(
    request: Request,
    status: str = '',
    q: str = '',
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    today = argentina_now().date()
    day_start = datetime.combine(today, datetime.min.time())
    day_end = datetime.combine(today, datetime.max.time())

    query = db.query(WaitingListEntry).filter(
        WaitingListEntry.arrival_time >= day_start,
        WaitingListEntry.arrival_time <= day_end
    )

    if status:
        query = query.filter(WaitingListEntry.status == status)

    entries = query.order_by(
        WaitingListEntry.finished_at.asc().nullsfirst(),
        WaitingListEntry.arrival_time.asc()
    ).all()

    if q:
        q_lower = q.lower()
        entries = [
            e for e in entries
            if (
                (e.patient and q_lower in e.patient.name.lower()) or
                (e.owner and q_lower in e.owner.name.lower()) or
                q_lower in (e.reason or '').lower() or
                q_lower in (e.notes or '').lower()
            )
        ]

    owners = db.query(Owner).order_by(Owner.name).limit(500).all()
    patients = db.query(Patient).order_by(Patient.name).limit(500).all()

    stats = {
        'total': len(entries),
        'waiting': len([e for e in entries if e.status == 'Esperando']),
        'consulting': len([e for e in entries if e.status == 'En consulta']),
        'finished': len([e for e in entries if e.status == 'Finalizado']),
    }

    return templates.TemplateResponse(
        'waitlist.html',
        {
            'request': request,
            'entries': entries,
            'owners': owners,
            'patients': patients,
            'stats': stats,
            'selected_status': status,
            'q': q,
            'today': today,
            'now': argentina_now().replace(tzinfo=None)
        }
    )


@app.post('/waitlist')
def waitlist_create(
    owner_id: str = Form(''),
    patient_id: str = Form(''),
    reason: str = Form(''),
    priority: str = Form('Normal'),
    notes: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    entry = WaitingListEntry(
        owner_id=int(owner_id) if owner_id else None,
        patient_id=int(patient_id) if patient_id else None,
        reason=reason or 'Consulta',
        priority=priority or 'Normal',
        notes=notes or '',
        status='Esperando',
        arrival_time=argentina_now(),
        created_by=user.username
    )

    db.add(entry)
    db.commit()

    return RedirectResponse('/waitlist', status_code=303)


@app.post('/waitlist/{entry_id}/enter')
def waitlist_enter_hc(
    entry_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    entry = db.get(WaitingListEntry, entry_id)

    if not entry:
        raise HTTPException(status_code=404, detail='Entrada no encontrada')

    if entry.status != 'Finalizado':
        entry.status = 'En consulta'
        entry.started_at = argentina_now()
        db.commit()

    request.session['active_waitlist_entry_id'] = entry.id

    if entry.patient_id:
        return RedirectResponse(f'/patients/{entry.patient_id}/v2', status_code=303)

    return RedirectResponse('/waitlist', status_code=303)
   


@app.post('/waitlist/{entry_id}/status')
def waitlist_update_status(
    entry_id: int,
    status: str = Form(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    entry = db.get(WaitingListEntry, entry_id)

    if not entry:
        raise HTTPException(status_code=404, detail='Entrada no encontrada')

    entry.status = status

    if status == 'En consulta':
        entry.started_at = argentina_now()

    if status == 'Finalizado':
        entry.finished_at = argentina_now()

    db.commit()

    return RedirectResponse('/waitlist', status_code=303)
@app.get('/agenda/{appointment_id}/edit', response_class=HTMLResponse)
def agenda_edit(
    request: Request,
    appointment_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    appointment = db.get(Appointment, appointment_id)

    if not appointment:
        raise HTTPException(status_code=404, detail='Turno no encontrado')

    owners = db.query(Owner).order_by(Owner.name).limit(300).all()
    patients = db.query(Patient).order_by(Patient.name).limit(300).all()

    return templates.TemplateResponse(
        'agenda_edit.html',
        {
            'request': request,
            'appointment': appointment,
            'owners': owners,
            'patients': patients
        }
    )
@app.post('/agenda/{appointment_id}/edit')
def agenda_update(
    appointment_id: int,
    service: str = Form(...),
    title: str = Form(''),
    appointment_date: str = Form(...),
    start_time: str = Form(''),
    end_time: str = Form(''),
    owner_id: str = Form(''),
    patient_id: str = Form(''),
    notes: str = Form(''),
    reminder_24h: str = Form(''),
    reminder_12h: str = Form(''),
    contact_whatsapp: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    appointment = db.get(Appointment, appointment_id)

    if not appointment:
        raise HTTPException(status_code=404, detail='Turno no encontrado')

    selected_day = datetime.strptime(appointment_date, '%Y-%m-%d').date()
    appointment_dt = datetime.combine(selected_day, datetime.min.time())

    appointment.service = service
    appointment.title = title
    appointment.appointment_date = appointment_dt
    appointment.start_time = start_time
    appointment.end_time = end_time
    appointment.owner_id = int(owner_id) if owner_id else None
    appointment.patient_id = int(patient_id) if patient_id else None
    appointment.reminder_12h = True if reminder_12h else False
    appointment.contact_whatsapp = contact_whatsapp
    appointment.notes = notes
    appointment.reminder_24h = True if reminder_24h else False

    db.commit()

    return RedirectResponse(f'/agenda?date={appointment_date}', status_code=303)
@app.get('/search', response_class=HTMLResponse)
def search(
    request: Request,
    q: str = '',
    quick: str = '',
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    results = []
    if quick:
        request.session['quick_event_type'] = quick
    if q:
        like = f'%{q}%'
        results = db.query(Patient).join(Owner).filter(or_(Patient.name.ilike(like), Owner.name.ilike(like), Owner.phone.ilike(like), Owner.whatsapp.ilike(like))).order_by(Patient.name).limit(100).all()
    return templates.TemplateResponse(
    'search.html',
    {
        'request': request,
        'q': q,
        'results': results,
        'today': argentina_now().date()
    }
)
@app.post('/patients/{patient_id}/weight')
def patient_update_weight(
    patient_id: int,
    weight: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404)

    patient.weight = float(weight.replace(',', '.')) if weight.strip() else None
    db.commit()

    return RedirectResponse(f'/patients/{patient.id}', status_code=303)
@app.get('/patients/{patient_id}/cart', response_class=HTMLResponse)
def patient_cart(
    request: Request,
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(
            status_code=404,
            detail='Paciente no encontrado'
        )

    products = (
        db.query(Product)
        .filter(Product.active == True)
        .order_by(Product.name)
        .all()
    )
    quick_names = [
        'Consulta clínica',
        'INY CONS',
        'INY FARM GRANDE',
        'INY CONS URGENCIA',
        'INY FARM URGENCIA',
        'DESP COMPLETO CONS',
        'DESP COMPLETO FARM',
        'INT CONS',
        'INT FARM',
        'ECG',
        'RX x1',
        'RX x2',
        'RX x3'
    ]
    
    quick_services = []
    for name in quick_names:
        product = (
            db.query(Product)
            .filter(Product.active == True, Product.name == name)
            .first()
        )
    
        if product:
            quick_services.append(product)
    return templates.TemplateResponse(
        'patient_cart.html',
        {
            'request': request,
            'patient': patient,
            'products': products,
            'quick_services': quick_services
        }
    )
@app.post('/patients/{patient_id}/cart/send')
def patient_cart_send(
    patient_id: int,
    cart_json: str = Form('[]'),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    import json

    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404)

    owner_id = patient.owner_id if patient.owner_id else None

    sale = Sale(
        status='pending',
        total=0,
        payment_method='Pendiente',
        patient_id=patient.id,
        owner_id=owner_id,
        notes='Generado desde carrito clínico'
    )

    db.add(sale)
    db.flush()

    total = 0
    cost_total = 0

    try:
        cart_items = json.loads(cart_json or '[]')
    except Exception:
        cart_items = []

    for cart_item in cart_items:
        product_id = cart_item.get('productId')
        name = cart_item.get('name', '')
        qty = float(cart_item.get('quantity') or 1)
        price = float(cart_item.get('price') or 0)

        if qty <= 0:
            continue

        subtotal = qty * price
        total += subtotal

        if product_id:
            product = db.get(Product, int(product_id))

            if not product:
                continue

            product_cost = product.cost_price or 0
            cost_total += qty * product_cost

            item = SaleItem(
                sale_id=sale.id,
                product_id=product.id,
                quantity=qty,
                unit_price=price,
                subtotal=subtotal
            )

            db.add(item)

            if product.stock is not None:
                product.stock = (product.stock or 0) - qty

        else:
            # Servicio rápido sin producto asociado:
            # por ahora se guarda como nota dentro de la venta
            sale.notes = (sale.notes or '') + f'\n- {name} x {qty} = $ {subtotal:.2f}'

    sale.total = total
    sale.cost_total = cost_total
    sale.profit_amount = total - cost_total
    sale.margin_percent = ((total - cost_total) / total * 100) if total > 0 else 0

    db.commit()
    db.refresh(sale)

    return RedirectResponse(
        url=f'/sales/{sale.id}',
        status_code=303
    )
DUE_ACTIVE_MARKER = '[DEUDA_ACTIVA]'
DUE_CLOSED_MARKER = '[DEUDA_CERRADA]'
WP_SENT_MARKER = '[WHATSAPP_AVISADO]'


def extract_due_metadata(event):
    description = event.description or ''

    sent_at = ''
    original_date = event.reminder_date

    sent_match = re.search(
        r'Recordatorio (?:enviado|marcado como enviado) por WhatsApp el '
        r'(\d{2}/\d{2}/\d{4} \d{2}:\d{2})',
        description
    )

    if sent_match:
        sent_at = sent_match.group(1)

    date_match = re.search(
        r'\[FECHA_PENDIENTE:(\d{4}-\d{2}-\d{2})\]',
        description
    )

    if date_match:
        try:
            original_date = datetime.strptime(
                date_match.group(1),
                '%Y-%m-%d'
            ).date()
        except ValueError:
            pass

    event.whatsapp_sent_at = sent_at
    event.original_reminder_date = original_date

    if original_date:
        event.days_overdue = max(
            0,
            (argentina_now().date() - original_date).days
        )
    else:
        event.days_overdue = 0

    return event


def mark_due_as_whatsapp_sent(event, user):
    now = argentina_now()
    description = event.description or ''

    additions = []

    if DUE_CLOSED_MARKER in description:
        return False

    if DUE_ACTIVE_MARKER not in description:
        additions.append(DUE_ACTIVE_MARKER)

    if event.reminder_date and '[FECHA_PENDIENTE:' not in description:
        additions.append(
            f'[FECHA_PENDIENTE:{event.reminder_date.strftime("%Y-%m-%d")}]'
        )

    if WP_SENT_MARKER not in description:
        additions.append(WP_SENT_MARKER)
        additions.append(
            f'📲 Recordatorio enviado por WhatsApp el '
            f'{now.strftime("%d/%m/%Y %H:%M")} por {user.username}.'
        )
    else:
        additions.append(
            f'📲 Recordatorio reenviado por WhatsApp el '
            f'{now.strftime("%d/%m/%Y %H:%M")} por {user.username}.'
        )

    if additions:
        event.description = (
            description.strip()
            + '\n\n'
            + '\n'.join(additions)
        ).strip()

    return True


def normalize_due_text(value):
    value = (value or '').lower()

    replacements = {
        'á': 'a',
        'é': 'e',
        'í': 'i',
        'ó': 'o',
        'ú': 'u',
        'ü': 'u',
        'ñ': 'n'
    }

    for old, new in replacements.items():
        value = value.replace(old, new)

    value = re.sub(r'[^a-z0-9 ]+', ' ', value)
    value = re.sub(r'\s+', ' ', value).strip()

    return value


def due_match_score(pending_event, real_event):
    pending_type = normalize_due_text(
        pending_event.event_type
    )

    real_type = normalize_due_text(
        real_event.event_type
    )

    pending_title = normalize_due_text(
        pending_event.title
    )

    real_vaccine = normalize_due_text(
        getattr(real_event, 'vaccine_name', '') or ''
    )

    real_dewormer = normalize_due_text(
        getattr(real_event, 'dewormer_product', '') or ''
    )

    real_text = normalize_due_text(
        f'{real_event.title or ""} '
        f'{real_event.description or ""} '
        f'{real_vaccine} '
        f'{real_dewormer}'
    )

    # Una vacuna o desparasitación puede estar cargada
    # dentro de una visita cuyo tipo principal sea Consulta clínica.
    effective_real_types = {real_type}

    if real_vaccine:
        effective_real_types.add('vacuna')

    if real_dewormer:
        effective_real_types.add('desparasitacion')

    if (
        'radiografia' in real_text
        or ' rx ' in f' {real_text} '
    ):
        effective_real_types.add('radiografia')

    if (
        'ecografia' in real_text
        or 'ecocardiografia' in real_text
    ):
        effective_real_types.add('ecografia')
        effective_real_types.add('ecocardiografia')

    if (
        'electrocardiograma' in real_text
        or ' ecg ' in f' {real_text} '
    ):
        effective_real_types.add('ecg')

    if (
        'laboratorio' in real_text
        or 'analisis' in real_text
    ):
        effective_real_types.add('laboratorio')

    score = 0

    if pending_type in effective_real_types:
        score += 70

    if pending_title and pending_title in real_text:
        score += 50

    pending_tokens = {
        token
        for token in pending_title.split()
        if len(token) >= 4
        and token not in {
            'pendiente',
            'vacuna',
            'control',
            'clinico',
            'desparasitacion',
            'consulta',
            'aplicacion',
            'antiparasitario'
        }
    }

    real_tokens = set(real_text.split())

    score += (
        len(
            pending_tokens.intersection(real_tokens)
        )
        * 15
    )

    if (
        (
            pending_type == 'vacuna'
            or 'vacun' in pending_title
        )
        and 'vacuna' in effective_real_types
    ):
        score += 40

    if (
        (
            pending_type == 'desparasitacion'
            or 'despar' in pending_title
            or 'antiparas' in pending_title
            or 'pipeta' in pending_title
        )
        and 'desparasitacion' in effective_real_types
    ):
        score += 40

    if (
        pending_type == 'control'
        and (
            'control' in effective_real_types
            or 'consulta clinica' in effective_real_types
        )
    ):
        score += 35

    if WP_SENT_MARKER in (
        pending_event.description or ''
    ):
        score += 5

    return score


def close_managed_due_events(
    db,
    patient,
    real_event,
    user
):
    active_due_events = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id
        )
        .filter(
            ClinicalEvent.description.ilike(
                f'%{DUE_ACTIVE_MARKER}%'
            )
        )
        .filter(
            ~ClinicalEvent.description.ilike(
                f'%{DUE_CLOSED_MARKER}%'
            )
        )
        .all()
    )

    candidates = []

    for due_event in active_due_events:
        score = due_match_score(
            due_event,
            real_event
        )

        if score >= 60:
            candidates.append(
                (score, due_event)
            )

    if not candidates:
        return 0

    candidates.sort(
        key=lambda item: (
            -item[0],
            item[1].reminder_date or date.max,
            item[1].id
        )
    )

    due_event = candidates[0][1]

    due_event.description = (
        due_event.description or ''
    ).replace(
        DUE_ACTIVE_MARKER,
        DUE_CLOSED_MARKER
    )

    due_event.description += (
        '\n\n'
        '✅ Pendiente cumplido automáticamente.\n'
        f'Evento realizado: '
        f'{real_event.title or real_event.event_type}\n'
        f'Fecha real: '
        f'{argentina_now().strftime("%d/%m/%Y %H:%M")}\n'
        f'Registrado por: {user.username}.'
    )

    due_event.reminder_date = None

    return 1


def reconcile_existing_due_events(
    db,
    patient,
    user
):
    active_due_events = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id
        )
        .filter(
            ClinicalEvent.description.ilike(
                f'%{DUE_ACTIVE_MARKER}%'
            )
        )
        .filter(
            ~ClinicalEvent.description.ilike(
                f'%{DUE_CLOSED_MARKER}%'
            )
        )
        .all()
    )

    if not active_due_events:
        return 0

    real_events = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id
        )
        .filter(
            ~ClinicalEvent.description.ilike(
                f'%{DUE_ACTIVE_MARKER}%'
            )
        )
        .filter(
            ~ClinicalEvent.description.ilike(
                f'%{DUE_CLOSED_MARKER}%'
            )
        )
        .order_by(
            ClinicalEvent.event_date.desc()
        )
        .all()
    )

    closed_count = 0

    for due_event in active_due_events:
        due_reference_date = (
            due_event.reminder_date
            or (
                due_event.event_date.date()
                if due_event.event_date
                else None
            )
        )

        matching_real_event = None
        best_score = 0

        for real_event in real_events:
            if not real_event.event_date:
                continue

            if (
                due_reference_date
                and real_event.event_date.date()
                < due_reference_date
            ):
                continue

            score = due_match_score(
                due_event,
                real_event
            )

            if score > best_score:
                best_score = score
                matching_real_event = real_event

        if (
            matching_real_event
            and best_score >= 60
        ):
            due_event.description = (
                due_event.description or ''
            ).replace(
                DUE_ACTIVE_MARKER,
                DUE_CLOSED_MARKER
            )

            due_event.description += (
                '\n\n'
                '✅ Pendiente cumplido automáticamente.\n'
                f'Evento realizado: '
                f'{matching_real_event.title or matching_real_event.event_type}\n'
                f'Fecha real: '
                f'{matching_real_event.event_date.strftime("%d/%m/%Y %H:%M")}\n'
                f'Conciliado automáticamente por: '
                f'{user.username}.'
            )

            due_event.reminder_date = None
            closed_count += 1

    if closed_count:
        db.commit()

    return closed_count
@app.post('/patients/{patient_id}/photo')
async def patient_upload_photo(
    patient_id: int,
    photo: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404)

    if not photo or not photo.filename:
        return RedirectResponse(f'/patients/{patient_id}', status_code=303)

    if supabase is None:
        raise HTTPException(
            status_code=500,
            detail="Supabase Storage no configurado."
        )

    original_name = os.path.basename(photo.filename)
    safe_name = safe_storage_filename(original_name)
    unique_name = f"{uuid.uuid4().hex}_{safe_name}"

    content = await photo.read()

    if not content:
        return RedirectResponse(f'/patients/{patient_id}', status_code=303)

    content_type = photo.content_type or mimetypes.guess_type(original_name)[0] or "application/octet-stream"

    storage_path = f"patient_{patient_id}/profile/{unique_name}"

    try:
        supabase.storage.from_("adjuntos").upload(
            path=storage_path,
            file=content,
            file_options={
                "content-type": content_type,
                "upsert": "true"
            }
        )

        public_url = supabase.storage.from_("adjuntos").get_public_url(storage_path)

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )

    patient.photo_url = public_url
    db.commit()

    return RedirectResponse(f'/patients/{patient_id}', status_code=303)
@app.get('/patients/{patient_id}/v2', response_class=HTMLResponse)
def patient_detail_v2(
    request: Request,
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404)

    today = argentina_now().date()

    reconcile_existing_due_events(
        db,
        patient,
        user
    )

    events = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .filter(
            ~ClinicalEvent.description.ilike(
                f'%{DUE_ACTIVE_MARKER}%'
            )
        )
        .filter(
            ~ClinicalEvent.description.ilike(
                f'%{DUE_CLOSED_MARKER}%'
            )
        )
        .order_by(ClinicalEvent.event_date.desc())
        .limit(20)
        .all()
    )

    upcoming_events = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id,
            ClinicalEvent.reminder_date != None,
            ClinicalEvent.reminder_date >= today
        )
        .order_by(ClinicalEvent.reminder_date.asc())
        .limit(10)
        .all()
    )

    upcoming_appointments = (
        db.query(Appointment)
        .filter(Appointment.patient_id == patient.id)
        .filter(Appointment.appointment_date >= datetime.combine(today, datetime.min.time()))
        .filter(Appointment.status.in_(['Pendiente', 'Confirmado']))
        .order_by(Appointment.appointment_date.asc(), Appointment.start_time.asc())
        .limit(10)
        .all()
    )
    whatsapp_pending_alerts = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id
        )
        .filter(
            ClinicalEvent.description.ilike(
                f'%{DUE_ACTIVE_MARKER}%'
            )
        )
        .filter(
            ClinicalEvent.description.ilike(
                f'%{WP_SENT_MARKER}%'
            )
        )
        .filter(
            ~ClinicalEvent.description.ilike(
                f'%{DUE_CLOSED_MARKER}%'
            )
        )
        .order_by(
            ClinicalEvent.reminder_date.asc().nullslast(),
            ClinicalEvent.id.asc()
        )
        .all()
    )

    whatsapp_pending_alerts = [
        extract_due_metadata(event)
        for event in whatsapp_pending_alerts
    ]
    upcoming_patient_items = []

    for e in upcoming_events:
        icon = '📌'
        if e.event_type == 'Vacuna':
            icon = '💉'
        elif e.event_type == 'Desparasitación':
            icon = '🪱'
        elif e.event_type == 'Control':
            icon = '🩺'

        upcoming_patient_items.append({
            'date': e.reminder_date,
            'time': '',
            'icon': icon,
            'title': e.title or e.event_type or 'Recordatorio',
            'detail': e.event_type or '',
            'source': 'Historia clínica'
        })

    for a in upcoming_appointments:
        upcoming_patient_items.append({
            'date': a.appointment_date.date() if a.appointment_date else None,
            'time': a.start_time or '',
            'icon': '📅',
            'title': a.title or a.service or 'Turno agendado',
            'detail': a.notes or '',
            'source': 'Agenda'
        })

    upcoming_patient_items = [
        item for item in upcoming_patient_items
        if item.get('date')
    ]

    upcoming_patient_items = sorted(
        upcoming_patient_items,
        key=lambda item: (item['date'], item.get('time') or '')
    )[:10]

    next_visit = upcoming_patient_items[0] if upcoming_patient_items else None

    vaccine_names = sorted({
        e.vaccine_name.strip()
        for e in db.query(ClinicalEvent).all()
        if e.vaccine_name and e.vaccine_name.strip()
    })

    dewormer_names = sorted({
        e.dewormer_product.strip()
        for e in db.query(ClinicalEvent).all()
        if e.dewormer_product and e.dewormer_product.strip()
    })

    return templates.TemplateResponse(
        'patient_detail_v2.html',
        {
            'request': request,
            'patient': patient,
            'today': today,
            'events': events,
            'upcoming_events': upcoming_events,
            'upcoming_appointments': upcoming_appointments,
            'upcoming_patient_items': upcoming_patient_items,
            'whatsapp_pending_alerts': whatsapp_pending_alerts,
            'next_visit': next_visit,
            'vaccine_names': vaccine_names,
            'dewormer_names': dewormer_names,
            'user': user
        }
    )
# ==========================================================
# CENTRO DE IA CLÍNICA - HISTORIA CLÍNICA V2
# FASE 3: integración clínica completa
# ==========================================================

def _ai_clean(value):
    return str(value or '').strip()


def _ai_collect_file_blocks(events, max_pdfs=5, max_images=6):
    content_blocks = []
    file_labels = []
    seen_urls = set()
    pdf_count = 0
    image_count = 0

    for event in events:
        for attachment in getattr(event, 'attachments', []) or []:
            file_url = _ai_clean(getattr(attachment, 'file_path', ''))
            file_name = _ai_clean(getattr(attachment, 'filename', '')) or 'archivo'
            file_name_lower = file_name.lower()

            if not file_url.startswith(('http://', 'https://')):
                continue

            if file_url in seen_urls:
                continue

            if file_name_lower.endswith('.pdf') and pdf_count < max_pdfs:
                content_blocks.append({
                    'type': 'input_file',
                    'file_url': file_url
                })
                file_labels.append(f'PDF: {file_name}')
                seen_urls.add(file_url)
                pdf_count += 1
                continue

            if (
                file_name_lower.endswith(('.jpg', '.jpeg', '.png', '.webp'))
                and image_count < max_images
            ):
                content_blocks.append({
                    'type': 'input_image',
                    'image_url': file_url
                })
                file_labels.append(f'Imagen: {file_name}')
                seen_urls.add(file_url)
                image_count += 1

            if pdf_count >= max_pdfs and image_count >= max_images:
                return content_blocks, file_labels

    return content_blocks, file_labels


def _ai_event_text(event):
    attachments = []
    for attachment in getattr(event, 'attachments', []) or []:
        attachments.append(
            f"{getattr(attachment, 'filename', '') or 'archivo'}: "
            f"{getattr(attachment, 'file_path', '') or ''}"
        )

    return f"""
Fecha: {event.event_date.strftime('%d/%m/%Y %H:%M') if event.event_date else ''}
Tipo: {_ai_clean(event.event_type)}
Título: {_ai_clean(event.title)}
Descripción: {_ai_clean(event.description)}
Anamnesis: {_ai_clean(event.anamnesis)}
Examen físico: {_ai_clean(event.physical_exam)}
Diagnóstico: {_ai_clean(event.diagnosis)}
Tratamiento: {_ai_clean(event.treatment)}
Peso: {_ai_clean(event.weight)} kg
Temperatura: {_ai_clean(event.temperature)} °C
FC: {_ai_clean(event.heart_rate)} lpm
FR: {_ai_clean(event.respiratory_rate)} rpm
Mucosas: {_ai_clean(event.mucous_membranes)}
TRC: {_ai_clean(event.crt)}
Hidratación: {_ai_clean(event.hydration)}
Vacuna: {_ai_clean(event.vaccine_name)}
Desparasitario: {_ai_clean(event.dewormer_product)}
ECG: FC {_ai_clean(event.ecg_hr)}; ritmo {_ai_clean(event.ecg_rhythm)}; P {_ai_clean(event.ecg_p)}; PR {_ai_clean(event.ecg_pr)}; QRS {_ai_clean(event.ecg_qrs)}; ST {_ai_clean(event.ecg_st)}; T {_ai_clean(event.ecg_t)}; QT {_ai_clean(event.ecg_qt)}; eje {_ai_clean(event.ecg_axis)}; interpretación {_ai_clean(event.ecg_interpretation)}
Ecocardiografía: AI/Ao {_ai_clean(event.eco_aiao)}; FS {_ai_clean(event.eco_fs)}; FE {_ai_clean(getattr(event, 'eco_fe', ''))}; EPSS {_ai_clean(getattr(event, 'eco_epss', ''))}; ACVIM {_ai_clean(event.eco_acvim)}; diagnóstico {_ai_clean(event.eco_diagnosis)}; tratamiento {_ai_clean(event.eco_treatment)}
Radiografía: VHS {_ai_clean(event.rx_vhs)}; VLAS {_ai_clean(event.rx_vlas)}; corazón {_ai_clean(getattr(event, 'rx_heart_size', ''))}; vasos {_ai_clean(getattr(event, 'rx_pulmonary_vessels', ''))}; patrón pulmonar {_ai_clean(event.rx_lung_pattern)}; edema {_ai_clean(event.rx_edema)}; congestión {_ai_clean(event.rx_congestion)}; observaciones {_ai_clean(event.rx_observations)}
Adjuntos: {' | '.join(attachments) if attachments else 'Sin adjuntos'}
"""


def _ai_numeric_trend(events, field, label, unit=''):
    points = []
    for event in reversed(events):
        value = getattr(event, field, None)
        if value in [None, '']:
            continue
        try:
            number = float(str(value).replace(',', '.'))
        except Exception:
            continue
        points.append((event.event_date, number))

    if not points:
        return ''

    first_date, first_value = points[0]
    last_date, last_value = points[-1]
    delta = last_value - first_value
    direction = 'sin cambios'
    if delta > 0:
        direction = f'aumentó {abs(delta):g}{unit}'
    elif delta < 0:
        direction = f'disminuyó {abs(delta):g}{unit}'

    values = ', '.join(
        f"{dt.strftime('%d/%m/%Y') if dt else '-'}: {value:g}{unit}"
        for dt, value in points[-8:]
    )
    return f"{label}: {values}. Tendencia global: {direction}."


def _ai_medication_context(db, hospitalization):
    if not hospitalization:
        return 'Sin internación activa.', [], []

    medications = (
        db.query(HospitalizationMedication)
        .filter(HospitalizationMedication.hospitalization_id == hospitalization.id)
        .order_by(HospitalizationMedication.scheduled_time.asc())
        .all()
    )
    fluids = (
        db.query(HospitalizationFluid)
        .filter(HospitalizationFluid.hospitalization_id == hospitalization.id)
        .order_by(HospitalizationFluid.created_at.desc())
        .all()
    )

    medication_lines = [
        f"- {m.medication_name or '-'} | dosis {m.dose or '-'} | vía {m.route or '-'} | "
        f"frecuencia {m.frequency or '-'} | horario {m.scheduled_time or '-'} | estado {m.status or '-'} | notas {m.notes or '-'}"
        for m in medications
    ]
    fluid_lines = [
        f"- {f.fluid_type or '-'} | ritmo {f.fluid_rate or '-'} ml/h | {f.ml_kg_h or '-'} ml/kg/h | "
        f"equipo {f.drip_set or '-'} | estado {f.status or '-'} | notas {f.notes or '-'}"
        for f in fluids
    ]

    text_context = f"""
Internación activa desde: {hospitalization.admission_date.strftime('%d/%m/%Y %H:%M') if hospitalization.admission_date else ''}
Jaula: {hospitalization.cage or '-'}
Veterinario responsable: {hospitalization.responsible_vet or '-'}
Motivo: {hospitalization.reason or '-'}
Diagnóstico: {hospitalization.diagnosis or '-'}
Plan terapéutico: {hospitalization.treatment_plan or '-'}
Notas: {hospitalization.notes or '-'}
Peso inicial: {hospitalization.initial_weight or '-'} kg
T° inicial: {hospitalization.initial_temperature or '-'}
FC inicial: {hospitalization.initial_heart_rate or '-'}
FR inicial: {hospitalization.initial_respiratory_rate or '-'}
Mucosas iniciales: {hospitalization.initial_mucous_membranes or '-'}
TRC inicial: {hospitalization.initial_crt or '-'}
Hidratación inicial: {hospitalization.initial_hydration or '-'}

Medicación de internación:
{chr(10).join(medication_lines) if medication_lines else '- Sin medicaciones cargadas'}

Fluidoterapia:
{chr(10).join(fluid_lines) if fluid_lines else '- Sin fluidoterapia cargada'}
"""
    return text_context, medications, fluids


def _ai_vademecum_context(db, medication_names):
    names = [name.strip() for name in medication_names if name and name.strip()]
    if not names:
        return 'Sin medicaciones para cruzar con el Vademécum.', []

    rows_found = []
    seen = set()

    for medication_name in names[:20]:
        search = f"%{medication_name}%"
        rows = db.execute(text("""
            SELECT
                a.name,
                a.dog_dose,
                a.cat_dose,
                a.route,
                a.frequency,
                a.indications,
                a.contraindications,
                a.interactions,
                a.warnings,
                a.observations,
                COALESCE(string_agg(DISTINCT b.brand_name, ', '), '') AS brands
            FROM vademecum_active_ingredients a
            LEFT JOIN vademecum_brands b ON b.active_ingredient_id = a.id AND b.active = TRUE
            WHERE a.active = TRUE
              AND (
                    a.name ILIKE :search
                    OR b.brand_name ILIKE :search
              )
            GROUP BY a.id
            LIMIT 5
        """), {'search': search}).mappings().all()

        for row in rows:
            key = (row.get('name') or '').lower()
            if key in seen:
                continue
            seen.add(key)
            rows_found.append(row)

    if not rows_found:
        return 'No se encontraron coincidencias de las medicaciones actuales en el Vademécum.', []

    lines = []
    for row in rows_found:
        lines.append(
            f"- Principio activo: {row.get('name') or '-'}; marcas: {row.get('brands') or '-'}; "
            f"dosis perro: {row.get('dog_dose') or '-'}; dosis gato: {row.get('cat_dose') or '-'}; "
            f"vía: {row.get('route') or '-'}; frecuencia: {row.get('frequency') or '-'}; "
            f"indicaciones: {row.get('indications') or '-'}; contraindicaciones: {row.get('contraindications') or '-'}; "
            f"interacciones: {row.get('interactions') or '-'}; advertencias: {row.get('warnings') or '-'}; "
            f"observaciones: {row.get('observations') or '-'}"
        )

    return '\n'.join(lines), rows_found



# ==========================================================
# LABORATORIO INTELIGENTE
# ==========================================================

def _lab_float(value):
    if value is None:
        return None
    text_value = str(value).strip().replace(',', '.')
    text_value = re.sub(r'[^0-9eE+\-.]', '', text_value)
    if not text_value:
        return None
    try:
        return float(text_value)
    except (TypeError, ValueError):
        return None


def _lab_normalize_analyte(value):
    text_value = unicodedata.normalize("NFKD", str(value or ""))
    text_value = "".join(c for c in text_value if not unicodedata.combining(c))
    text_value = re.sub(r'[^a-zA-Z0-9]+', ' ', text_value.lower()).strip()
    aliases = {
        'hematocrito': 'hematocrito', 'hct': 'hematocrito', 'hto': 'hematocrito',
        'hemoglobina': 'hemoglobina', 'hgb': 'hemoglobina',
        'eritrocitos': 'eritrocitos', 'rbc': 'eritrocitos',
        'leucocitos': 'leucocitos', 'wbc': 'leucocitos',
        'plaquetas': 'plaquetas', 'plt': 'plaquetas',
        'urea': 'urea', 'bun': 'bun',
        'creatinina': 'creatinina', 'crea': 'creatinina',
        'fosforo': 'fosforo', 'phosphorus': 'fosforo',
        'alt': 'alt', 'gpt': 'alt', 'alat': 'alt',
        'ast': 'ast', 'got': 'ast', 'asat': 'ast',
        'fosfatasa alcalina': 'fosfatasa alcalina', 'alp': 'fosfatasa alcalina', 'fal': 'fosfatasa alcalina',
        'glucosa': 'glucosa', 'glucose': 'glucosa',
        'albumina': 'albumina', 'proteinas totales': 'proteinas totales',
        'sodio': 'sodio', 'na': 'sodio', 'potasio': 'potasio', 'k': 'potasio',
        'calcio': 'calcio', 'bilirrubina total': 'bilirrubina total'
    }
    return aliases.get(text_value, text_value)


def _lab_results_context(db, patient_id, limit=120):
    rows = (
        db.query(LaboratoryResult)
        .filter(LaboratoryResult.patient_id == patient_id)
        .order_by(LaboratoryResult.sample_date.desc().nullslast(), LaboratoryResult.id.desc())
        .limit(limit)
        .all()
    )
    if not rows:
        return "Sin resultados de laboratorio estructurados.", []

    lines = []
    for row in rows:
        date_text = row.sample_date.strftime('%d/%m/%Y') if row.sample_date else '-'
        ref = row.reference_text or (
            f"{row.reference_low if row.reference_low is not None else ''}"
            f"–{row.reference_high if row.reference_high is not None else ''}"
        ).strip('–')
        lines.append(
            f"{date_text} | {row.analyte}: {row.value_text or row.value_numeric or '-'} "
            f"{row.unit or ''} | Ref: {ref or '-'} | Estado: {row.flag or 'unknown'}"
        )
    return "\\n".join(lines), rows


@app.post('/patients/{patient_id}/laboratory/analyze')
async def laboratory_analyze(
    patient_id: int,
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail='Paciente no encontrado')

    api_key = os.getenv('OPENAI_API_KEY', '').strip()
    if not api_key:
        raise HTTPException(status_code=500, detail='Falta OPENAI_API_KEY en Render.')

    request_content = [{
        'type': 'input_text',
        'text': f"""
Extraé resultados de laboratorio veterinario de los archivos adjuntos.
Paciente esperado: {patient.name}. Especie: {patient.species or '-'}.

Devolvé EXCLUSIVAMENTE JSON válido, sin markdown, con esta estructura:
{{
  "patient_match": "yes|uncertain|no",
  "sample_date": "YYYY-MM-DD o vacío",
  "laboratory": "",
  "warnings": [],
  "results": [
    {{
      "panel": "Hemograma|Bioquímica|Ionograma|Orina|Otro",
      "analyte": "nombre tal como figura",
      "value_text": "valor tal como figura",
      "value_numeric": 0.0,
      "unit": "",
      "reference_low": null,
      "reference_high": null,
      "reference_text": "",
      "flag": "low|high|normal|unknown"
    }}
  ]
}}

Reglas:
- No inventes valores ni rangos.
- Si un valor es textual, dejá value_numeric en null.
- Usá el rango de referencia impreso en el informe, no uno genérico.
- Si el informe ya marca alto/bajo, respetalo.
- Si no puede determinarse el estado, flag=unknown.
- Extraé todos los analitos legibles.
"""
    }]

    filenames = []
    for upload in files[:5]:
        if not upload or not upload.filename:
            continue
        content = await upload.read()
        if not content:
            continue
        filenames.append(os.path.basename(upload.filename))
        content_type = upload.content_type or mimetypes.guess_type(upload.filename)[0] or 'application/octet-stream'
        encoded = base64.b64encode(content).decode('ascii')

        if content_type.startswith('image/'):
            request_content.append({
                'type': 'input_image',
                'image_url': f'data:{content_type};base64,{encoded}'
            })
        else:
            request_content.append({
                'type': 'input_file',
                'filename': os.path.basename(upload.filename),
                'file_data': f'data:{content_type};base64,{encoded}'
            })

    if len(request_content) == 1:
        raise HTTPException(status_code=400, detail='No se recibieron archivos válidos.')

    payload = {
        'model': os.getenv('OPENAI_MODEL', 'gpt-5.4-mini'),
        'input': [{'role': 'user', 'content': request_content}],
        'max_output_tokens': 7000
    }

    req = urllib.request.Request(
        'https://api.openai.com/v1/responses',
        data=json.dumps(payload).encode('utf-8'),
        headers={'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'},
        method='POST'
    )

    try:
        with urllib.request.urlopen(req, timeout=120) as response:
            data = json.loads(response.read().decode('utf-8'))
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f'No se pudo analizar el laboratorio: {exc}')

    output_text = data.get('output_text', '')
    if not output_text:
        parts = []
        for item in data.get('output', []) or []:
            for content_item in item.get('content', []) or []:
                if content_item.get('type') == 'output_text':
                    parts.append(content_item.get('text', ''))
        output_text = ''.join(parts)

    output_text = output_text.strip()
    output_text = re.sub(r'^```(?:json)?\s*|\s*```$', '', output_text, flags=re.I | re.S).strip()

    try:
        parsed = json.loads(output_text)
    except json.JSONDecodeError:
        raise HTTPException(status_code=502, detail='La IA no devolvió un JSON de laboratorio válido.')

    cleaned = []
    for row in parsed.get('results', []) or []:
        analyte = str(row.get('analyte') or '').strip()
        if not analyte:
            continue
        numeric = _lab_float(row.get('value_numeric'))
        low = _lab_float(row.get('reference_low'))
        high = _lab_float(row.get('reference_high'))
        flag = str(row.get('flag') or 'unknown').lower()
        if flag not in {'low', 'high', 'normal', 'unknown'}:
            flag = 'unknown'
        cleaned.append({
            'panel': str(row.get('panel') or 'Otro')[:120],
            'analyte': analyte[:180],
            'normalized_analyte': _lab_normalize_analyte(analyte)[:180],
            'value_text': str(row.get('value_text') or '')[:120],
            'value_numeric': numeric,
            'unit': str(row.get('unit') or '')[:80],
            'reference_low': low,
            'reference_high': high,
            'reference_text': str(row.get('reference_text') or '')[:120],
            'flag': flag
        })

    return JSONResponse({
        'ok': True,
        'patient_match': parsed.get('patient_match', 'uncertain'),
        'sample_date': parsed.get('sample_date', ''),
        'laboratory': parsed.get('laboratory', ''),
        'warnings': parsed.get('warnings', []),
        'filenames': filenames,
        'results': cleaned
    })


@app.get('/patients/{patient_id}/laboratory/evolution')
def laboratory_evolution(
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail='Paciente no encontrado')

    rows = (
        db.query(LaboratoryResult)
        .filter(LaboratoryResult.patient_id == patient_id)
        .order_by(LaboratoryResult.sample_date.asc().nullslast(), LaboratoryResult.id.asc())
        .all()
    )

    grouped = {}
    for row in rows:
        key = row.normalized_analyte or _lab_normalize_analyte(row.analyte)
        grouped.setdefault(key, {
            'analyte': row.analyte,
            'points': []
        })
        grouped[key]['points'].append({
            'date': row.sample_date.isoformat() if row.sample_date else '',
            'value_text': row.value_text,
            'value_numeric': row.value_numeric,
            'unit': row.unit,
            'flag': row.flag,
            'reference_text': row.reference_text
        })

    return JSONResponse({'ok': True, 'series': list(grouped.values())})


@app.post('/patients/{patient_id}/ai-center')
async def patient_ai_center(
    patient_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    import json
    import urllib.request

    patient = db.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail='Paciente no encontrado')

    payload_in = await request.json()
    action = _ai_clean(payload_in.get('action'))
    form_data = payload_in.get('form_data') or {}
    custom_question = _ai_clean(payload_in.get('question'))

    events = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .order_by(ClinicalEvent.event_date.desc())
        .limit(40)
        .all()
    )

    real_events = [
        event for event in events
        if DUE_ACTIVE_MARKER not in (event.description or '')
        and DUE_CLOSED_MARKER not in (event.description or '')
    ]

    active_hospitalization = (
        db.query(Hospitalization)
        .filter(
            Hospitalization.patient_id == patient.id,
            Hospitalization.status == 'Internado'
        )
        .order_by(Hospitalization.admission_date.desc())
        .first()
    )

    hospitalization_context, hospitalization_medications, hospitalization_fluids = (
        _ai_medication_context(db, active_hospitalization)
    )

    medication_names = [m.medication_name for m in hospitalization_medications]

    current_treatment = _ai_clean(form_data.get('treatment'))
    if current_treatment:
        for line in current_treatment.replace(',', '\n').splitlines():
            clean_line = line.strip(' -•\t')
            if clean_line:
                medication_names.append(clean_line[:120])

    vademecum_context, vademecum_rows = _ai_vademecum_context(db, medication_names)

    ecg_events = [e for e in real_events if (e.event_type or '').lower() == 'ecg'][:8]
    eco_events = [e for e in real_events if 'ecocardi' in (e.event_type or '').lower()][:8]
    rx_events = [e for e in real_events if 'radiograf' in (e.event_type or '').lower()][:8]
    lab_events = [e for e in real_events if 'laboratorio' in (e.event_type or '').lower()][:10]
    hospitalization_events = [e for e in real_events if 'internaci' in (e.event_type or '').lower()][:12]

    trend_lines = [
        _ai_numeric_trend(real_events[:20], 'weight', 'Peso', ' kg'),
        _ai_numeric_trend(real_events[:20], 'temperature', 'Temperatura', ' °C'),
        _ai_numeric_trend(real_events[:20], 'heart_rate', 'Frecuencia cardíaca', ' lpm'),
        _ai_numeric_trend(real_events[:20], 'respiratory_rate', 'Frecuencia respiratoria', ' rpm'),
        _ai_numeric_trend(ecg_events, 'ecg_hr', 'FC en ECG', ' lpm'),
        _ai_numeric_trend(ecg_events, 'ecg_axis', 'Eje eléctrico', '°'),
        _ai_numeric_trend(eco_events, 'eco_aiao', 'AI/Ao'),
        _ai_numeric_trend(eco_events, 'eco_fs', 'Fracción de acortamiento', '%'),
        _ai_numeric_trend(rx_events, 'rx_vhs', 'VHS'),
        _ai_numeric_trend(rx_events, 'rx_vlas', 'VLAS')
    ]
    trend_context = '\n'.join(line for line in trend_lines if line) or 'No hay datos numéricos suficientes para calcular tendencias.'

    file_content_blocks, file_labels = _ai_collect_file_blocks(real_events[:25])

    history_context = '\n---\n'.join(_ai_event_text(event) for event in real_events[:25])
    lab_context = '\n---\n'.join(_ai_event_text(event) for event in lab_events) or 'Sin eventos de laboratorio cargados.'
    structured_lab_context, structured_lab_rows = _lab_results_context(db, patient.id)
    cardio_context = '\n---\n'.join(_ai_event_text(event) for event in (ecg_events + eco_events + rx_events)[:20]) or 'Sin estudios cardiológicos cargados.'
    hospitalization_events_context = '\n---\n'.join(_ai_event_text(event) for event in hospitalization_events) or 'Sin evoluciones de internación en la historia clínica.'

    source_labels = []
    if real_events:
        source_labels.append(f'Historia clínica ({len(real_events[:25])} eventos)')
    if active_hospitalization:
        source_labels.append('Internación activa')
    if ecg_events:
        source_labels.append(f'ECG ({len(ecg_events)})')
    if eco_events:
        source_labels.append(f'Ecocardiografía ({len(eco_events)})')
    if rx_events:
        source_labels.append(f'Radiografías ({len(rx_events)})')
    if lab_events:
        source_labels.append(f'Laboratorio ({len(lab_events)})')
    if structured_lab_rows:
        source_labels.append(f'Laboratorio estructurado ({len(structured_lab_rows)} resultados)')
    if hospitalization_medications:
        source_labels.append(f'Medicaciones ({len(hospitalization_medications)})')
    if hospitalization_fluids:
        source_labels.append(f'Fluidoterapia ({len(hospitalization_fluids)})')
    if vademecum_rows:
        source_labels.append(f'Vademécum ({len(vademecum_rows)} coincidencias)')
    if file_labels:
        source_labels.append(f'Archivos leídos ({len(file_labels)})')

    clinical_context = f"""
DATOS DEL PACIENTE
Nombre: {patient.name or ''}
Especie: {patient.species or ''}
Raza: {patient.breed or ''}
Sexo: {patient.sex or ''}
Fecha de nacimiento: {patient.birth_date.strftime('%d/%m/%Y') if patient.birth_date else ''}
Peso actual registrado: {patient.weight or ''} kg
Alertas permanentes: {patient.alerts or ''}
Notas permanentes: {patient.notes or ''}

CONSULTA ACTUAL CARGADA EN PANTALLA
Fecha: {form_data.get('event_date', '')}
Tipo: {form_data.get('event_type', '')}
Peso: {form_data.get('weight', '')}
Temperatura: {form_data.get('temperature', '')}
FC: {form_data.get('heart_rate', '')}
FR: {form_data.get('respiratory_rate', '')}
Motivo / título: {form_data.get('title', '')}
Descripción: {form_data.get('description', '')}
Anamnesis: {form_data.get('anamnesis', '')}
Examen físico: {form_data.get('physical_exam', '')}
Diagnóstico / conclusión: {form_data.get('diagnosis', '')}
Tratamiento / indicaciones: {form_data.get('treatment', '')}

INTERNACIÓN ACTIVA
{hospitalization_context}

EVOLUCIONES DE INTERNACIÓN
{hospitalization_events_context}

TENDENCIAS CLÍNICAS Y DE ESTUDIOS
{trend_context}

ECG, ECOCARDIOGRAFÍA Y RADIOGRAFÍAS
{cardio_context}

LABORATORIO - EVENTOS
{lab_context}

LABORATORIO - RESULTADOS ESTRUCTURADOS Y EVOLUCIÓN
{structured_lab_context}

VADEMÉCUM, CONTRAINDICACIONES E INTERACCIONES
{vademecum_context}

HISTORIA CLÍNICA RECIENTE
{history_context or 'Sin historia clínica previa.'}
"""

    action_lower = action.lower()
    if 'resumir' in action_lower:
        task_instruction = """
Generá un resumen clínico integrado y breve.
Separá: motivo actual, hallazgos, problemas activos, evolución/tendencias, tratamiento actual y próximos pasos.
"""
    elif 'diferenciales' in action_lower:
        task_instruction = """
Proponé diagnósticos diferenciales jerarquizados.
Para cada uno indicá datos a favor, datos en contra o faltantes y nivel de confianza.
No afirmes diagnósticos definitivos no sustentados.
"""
    elif 'tratamiento' in action_lower:
        task_instruction = """
Sugerí un plan terapéutico orientativo integrado con el peso, la especie, la historia, los estudios, la internación y el Vademécum disponible.
Incluí dosis en mg/kg solo cuando estén sustentadas, vía, frecuencia, duración, monitoreo, ajustes y precauciones.
Revisá expresamente duplicaciones, contraindicaciones e interacciones.
"""
    elif 'estudios' in action_lower:
        task_instruction = """
Priorizá estudios complementarios según urgencia y rendimiento diagnóstico.
Indicá qué pregunta clínica responde cada estudio y qué hallazgo se busca confirmar o descartar.
Tené en cuenta los estudios ya realizados para no repetirlos sin justificación.
"""
    elif 'alertas' in action_lower or 'estado clínico' in action_lower:
        task_instruction = """
Detectá alertas clínicas, tendencias desfavorables, datos incompatibles, omisiones importantes, riesgos de medicación, interacciones y controles vencidos.
Finalizá obligatoriamente con una sola clasificación exacta:
🟢 Sin alertas importantes
🟡 Observaciones encontradas
🔴 Revisar este paciente
"""
    elif 'alta' in action_lower:
        task_instruction = """
Redactá indicaciones de alta para el propietario en lenguaje simple.
Incluí cuidados, medicación con horarios si está disponible, alimentación, controles y signos de alarma.
No uses jerga técnica sin explicarla.
"""
    elif 'derivación' in action_lower:
        task_instruction = """
Redactá un informe profesional de derivación con antecedentes, motivo, hallazgos, evolución, estudios, tratamientos, respuesta y pregunta clínica al colega receptor.
"""
    elif 'evidencia' in action_lower:
        task_instruction = """
Explicá el razonamiento clínico y el grado de evidencia de cada conclusión.
Diferenciá datos observados, inferencias y recomendaciones.
No inventes artículos, autores ni citas.
"""
    else:
        task_instruction = f"""
Respondé esta consulta clínica concreta de la veterinaria:
{custom_question or action}
Dá una respuesta directa, razonamiento, recomendaciones y nivel de confianza.
"""

    api_key = os.getenv('OPENAI_API_KEY', '').strip()
    if not api_key:
        return JSONResponse({
            'ok': False,
            'result': 'Falta configurar OPENAI_API_KEY en Render.',
            'sources': source_labels,
            'status': 'error',
            'files_read': file_labels
        })

    prompt = f"""
Sos un asistente clínico veterinario experto en pequeños animales integrado a Aromos Cloud.

ACCIÓN SOLICITADA
{action}

{task_instruction}

REGLAS OBLIGATORIAS
- Usá toda la información disponible, incluyendo historia, internación, estudios, tendencias, medicaciones y Vademécum.
- No inventes datos, resultados, imágenes ni valores.
- Diferenciá con claridad hallazgos, sospechas y recomendaciones.
- Si faltan datos, indicá exactamente cuáles.
- Correlacioná estudios con anamnesis, examen físico y evolución.
- Para tratamientos, verificá especie, peso, vía, dosis, duplicaciones, contraindicaciones e interacciones.
- No reemplazás el criterio de la veterinaria tratante.
- Respondé en español con terminología veterinaria argentina.
- No agregues bibliografía ficticia.

CONTEXTO CLÍNICO INTEGRADO
{clinical_context}
"""

    try:
        request_content = [
            {
                'type': 'input_text',
                'text': prompt
            }
        ]
        request_content.extend(file_content_blocks)

        payload = {
            'model': os.getenv('OPENAI_MODEL', 'gpt-5.4-mini'),
            'input': [
                {
                    'role': 'user',
                    'content': request_content
                }
            ],
            'max_output_tokens': 6000
        }
        req = urllib.request.Request(
            'https://api.openai.com/v1/responses',
            data=json.dumps(payload).encode('utf-8'),
            headers={
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json'
            },
            method='POST'
        )
        with urllib.request.urlopen(req, timeout=120) as response:
            data = json.loads(response.read().decode('utf-8'))

        result = data.get('output_text', '')
        if not result:
            parts = []
            for item in data.get('output', []):
                for content in item.get('content', []):
                    if content.get('type') in ['output_text', 'text']:
                        parts.append(content.get('text', ''))
            result = '\n'.join(parts).strip()

        result_lower = (result or '').lower()
        status = 'green'
        if '🔴' in result or 'revisar este paciente' in result_lower or 'urgente' in result_lower:
            status = 'red'
        elif '🟡' in result or 'observaciones encontradas' in result_lower or 'precauc' in result_lower:
            status = 'yellow'

        return JSONResponse({
            'ok': True,
            'result': result or 'La IA no devolvió contenido.',
            'sources': source_labels,
            'status': status,
            'integrated': True,
            'files_read': file_labels
        })

    except Exception as e:
        print('ERROR CENTRO IA CLÍNICA FASE 3:', str(e))
        return JSONResponse({
            'ok': False,
            'result': f'No se pudo consultar la IA: {str(e)}',
            'sources': source_labels,
            'status': 'error',
            'files_read': file_labels
        })
# ==========================================================
# IMPORTADOR INTELIGENTE DE HISTORIA CLÍNICA
# PDF, imágenes, ZIP, DOCX, XLSX y EML
# ==========================================================

IMPORT_ALLOWED_EXTENSIONS = {
    '.pdf', '.jpg', '.jpeg', '.png', '.webp',
    '.zip', '.docx', '.xlsx', '.eml'
}
IMPORT_MAX_FILE_BYTES = 25 * 1024 * 1024
IMPORT_MAX_TOTAL_BYTES = 80 * 1024 * 1024
IMPORT_MAX_EXPANDED_FILES = 100
IMPORT_MAX_ZIP_UNCOMPRESSED = 80 * 1024 * 1024
IMPORT_TEXT_LIMIT = 30000


def _import_safe_filename(filename: str) -> str:
    import unicodedata

    name = os.path.basename(filename or 'archivo').strip() or 'archivo'

    # Convierte Analía → Analia, Muñoz → Munoz, etc.
    name = unicodedata.normalize('NFKD', name)
    name = name.encode('ascii', 'ignore').decode('ascii')

    # Supabase Storage: dejar solamente caracteres seguros.
    name = re.sub(r'[^A-Za-z0-9._()-]+', '_', name)

    # Evitar nombres que comiencen o terminen con puntos o guiones bajos.
    name = name.strip('._')

    if not name:
        name = 'archivo'

    # Conservar la extensión y limitar longitud.
    base, extension = os.path.splitext(name)
    extension = re.sub(r'[^A-Za-z0-9.]', '', extension.lower())

    base = base[:150].strip('._') or 'archivo'

    return f'{base}{extension}'
def _import_extension(filename: str) -> str:
    return os.path.splitext(filename or '')[1].lower()


def _import_date_from_text(text_value: str):
    text_value = text_value or ''
    patterns = [
        r'(?<!\d)([0-3]?\d)[/.-]([01]?\d)[/.-]((?:19|20)\d{2})(?!\d)',
        r'(?<!\d)((?:19|20)\d{2})[/.-]([01]?\d)[/.-]([0-3]?\d)(?!\d)',
    ]
    for index, pattern in enumerate(patterns):
        for match in re.finditer(pattern, text_value):
            try:
                if index == 0:
                    day, month, year = map(int, match.groups())
                else:
                    year, month, day = map(int, match.groups())
                parsed = date(year, month, day)
                if date(1990, 1, 1) <= parsed <= argentina_now().date() + timedelta(days=7):
                    return parsed
            except ValueError:
                continue
    return None


def _import_guess_type(text_value: str, filename: str) -> str:
    haystack = f'{filename} {text_value}'.lower()
    rules = [
        ('Ecocardiografía', ['ecocardiograf', 'aiao', 'ai/ao', 'fracción de acortamiento']),
        ('ECG', ['electrocardiograma', 'intervalo pr', 'qrs', 'eje eléctrico']),
        ('Radiografía', ['radiograf', 'rayos x', 'vhs', 'proyección laterolateral', 'proyección ventrodorsal']),
        ('Ecografía', ['ecograf', 'ultrasonograf', 'vesícula biliar', 'parénquima hepático']),
        ('Laboratorio', ['hemograma', 'hematocrito', 'bioquímica', 'creatinina', 'urea', 'leucocitos']),
        ('Cirugía', ['cirugía', 'quirúrgic', 'procedimiento operatorio']),
        ('Internación', ['internación', 'hospitalización', 'hospitalizado']),
        ('Alta', ['alta médica', 'indicaciones al alta']),
        ('Vacuna', ['vacuna', 'vacunación', 'séxtuple', 'antirrábica']),
        ('Desparasitación', ['desparasitación', 'antiparasitario', 'pipeta']),
    ]
    for event_type, keywords in rules:
        if any(keyword in haystack for keyword in keywords):
            return event_type
    return 'Consulta clínica'


def _import_image_data_url(content: bytes, mime_type: str) -> str:
    encoded = base64.b64encode(content).decode('ascii')
    return f'data:{mime_type};base64,{encoded}'


def _import_pdf_content(content: bytes):
    text_parts = []
    images = []

    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(content))
        for page in reader.pages[:40]:
            try:
                text_parts.append(page.extract_text() or '')
            except Exception:
                continue
    except Exception as exc:
        print('IMPORTADOR PDF TEXTO:', str(exc))

    try:
        import fitz
        document = fitz.open(stream=content, filetype='pdf')
        for page_index in range(min(len(document), 4)):
            page = document.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            png = pix.tobytes('png')
            images.append(_import_image_data_url(png, 'image/png'))
        document.close()
    except Exception as exc:
        print('IMPORTADOR PDF IMÁGENES:', str(exc))

    return '\n'.join(text_parts).strip(), images


def _import_docx_text(content: bytes) -> str:
    from docx import Document
    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(' | '.join(cell.text.strip() for cell in row.cells))
    return '\n'.join(parts)


def _import_xlsx_text(content: bytes) -> str:
    workbook = load_workbook(BytesIO(content), data_only=True, read_only=True)
    parts = []
    for worksheet in workbook.worksheets[:15]:
        parts.append(f'HOJA: {worksheet.title}')
        for row_index, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
            if row_index > 300:
                break
            values = ['' if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                parts.append(' | '.join(values))
    workbook.close()
    return '\n'.join(parts)


def _import_eml_parts(filename: str, content: bytes):
    from email import policy
    from email.parser import BytesParser

    message = BytesParser(policy=policy.default).parsebytes(content)
    subject = str(message.get('subject') or '')
    sent_date = str(message.get('date') or '')
    sender = str(message.get('from') or '')
    body_parts = [f'Asunto: {subject}', f'Fecha del correo: {sent_date}', f'Remitente: {sender}']
    children = []

    if message.is_multipart():
        for part in message.walk():
            disposition = part.get_content_disposition()
            child_name = part.get_filename()
            if disposition == 'attachment' and child_name:
                payload = part.get_payload(decode=True) or b''
                if payload:
                    children.append((_import_safe_filename(child_name), payload, part.get_content_type()))
            elif part.get_content_type() == 'text/plain' and disposition != 'attachment':
                try:
                    body_parts.append(part.get_content())
                except Exception:
                    pass
    else:
        try:
            body_parts.append(message.get_content())
        except Exception:
            pass

    return '\n'.join(body_parts), children


def _import_expand_uploaded_files(uploaded_files):
    expanded = []
    total_bytes = 0

    for original_name, content, mime_type in uploaded_files:
        if len(expanded) >= IMPORT_MAX_EXPANDED_FILES:
            break
        if not content:
            continue
        if len(content) > IMPORT_MAX_FILE_BYTES:
            raise HTTPException(status_code=413, detail=f'{original_name}: supera 25 MB.')

        total_bytes += len(content)
        if total_bytes > IMPORT_MAX_TOTAL_BYTES:
            raise HTTPException(status_code=413, detail='La carga total supera 80 MB.')

        extension = _import_extension(original_name)
        if extension not in IMPORT_ALLOWED_EXTENSIONS:
            continue

        if extension == '.zip':
            try:
                with zipfile.ZipFile(BytesIO(content)) as archive:
                    members = [member for member in archive.infolist() if not member.is_dir()]
                    if sum(member.file_size for member in members) > IMPORT_MAX_ZIP_UNCOMPRESSED:
                        raise HTTPException(status_code=413, detail=f'{original_name}: ZIP demasiado grande.')
                    for member in members:
                        if len(expanded) >= IMPORT_MAX_EXPANDED_FILES:
                            break
                        child_name = _import_safe_filename(member.filename)
                        child_extension = _import_extension(child_name)
                        if child_extension not in IMPORT_ALLOWED_EXTENSIONS or child_extension == '.zip':
                            continue
                        child_content = archive.read(member)
                        if len(child_content) > IMPORT_MAX_FILE_BYTES:
                            continue
                        expanded.append((child_name, child_content, mimetypes.guess_type(child_name)[0] or 'application/octet-stream'))
            except zipfile.BadZipFile:
                raise HTTPException(status_code=400, detail=f'{original_name}: ZIP inválido.')
            continue

        if extension == '.eml':
            body_text, children = _import_eml_parts(original_name, content)
            expanded.append((original_name, content, mime_type or 'message/rfc822', body_text))
            for child_name, child_content, child_mime in children:
                if len(expanded) >= IMPORT_MAX_EXPANDED_FILES:
                    break
                if _import_extension(child_name) in IMPORT_ALLOWED_EXTENSIONS - {'.zip', '.eml'}:
                    expanded.append((child_name, child_content, child_mime))
            continue

        expanded.append((original_name, content, mime_type))

    return expanded


def _import_extract_content(filename: str, content: bytes, mime_type: str, preloaded_text: str = ''):
    extension = _import_extension(filename)
    text_value = preloaded_text or ''
    images = []

    try:
        if extension == '.pdf':
            text_value, images = _import_pdf_content(content)
        elif extension in {'.jpg', '.jpeg', '.png', '.webp'}:
            image_mime = mime_type or mimetypes.guess_type(filename)[0] or 'image/jpeg'
            images = [_import_image_data_url(content, image_mime)]
        elif extension == '.docx':
            text_value = _import_docx_text(content)
        elif extension == '.xlsx':
            text_value = _import_xlsx_text(content)
    except Exception as exc:
        print('IMPORTADOR EXTRACCIÓN:', filename, str(exc))

    return (text_value or '')[:IMPORT_TEXT_LIMIT], images[:4]


def _import_ai_analysis(patient, filename: str, text_value: str, images):
    api_key = os.getenv('OPENAI_API_KEY', '').strip()
    fallback_date = _import_date_from_text(text_value)
    fallback_type = _import_guess_type(text_value, filename)
    fallback = {
        'date': fallback_date.isoformat() if fallback_date else argentina_now().date().isoformat(),
        'event_type': fallback_type,
        'title': os.path.splitext(filename)[0][:140],
        'summary': (text_value[:900].strip() if text_value.strip() else 'Documento importado sin texto extraíble.'),
        'patient_match': 'unknown',
        'warnings': []
    }

    if not api_key:
        fallback['warnings'].append('La IA no está configurada; revisar fecha, tipo y resumen antes de importar.')
        return fallback

    allowed_types = ', '.join(EVENT_TYPES)
    prompt = f"""
Sos un asistente veterinario que clasifica documentos para importar una historia clínica.
No inventes diagnósticos ni datos ausentes.
Elegí como fecha la fecha clínica del estudio o consulta, no la fecha de impresión, salvo que sea la única disponible.
Comprobá si el documento parece pertenecer al paciente indicado.

Paciente de destino:
Nombre: {patient.name or ''}
Especie: {patient.species or ''}
Raza: {patient.breed or ''}
Propietario: {patient.owner.name if patient.owner else ''}

Archivo: {filename}
Tipos permitidos: {allowed_types}
Texto extraído:
{text_value}

Respondé únicamente JSON válido:
{{
  "date":"YYYY-MM-DD",
  "event_type":"uno de los tipos permitidos",
  "title":"título breve",
  "summary":"resumen fiel de hasta 1200 caracteres",
  "patient_match":"yes|no|unknown",
  "warnings":[]
}}
"""

    content_blocks = [{'type': 'input_text', 'text': prompt}]
    for image_url in images[:4]:
        content_blocks.append({'type': 'input_image', 'image_url': image_url})

    payload = {
        'model': os.getenv('OPENAI_MODEL', 'gpt-5.4-mini'),
        'input': [{'role': 'user', 'content': content_blocks}],
        'max_output_tokens': 900
    }

    try:
        request_data = urllib.request.Request(
            'https://api.openai.com/v1/responses',
            data=json.dumps(payload).encode('utf-8'),
            headers={
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json'
            },
            method='POST'
        )
        with urllib.request.urlopen(request_data, timeout=60) as response:
            response_data = json.loads(response.read().decode('utf-8'))

        output_text = response_data.get('output_text', '')
        if not output_text:
            pieces = []
            for output_item in response_data.get('output', []):
                for content_item in output_item.get('content', []):
                    if content_item.get('type') in {'output_text', 'text'}:
                        pieces.append(content_item.get('text', ''))
            output_text = '\n'.join(pieces)

        cleaned = output_text.strip().replace('```json', '').replace('```', '').strip()
        result = json.loads(cleaned)

        parsed_date = None
        try:
            parsed_date = datetime.strptime(str(result.get('date', '')), '%Y-%m-%d').date()
        except ValueError:
            parsed_date = fallback_date or argentina_now().date()

        event_type = result.get('event_type') or fallback_type
        if event_type not in EVENT_TYPES:
            event_type = fallback_type

        return {
            'date': parsed_date.isoformat(),
            'event_type': event_type,
            'title': str(result.get('title') or fallback['title'])[:180],
            'summary': str(result.get('summary') or fallback['summary'])[:2500],
            'patient_match': result.get('patient_match') if result.get('patient_match') in {'yes', 'no', 'unknown'} else 'unknown',
            'warnings': result.get('warnings') if isinstance(result.get('warnings'), list) else []
        }
    except Exception as exc:
        print('IMPORTADOR IA:', filename, str(exc))
        fallback['warnings'].append('No se pudo completar el análisis IA; revisar manualmente.')
        return fallback


@app.post('/patients/{patient_id}/history-import/analyze')
async def patient_history_import_analyze(
    patient_id: int,
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        return JSONResponse(
            status_code=404,
            content={
                'ok': False,
                'detail': 'Paciente no encontrado.',
                'items': []
            }
        )

    if supabase is None:
        return JSONResponse(
            status_code=500,
            content={
                'ok': False,
                'detail': 'Supabase Storage no configurado.',
                'items': []
            }
        )

    try:
        uploaded_files = []

        for uploaded in files:
            original_filename = uploaded.filename or 'archivo'
            safe_filename = _import_safe_filename(original_filename)
            content = await uploaded.read()

            mime_type = (
                uploaded.content_type
                or mimetypes.guess_type(safe_filename)[0]
                or 'application/octet-stream'
            )

            uploaded_files.append(
                (safe_filename, content, mime_type)
            )

        expanded_files = _import_expand_uploaded_files(uploaded_files)

        if not expanded_files:
            return JSONResponse(
                status_code=400,
                content={
                    'ok': False,
                    'detail': 'No se encontraron archivos compatibles.',
                    'items': []
                }
            )

        items = []
        failed_files = []

        for entry in expanded_files:
            if len(entry) == 4:
                filename, content, mime_type, preloaded_text = entry
            else:
                filename, content, mime_type = entry
                preloaded_text = ''

            safe_name = _import_safe_filename(filename)

            try:
                unique_name = f'{uuid.uuid4().hex}_{safe_name}'
                storage_path = (
                    f'patient_{patient.id}/imports/{unique_name}'
                )

                supabase.storage.from_('adjuntos').upload(
                    path=storage_path,
                    file=content,
                    file_options={
                        'content-type': (
                            mime_type
                            or 'application/octet-stream'
                        ),
                        'upsert': 'false'
                    }
                )

                text_value, images = _import_extract_content(
                    safe_name,
                    content,
                    mime_type,
                    preloaded_text
                )

                analysis = _import_ai_analysis(
                    patient,
                    safe_name,
                    text_value,
                    images
                )

                analysis.update({
                    'filename': safe_name,
                    'storage_path': storage_path,
                    'mime_type': (
                        mime_type
                        or 'application/octet-stream'
                    ),
                    'size_bytes': len(content),
                    'selected': (
                        analysis.get('patient_match') != 'no'
                    )
                })

                items.append(analysis)

            except Exception as file_error:
                print(
                    'ERROR IMPORTANDO ARCHIVO:',
                    safe_name,
                    str(file_error)
                )

                failed_files.append({
                    'filename': safe_name,
                    'error': str(file_error)
                })

        items.sort(
            key=lambda item: (
                item.get('date') or '',
                item.get('filename') or ''
            )
        )

        if not items and failed_files:
            return JSONResponse(
                status_code=400,
                content={
                    'ok': False,
                    'detail': (
                        'No se pudo procesar ninguno de los archivos. '
                        + ' | '.join(
                            f"{item['filename']}: {item['error']}"
                            for item in failed_files
                        )
                    ),
                    'items': [],
                    'failed_files': failed_files
                }
            )

        return JSONResponse({
            'ok': True,
            'count': len(items),
            'items': items,
            'failed_files': failed_files,
            'warnings': [
                (
                    f"{item['filename']}: no pudo procesarse."
                )
                for item in failed_files
            ]
        })

    except HTTPException as exc:
        return JSONResponse(
            status_code=exc.status_code,
            content={
                'ok': False,
                'detail': str(exc.detail),
                'items': []
            }
        )

    except Exception as exc:
        print('ERROR GENERAL IMPORTADOR:', str(exc))

        return JSONResponse(
            status_code=500,
            content={
                'ok': False,
                'detail': (
                    'Ocurrió un error al analizar los documentos: '
                    f'{str(exc)}'
                ),
                'items': []
            }
        )
@app.post('/patients/{patient_id}/history-import/confirm')
async def patient_history_import_confirm(
    patient_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail='Paciente no encontrado.')
    if supabase is None:
        raise HTTPException(status_code=500, detail='Supabase Storage no configurado.')

    payload = await request.json()
    items = payload.get('items', []) if isinstance(payload, dict) else []
    imported = 0
    skipped = 0
    errors = []

    for item in items:
        if not item.get('selected', True):
            skipped += 1
            continue

        filename = _import_safe_filename(str(item.get('filename') or 'archivo'))
        storage_path = str(item.get('storage_path') or '')
        expected_prefix = f'patient_{patient.id}/imports/'
        if not storage_path.startswith(expected_prefix):
            errors.append(f'{filename}: ubicación de archivo inválida.')
            continue

        try:
            event_day = datetime.strptime(str(item.get('date') or ''), '%Y-%m-%d').date()
        except ValueError:
            errors.append(f'{filename}: fecha inválida.')
            continue

        day_start = datetime.combine(event_day, datetime.min.time())
        day_end = day_start + timedelta(days=1)
        existing = (
            db.query(EventAttachment)
            .join(ClinicalEvent, EventAttachment.event_id == ClinicalEvent.id)
            .filter(ClinicalEvent.patient_id == patient.id)
            .filter(ClinicalEvent.event_date >= day_start)
            .filter(ClinicalEvent.event_date < day_end)
            .filter(EventAttachment.filename == filename)
            .first()
        )
        if existing:
            skipped += 1
            continue

        event_type = str(item.get('event_type') or 'Consulta clínica')
        if event_type not in EVENT_TYPES:
            event_type = 'Consulta clínica'

        title = str(item.get('title') or os.path.splitext(filename)[0])[:180]
        summary = str(item.get('summary') or 'Documento importado.')[:5000]
        event = ClinicalEvent(
            patient_id=patient.id,
            event_date=datetime.combine(event_day, datetime.min.time().replace(hour=12)),
            event_type=event_type,
            title=title,
            description=f'📥 Importado desde historia clínica externa\nArchivo original: {filename}\n\n{summary}',
            created_by=user.username
        )
        db.add(event)
        db.flush()

        public_url = supabase.storage.from_('adjuntos').get_public_url(storage_path)
        db.add(EventAttachment(event_id=event.id, filename=filename, file_path=public_url))
        imported += 1

    db.commit()
    return JSONResponse({
        'ok': len(errors) == 0,
        'imported': imported,
        'skipped': skipped,
        'errors': errors
    })
@app.post('/patients/{patient_id}/visits')
async def patient_visit_create(
    patient_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404)

    form = await request.form()

    def get(name, default=''):
        value = form.get(name)
        return value if value is not None else default

    def getlist(name):
        return form.getlist(name)

    def to_float(value):
        try:
            return float(str(value).replace(',', '.')) if value and str(value).strip() else None
        except ValueError:
            return None

    def to_int(value):
        try:
            return int(float(str(value).replace(',', '.'))) if value and str(value).strip() else None
        except ValueError:
            return None

    def parse_date(value):
        if value and str(value).strip():
            try:
                return datetime.strptime(str(value).strip(), '%Y-%m-%d').date()
            except ValueError:
                return None
        return None

    # Fecha elegida en "Fecha del evento"
    visit_date_value = get('visit_date')
    event_datetime = argentina_now()

    if visit_date_value and visit_date_value.strip():
        selected_date = parse_date(visit_date_value)

        if selected_date:
            event_datetime = datetime.combine(
                selected_date,
                argentina_now().time()
            )
    vaccine_name = get('vaccine_name')
    vaccine_lot = get('vaccine_lot')
    vaccine_expiration = get('vaccine_expiration')
    next_vaccine_name = get('next_vaccine_name')
    next_vaccine_date = get('next_vaccine_date')
    
    dewormer_product = get('dewormer_product')
    dewormer_drug = get('dewormer_drug')
    dewormer_dose = get('dewormer_dose')
    next_dewormer_product = get('next_dewormer_product')
    next_deworming_date = get('next_deworming_date')
    has_vaccine = any([
        vaccine_name.strip(),
        vaccine_lot.strip(),
        vaccine_expiration.strip(),
        next_vaccine_date.strip()
    ])

    has_deworming = any([
        dewormer_product.strip(),
        dewormer_drug.strip(),
        dewormer_dose.strip(),
        next_deworming_date.strip()
    ])

    description_parts = []

    if get('description').strip():
        description_parts.append(get('description').strip())

    if has_vaccine:
        description_parts.append(
            "Vacunación en esta visita:\n"
            f"Vacuna: {vaccine_name or '-'}\n"
            f"Lote: {vaccine_lot or '-'}\n"
            f"Vencimiento: {vaccine_expiration or '-'}\n"
            f"Próxima vacuna: {next_vaccine_date or '-'}"
        )

    if has_deworming:
        description_parts.append(
            "Desparasitación en esta visita:\n"
            f"Producto: {dewormer_product or '-'}\n"
            f"Droga: {dewormer_drug or '-'}\n"
            f"Dosis: {dewormer_dose or '-'}\n"
            f"Próxima desparasitación: {next_deworming_date or '-'}"
        )

    studies = getlist('studies')
    studies = [s for s in studies if s and str(s).strip()]
    study_map = [
        ('study_rx', 'Radiografía'),
        ('study_ecg', 'ECG'),
        ('study_eco', 'Ecocardiografía'),
        ('study_lab', 'Laboratorio'),
    ]
    for field_name, label in study_map:
        if get(field_name).strip() and label not in studies:
            studies.append(label)
    reminder_types = getlist('reminder_type') or getlist('reminder_type[]') or []
    reminder_titles = getlist('reminder_title') or getlist('reminder_title[]') or []
    reminder_dates = getlist('reminder_date') or getlist('reminder_date[]') or []
    if studies:
        description_parts.append(
            "Estudios solicitados / realizados:\n" + "\n".join(f"- {s}" for s in studies)
        )

    event = ClinicalEvent(
        patient_id=patient.id,
        event_date=event_datetime,
        event_type=get('event_type') or 'Consulta clínica',
        title=get('title') or 'Consulta / visita',
        description="\n\n".join(description_parts),
        anamnesis=get('anamnesis'),
        physical_exam=get('physical_exam'),
        diagnosis=get('diagnosis'),
        treatment=get('treatment'),

        weight=to_float(get('weight')),
        temperature=to_float(get('temperature')),
        heart_rate=to_int(get('heart_rate')),
        respiratory_rate=to_int(get('respiratory_rate')),
        mucous_membranes=get('mucous_membranes'),
        crt=get('crt'),
        hydration=get('hydration'),

        vaccine_name=vaccine_name if has_vaccine else '',
        vaccine_lot=vaccine_lot if has_vaccine else '',
        vaccine_expiration=vaccine_expiration if has_vaccine else '',
        next_vaccine_date=next_vaccine_date if has_vaccine else '',

        dewormer_product=dewormer_product if has_deworming else '',
        dewormer_drug=dewormer_drug if has_deworming else '',
        dewormer_dose=dewormer_dose if has_deworming else '',
        next_deworming_date=next_deworming_date if has_deworming else '',

        ecg_hr=get('ecg_hr'),
        ecg_rhythm=get('ecg_rhythm'),
        ecg_p=get('ecg_p'),
        ecg_pr=get('ecg_pr'),
        ecg_qrs=get('ecg_qrs'),
        ecg_st=get('ecg_st'),
        ecg_t=get('ecg_t'),
        ecg_qt=get('ecg_qt'),
        ecg_axis=get('ecg_axis'),
        ecg_interpretation=get('ecg_interpretation'),

        eco_aiao=get('eco_aiao'),
        eco_fs=get('eco_fs'),
        eco_acvim=get('eco_acvim'),
        eco_diagnosis=get('eco_diagnosis'),
        eco_treatment=get('eco_treatment'),

        rx_vhs=get('rx_vhs'),
        rx_vlas=get('rx_vlas'),
        rx_lung_pattern=get('rx_lung_pattern'),
        rx_edema=get('rx_edema'),
        rx_congestion=get('rx_congestion'),
        rx_observations=get('rx_observations'),

        created_by=user.username
    )

    db.add(event)
    db.flush()

    # Resultados de laboratorio revisados en la vista previa.
    lab_results_raw = get('lab_results_json')
    if lab_results_raw and str(lab_results_raw).strip():
        try:
            lab_payload = json.loads(lab_results_raw)
        except json.JSONDecodeError:
            lab_payload = {}

        sample_date = parse_date(lab_payload.get('sample_date', '')) or event_datetime.date()
        source_files = ', '.join(lab_payload.get('filenames', []) or [])[:255]

        for row in lab_payload.get('results', []) or []:
            analyte = str(row.get('analyte') or '').strip()
            if not analyte:
                continue
            db.add(LaboratoryResult(
                event_id=event.id,
                patient_id=patient.id,
                sample_date=sample_date,
                panel=str(row.get('panel') or 'Otro')[:120],
                analyte=analyte[:180],
                normalized_analyte=_lab_normalize_analyte(analyte)[:180],
                value_text=str(row.get('value_text') or '')[:120],
                value_numeric=_lab_float(row.get('value_numeric')),
                unit=str(row.get('unit') or '')[:80],
                reference_low=_lab_float(row.get('reference_low')),
                reference_high=_lab_float(row.get('reference_high')),
                reference_text=str(row.get('reference_text') or '')[:120],
                flag=str(row.get('flag') or 'unknown')[:20],
                source_filename=source_files
            ))

    if get('weight') and str(get('weight')).strip():
        patient.weight = to_float(get('weight'))

    attachments = getlist('attachments')

    for file in attachments:
        if not file or not getattr(file, 'filename', ''):
            continue

        if supabase is None:
            raise HTTPException(
                status_code=500,
                detail="Supabase Storage no configurado."
            )

        original_name = os.path.basename(file.filename)
        safe_name = safe_storage_filename(original_name)
        unique_name = f"{uuid.uuid4().hex}_{safe_name}"

        content = await file.read()

        if not content:
            continue

        content_type = file.content_type or mimetypes.guess_type(original_name)[0] or "application/octet-stream"
        storage_path = f"patient_{patient.id}/event_{event.id}/{unique_name}"

        supabase.storage.from_("adjuntos").upload(
            path=storage_path,
            file=content,
            file_options={
                "content-type": content_type,
                "upsert": "false"
            }
        )

        public_url = supabase.storage.from_("adjuntos").get_public_url(storage_path)

        db.add(EventAttachment(
            event_id=event.id,
            filename=original_name,
            file_path=public_url
        ))
    close_managed_due_events(db, patient, event, user)
    if next_vaccine_date and next_vaccine_date.strip():
        db.add(ClinicalEvent(
            patient_id=patient.id,
            event_date=event_datetime,
            event_type='Vacuna',
            title=next_vaccine_name or 'Vacuna pendiente',
            description=(
                f"{DUE_ACTIVE_MARKER}\n"
                f"Recordatorio creado desde visita #{event.id}\n"
                f"Vacuna aplicada hoy: {vaccine_name or '-'}\n"
                f"Próxima vacuna: {next_vaccine_name or '-'}"
            ),
            reminder_date=parse_date(next_vaccine_date),
            created_by=user.username
        ))
    
    if next_deworming_date and next_deworming_date.strip():
        db.add(ClinicalEvent(
            patient_id=patient.id,
            event_date=event_datetime,
            event_type='Desparasitación',
            title=next_dewormer_product or 'Desparasitación pendiente',
            description=(
                f"{DUE_ACTIVE_MARKER}\n"
                f"Recordatorio creado desde visita #{event.id}\n"
                f"Desparasitación aplicada hoy: {dewormer_product or '-'}\n"
                f"Próximo desparasitario: {next_dewormer_product or '-'}"
            ),
            reminder_date=parse_date(next_deworming_date),
            created_by=user.username
        ))
    for r_type, r_title, r_date in zip(reminder_types, reminder_titles, reminder_dates):
        r_date_parsed = parse_date(r_date)

        if not r_date_parsed:
            continue

        r_type = r_type or 'Control'
        r_title = r_title or r_type

        reminder_event = ClinicalEvent(
            patient_id=patient.id,
            event_date=event_datetime,
            event_type=r_type,
            title=r_title,
            description=(
                f"{DUE_ACTIVE_MARKER}\n"
                f"Recordatorio creado desde visita #{event.id}\n"
                f"Tipo: {r_type}\n"
                f"Título: {r_title}"
            ),
            reminder_date=r_date_parsed,
            created_by=user.username
        )

        db.add(reminder_event)
    active_waitlist_entry_id = request.session.pop('active_waitlist_entry_id', None)

    active_waiting_entries = []

    if active_waitlist_entry_id:
        active_entry = db.get(WaitingListEntry, int(active_waitlist_entry_id))
        if active_entry and active_entry.status in ['Esperando', 'En consulta']:
            active_waiting_entries.append(active_entry)

    if not active_waiting_entries:
        active_waiting_entries = (
            db.query(WaitingListEntry)
            .filter(WaitingListEntry.patient_id == patient.id)
            .filter(WaitingListEntry.status.in_(['Esperando', 'En consulta']))
            .all()
        )

    for waiting_entry in active_waiting_entries:
        waiting_entry.status = 'Finalizado'
        waiting_entry.finished_at = argentina_now()

    db.commit()

    return RedirectResponse(
        f'/patients/{patient.id}/v2',
        status_code=303
    )
@app.get('/patients/{patient_id}', response_class=HTMLResponse)
def patient_detail(request: Request, patient_id: int, db: Session = Depends(get_db), user: User = Depends(require_user)):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(404)
    return RedirectResponse(
        f'/patients/{patient.id}/v2',
        status_code=303
    )
    events = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .filter(~ClinicalEvent.description.ilike(f'%{DUE_ACTIVE_MARKER}%'))
        .filter(~ClinicalEvent.description.ilike(f'%{DUE_CLOSED_MARKER}%'))
        .order_by(ClinicalEvent.event_date.desc())
        .limit(20)
        .all()
    )

    today = argentina_now().date()
    pending_actions = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .filter(ClinicalEvent.description.ilike(f'%{DUE_ACTIVE_MARKER}%'))
        .order_by(ClinicalEvent.event_date.desc())
        .all()
    )
    upcoming_events = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id,
            ClinicalEvent.reminder_date != None,
            ClinicalEvent.reminder_date >= today
        )
        .order_by(ClinicalEvent.reminder_date.asc())
        .limit(5)
        .all()
    )
    
    clinical_ai = {}    
    
    last_anesthesia = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id,
            ClinicalEvent.event_type == 'Anestesia'
        )
        .order_by(ClinicalEvent.event_date.desc())
        .first()
    )

    anesthesia_history = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id,
            ClinicalEvent.event_type == 'Anestesia'
        )
        .order_by(ClinicalEvent.event_date.desc())
        .all()
    )
    cardiology_ecgs = (
    db.query(ClinicalEvent)
    .filter(
        ClinicalEvent.patient_id == patient.id,
        ClinicalEvent.event_type == "ECG"
    )
    .order_by(ClinicalEvent.event_date.desc())
    .all()
)

    last_ecg = cardiology_ecgs[0] if cardiology_ecgs else None
    ecg_count = len(cardiology_ecgs)

    cardiology_ecos = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id,
            ClinicalEvent.event_type == "Ecocardiografía"
        )
        .order_by(ClinicalEvent.event_date.desc())
        .all()
    )

    last_eco = cardiology_ecos[0] if cardiology_ecos else None
    eco_count = len(cardiology_ecos)

    next_visit = upcoming_events[0] if upcoming_events else None
    vaccine_names = sorted(
        list(
            {
                e.vaccine_name.strip()
                for e in db.query(ClinicalEvent).all()
                if e.vaccine_name and e.vaccine_name.strip()
            }
        )
    )

    dewormer_names = sorted(
        list(
            {
                e.dewormer_product.strip()
                for e in db.query(ClinicalEvent).all()
                if e.dewormer_product and e.dewormer_product.strip()
            }
        )
    )
    
    active_hospitalization = (
        db.query(Hospitalization)
        .filter(
            Hospitalization.patient_id == patient.id,
            Hospitalization.status == 'Internado'
        )
        .order_by(Hospitalization.admission_date.desc())
        .first()
    )
    return templates.TemplateResponse(
        'patient_detail.html',
        {
            'request': request,
            'patient': patient,
            'today': today,
            'events': events,
            'clinical_ai': clinical_ai,
            'event_types': EVENT_TYPES,
            'upcoming_events': upcoming_events,
            'pending_actions': pending_actions,
            'timedelta': timedelta,
            'last_anesthesia': last_anesthesia,
            'anesthesia_history': anesthesia_history,
            'next_visit': next_visit,
            'last_ecg': last_ecg,
            'ecg_count': ecg_count,
            'last_eco': last_eco,
            'eco_count': eco_count,
            "vaccine_names": vaccine_names,
            "dewormer_names": dewormer_names,
            "active_hospitalization": active_hospitalization
        }
    )
@app.get('/patients/{patient_id}/events/{event_id}/edit', response_class=HTMLResponse)
def edit_clinical_event_form(
    request: Request,
    patient_id: int,
    event_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    event = db.get(ClinicalEvent, event_id)

    if not patient or not event or event.patient_id != patient.id:
        raise HTTPException(status_code=404)

    return templates.TemplateResponse(
        'patient_event_edit.html',
        {
            'request': request,
            'patient': patient,
            'event': event,
            'event_types': EVENT_TYPES,
        }
    )


@app.post('/patients/{patient_id}/events/{event_id}/edit')
async def edit_clinical_event_save(
    patient_id: int,
    event_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    event = db.get(ClinicalEvent, event_id)

    if not patient or not event or event.patient_id != patient.id:
        raise HTTPException(status_code=404)

    form = await request.form()

    fields = [
        'event_type', 'title', 'description', 'anamnesis', 'physical_exam',
        'diagnosis', 'treatment',
        'weight', 'temperature', 'heart_rate', 'respiratory_rate',
        'mucous_membranes', 'crt', 'hydration',

        'ecg_hr', 'ecg_rhythm', 'ecg_p', 'ecg_pr', 'ecg_qrs',
        'ecg_st', 'ecg_t', 'ecg_qt', 'ecg_axis', 'ecg_interpretation',

        'ecg_p_mv', 'ecg_qrs_mv', 'ecg_t_mv', 'ecg_qtc',
        'ecg_polarity', 'ecg_arrhythmia', 'ecg_conduction', 'ecg_notes',

        'eco_aiao', 'eco_fs', 'eco_acvim', 'eco_diagnosis', 'eco_treatment',

        'eco_epss', 'eco_lvidd', 'eco_lvids', 'eco_ivsd', 'eco_ivss',
        'eco_lvpwd', 'eco_lvpws', 'eco_fe',
        'eco_la_size', 'eco_lv_size', 'eco_rv_size', 'eco_ra_size',
        'eco_mitral', 'eco_tricuspid', 'eco_aortic', 'eco_pulmonary',
        'eco_pulmonary_htn', 'eco_pericardium',
        'eco_doppler', 'eco_observations',

        'rx_vhs', 'rx_vlas', 'rx_heart_size', 'rx_left_atrium',
        'rx_left_heart', 'rx_right_heart',
        'rx_pulmonary_vessels', 'rx_lung_pattern',
        'rx_edema', 'rx_congestion', 'rx_trachea',
        'rx_observations',

        'vaccine_name', 'vaccine_lot', 'vaccine_expiration', 'next_vaccine_date',
        'dewormer_product', 'dewormer_drug', 'dewormer_dose', 'next_deworming_date',
        'reminder_date'
    ]
    for field in fields:
        if hasattr(event, field):
            value = form.get(field)

            if value == '':
                value = None

            if field in ['diagnosis', 'treatment', 'description', 'anamnesis', 'physical_exam']:
                value = value or ''

            if field in ['next_vaccine_date', 'next_deworming_date', 'reminder_date'] and value:
                value = datetime.strptime(value, '%Y-%m-%d').date()

            setattr(event, field, value)
    attachments = form.getlist("attachments")
    upload_dir = "/opt/render/project/src/app/uploads"
    os.makedirs(upload_dir, exist_ok=True)

    for file in attachments:
        if file and file.filename:
            safe_filename = os.path.basename(file.filename)
            full_path = os.path.join(upload_dir, safe_filename)

            with open(full_path, "wb") as buffer:
                buffer.write(await file.read())

            attachment = EventAttachment(
                event_id=event.id,
                filename=safe_filename,
                file_path="/uploads/" + safe_filename
            )

            db.add(attachment)

  
    db.commit()

    return RedirectResponse(
        url=f'/patients/{patient_id}',
        status_code=303
    )
@app.post('/patients/{patient_id}/events/{event_id}/attachments/{attachment_id}/delete')
def delete_event_attachment_from_edit(
    patient_id: int,
    event_id: int,
    attachment_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    attachment = db.get(EventAttachment, attachment_id)

    if not attachment:
        raise HTTPException(status_code=404)

    db.delete(attachment)
    db.commit()

    return RedirectResponse(
        url=f'/patients/{patient_id}/events/{event_id}/edit',
        status_code=303
    )
@app.get('/patients/{patient_id}/cardiology/ecg', response_class=HTMLResponse)
def patient_ecg_list(
    request: Request,
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404)

    ecg_events = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id,
            ClinicalEvent.event_type == "ECG"
        )
        .order_by(ClinicalEvent.event_date.desc())
        .all()
    )

    return templates.TemplateResponse(
        'patient_ecg_list.html',
        {
            'request': request,
            'patient': patient,
            'ecg_events': ecg_events
        }
    )
def build_cardio_ai(last_ecg, previous_ecg, last_eco, previous_eco, last_rx, previous_rx):
    def to_float(value):
        try:
            if value is None:
                return None
            value = str(value).replace(",", ".").strip()
            if not value:
                return None
            return float(value)
        except Exception:
            return None

    score = 100
    alerts = []
    sections = []
    recommendations = []

    hr = to_float(getattr(last_ecg, "ecg_hr", None)) if last_ecg else None
    axis = to_float(getattr(last_ecg, "ecg_axis", None)) if last_ecg else None
    rhythm = ((getattr(last_ecg, "ecg_rhythm", "") or "") + " " + (getattr(last_ecg, "ecg_arrhythmia", "") or "")).lower() if last_ecg else ""

    aiao = to_float(getattr(last_eco, "eco_aiao", None)) if last_eco else None
    fs = to_float(getattr(last_eco, "eco_fs", None)) if last_eco else None
    epss = to_float(getattr(last_eco, "eco_epss", None)) if last_eco else None
    fe = to_float(getattr(last_eco, "eco_fe", None)) if last_eco else None
    acvim = (getattr(last_eco, "eco_acvim", "") or "").upper() if last_eco else ""

    vhs = to_float(getattr(last_rx, "rx_vhs", None)) if last_rx else None
    vlas = to_float(getattr(last_rx, "rx_vlas", None)) if last_rx else None
    rx_text = " ".join([
        getattr(last_rx, "rx_edema", "") or "",
        getattr(last_rx, "rx_congestion", "") or "",
        getattr(last_rx, "rx_pulmonary_vessels", "") or "",
        getattr(last_rx, "rx_lung_pattern", "") or "",
        getattr(last_rx, "rx_observations", "") or "",
        getattr(last_rx, "description", "") or "",
    ]).lower() if last_rx else ""

    if last_ecg:
        ecg_notes = []

        if hr is not None:
            if hr > 180:
                score -= 10
                ecg_notes.append(f"FC elevada ({hr:g} lpm).")
                alerts.append("FC elevada en el último ECG.")
            elif hr < 60:
                score -= 10
                ecg_notes.append(f"FC baja ({hr:g} lpm).")
                alerts.append("FC baja en el último ECG.")
            else:
                ecg_notes.append(f"FC dentro de rango clínico orientativo ({hr:g} lpm).")

        if "fibril" in rhythm or "bloqueo" in rhythm or "taqui" in rhythm or "extra" in rhythm:
            score -= 12
            ecg_notes.append("Se describe arritmia o alteración de conducción.")
            alerts.append("Revisar arritmia/conducción descrita en ECG.")
        elif rhythm.strip():
            ecg_notes.append("Ritmo sin alteraciones mayores descritas.")

        if axis is not None:
            if axis < 40 or axis > 100:
                score -= 6
                ecg_notes.append(f"Eje eléctrico fuera del rango esperado ({axis:g}°).")
            else:
                ecg_notes.append(f"Eje eléctrico conservado ({axis:g}°).")

        sections.append({"title": "ECG", "items": ecg_notes or ["ECG cargado sin datos suficientes para interpretación automática."]})
    else:
        score -= 8
        sections.append({"title": "ECG", "items": ["Sin ECG cargado."]})

    if last_eco:
        eco_notes = []

        if aiao is not None:
            if aiao >= 1.9:
                score -= 18
                eco_notes.append(f"AI/Ao aumentado de forma importante ({aiao:g}).")
                alerts.append("AI/Ao elevado: sugiere dilatación auricular izquierda.")
            elif aiao >= 1.6:
                score -= 10
                eco_notes.append(f"AI/Ao aumentado ({aiao:g}).")
            else:
                eco_notes.append(f"AI/Ao dentro de rango orientativo ({aiao:g}).")

        if fs is not None:
            if fs < 20:
                score -= 18
                eco_notes.append(f"FS baja ({fs:g}%). Posible disfunción sistólica.")
                alerts.append("FS baja: revisar función sistólica.")
            elif fs > 45:
                eco_notes.append(f"FS aumentada/hiperdinámica ({fs:g}%).")
            else:
                eco_notes.append(f"FS conservada ({fs:g}%).")

        if epss is not None:
            if epss > 7:
                score -= 8
                eco_notes.append(f"EPSS aumentado ({epss:g} mm).")
            else:
                eco_notes.append(f"EPSS sin aumento relevante ({epss:g} mm).")

        if fe is not None:
            if fe < 45:
                score -= 12
                eco_notes.append(f"FE disminuida ({fe:g}%).")
            else:
                eco_notes.append(f"FE conservada ({fe:g}%).")

        if acvim:
            if acvim == "B2":
                score -= 12
                eco_notes.append("Clasificación ACVIM B2: requiere seguimiento cardiológico.")
            elif acvim in ["C", "D"]:
                score -= 25
                eco_notes.append(f"Clasificación ACVIM {acvim}: paciente de control estricto.")
                alerts.append(f"ACVIM {acvim}: riesgo cardiológico elevado.")

        sections.append({"title": "Ecocardiografía", "items": eco_notes or ["Eco cargada sin datos suficientes para interpretación automática."]})
    else:
        score -= 8
        sections.append({"title": "Ecocardiografía", "items": ["Sin ecocardiografía cargada."]})

    if last_rx:
        rx_notes = []

        if vhs is not None:
            if vhs > 11.5:
                score -= 10
                rx_notes.append(f"VHS aumentado ({vhs:g}). Compatible con cardiomegalia.")
                alerts.append("VHS elevado en RX.")
            else:
                rx_notes.append(f"VHS sin aumento relevante ({vhs:g}).")

        if vlas is not None:
            if vlas > 3:
                score -= 8
                rx_notes.append(f"VLAS aumentado ({vlas:g}). Sugiere aumento auricular izquierdo.")
            else:
                rx_notes.append(f"VLAS sin aumento relevante ({vlas:g}).")

        if "edema" in rx_text:
            score -= 20
            rx_notes.append("Se menciona edema pulmonar.")
            alerts.append("RX con mención de edema pulmonar.")
        if "congest" in rx_text:
            score -= 12
            rx_notes.append("Se menciona congestión vascular/pulmonar.")
        if "normal" in rx_text and not rx_notes:
            rx_notes.append("RX descrita sin alteraciones cardiopulmonares relevantes.")

        sections.append({"title": "Radiografía", "items": rx_notes or ["RX cargada sin datos suficientes para interpretación automática."]})
    else:
        sections.append({"title": "Radiografía", "items": ["Sin RX cardíaca cargada."]})

    if aiao is not None and vhs is not None:
        if aiao >= 1.6 and vhs <= 11.5:
            alerts.append("Eco sugiere AI aumentada, pero VHS no acompaña: revisar concordancia RX/Eco.")
        if aiao < 1.6 and vhs > 11.5:
            alerts.append("VHS elevado con AI/Ao normal: revisar otras causas de cardiomegalia.")

    if acvim in ["C", "D"] or "edema" in rx_text:
        recommendations.append("Control cardiológico estricto y reevaluar signos de insuficiencia cardíaca congestiva.")
    elif acvim == "B2" or (aiao is not None and aiao >= 1.6):
        recommendations.append("Seguimiento cardiológico periódico y control ecocardiográfico.")
    else:
        recommendations.append("Continuar controles periódicos según evolución clínica.")

    if not last_ecg:
        recommendations.append("Cargar ECG para completar interpretación eléctrica.")
    if not last_eco:
        recommendations.append("Cargar ecocardiografía para estadificación cardiológica.")
    if not last_rx:
        recommendations.append("Cargar RX torácica/cardiológica si hay tos, disnea, soplo o sospecha de congestión.")

    score = max(0, min(100, int(score)))

    if score >= 85:
        label = "🟢 Estable"
        conclusion = "Paciente cardiológicamente estable con los datos disponibles."
    elif score >= 60:
        label = "🟡 Requiere seguimiento"
        conclusion = "Paciente con hallazgos que justifican seguimiento cardiológico."
    elif score >= 40:
        label = "🟠 Riesgo moderado"
        conclusion = "Paciente con alteraciones cardiológicas relevantes; conviene control cercano."
    else:
        label = "🔴 Alto riesgo"
        conclusion = "Paciente con indicadores de alto riesgo cardiológico o posible descompensación."

    return {
        "score": score,
        "label": label,
        "conclusion": conclusion,
        "sections": sections,
        "alerts": alerts,
        "recommendations": recommendations,
    }
@app.get('/patients/{patient_id}/cardiology', response_class=HTMLResponse)
def patient_cardiology(
    request: Request,
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404)

    cardiology_events = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id,
            ClinicalEvent.event_type.in_(
                [
                    "ECG",
                    "Ecocardiografía",
                    "Radiografía"
                ]
            )
        )
        .order_by(ClinicalEvent.event_date.desc())
        .all()
    )

    ecg_events = [
        e for e in cardiology_events
        if e.event_type == "ECG"
    ]

    eco_events = [
        e for e in cardiology_events
        if e.event_type == "Ecocardiografía"
    ]

    rx_events = [
        e for e in cardiology_events
        if e.event_type == "Radiografía"
    ]

    last_ecg = ecg_events[0] if ecg_events else None
    previous_ecg = ecg_events[1] if len(ecg_events) > 1 else None

    last_eco = eco_events[0] if eco_events else None
    previous_eco = eco_events[1] if len(eco_events) > 1 else None

    last_rx = rx_events[0] if rx_events else None
    previous_rx = rx_events[1] if len(rx_events) > 1 else None

    def to_float(value):
        try:
            if value is None:
                return None

            value = str(value).replace(",", ".").strip()

            if value == "":
                return None

            return float(value)

        except Exception:
            return None

    def delta(current, previous):

        current_value = to_float(current)
        previous_value = to_float(previous)

        if current_value is None:
            return None

        if previous_value is None:
            return None

        return round(current_value - previous_value, 2)

    def trend_width(value, maximum):

        number = to_float(value)

        if number is None:
            return 0

        width = int((number / maximum) * 100)

        if width < 4:
            width = 4

        if width > 100:
            width = 100

        return width

    fc_delta = delta(
        last_ecg.ecg_hr if last_ecg else None,
        previous_ecg.ecg_hr if previous_ecg else None
    )

    axis_delta = delta(
        last_ecg.ecg_axis if last_ecg else None,
        previous_ecg.ecg_axis if previous_ecg else None
    )

    aiao_delta = delta(
        last_eco.eco_aiao if last_eco else None,
        previous_eco.eco_aiao if previous_eco else None
    )

    fs_delta = delta(
        last_eco.eco_fs if last_eco else None,
        previous_eco.eco_fs if previous_eco else None
    )
    acvim = (last_eco.eco_acvim if last_eco else "") or ""

    cardio_status = {
        "label": "Estable",
        "class": "ok",
        "icon": "🟢"
    }

    if acvim in ["B2"]:
        cardio_status = {
            "label": "En seguimiento",
            "class": "warn",
            "icon": "🟡"
        }

    if acvim in ["C", "D"]:
        cardio_status = {
            "label": "Control estricto",
            "class": "danger",
            "icon": "🔴"
        }

    active_treatment = ""

    if last_eco:
        active_treatment = (
            getattr(last_eco, "eco_treatment", "")
            or ""
        )

    if not active_treatment and last_ecg:
        active_treatment = (
            getattr(last_ecg, "treatment", "")
            or ""
        )

    if not active_treatment:
        active_treatment = "Sin tratamiento registrado"

    next_control = None

    if last_eco:
        stage = (last_eco.eco_acvim or "").upper()

        if stage in ["C", "D"]:
            next_control = 30

        elif stage == "B2":
            next_control = 90

        else:
            next_control = 180

    elif last_ecg:
        next_control = 180

    else:
        next_control = None
    fc_trend = [
        {
            "label": e.event_date.strftime("%d/%m"),
            "value": e.ecg_hr,
            "width": trend_width(e.ecg_hr, 220),
            "event_id": e.id
        }
        for e in reversed(ecg_events[:8])
        if e.ecg_hr
    ]

    aiao_trend = [
        {
            "label": e.event_date.strftime("%d/%m"),
            "value": e.eco_aiao,
            "width": trend_width(e.eco_aiao, 3),
            "event_id": e.id
        }
        for e in reversed(eco_events[:8])
        if e.eco_aiao
    ]

    fs_trend = [
        {
            "label": e.event_date.strftime("%d/%m"),
            "value": e.eco_fs,
            "width": trend_width(e.eco_fs, 60),
            "event_id": e.id
        }
        for e in reversed(eco_events[:8])
        if e.eco_fs
    ]

    comparison_notes = []

    if fc_delta is not None:

        if fc_delta < 0:
            comparison_notes.append(
                f"FC disminuyó {abs(fc_delta):g} lpm respecto al ECG previo."
            )

        elif fc_delta > 0:
            comparison_notes.append(
                f"FC aumentó {fc_delta:g} lpm respecto al ECG previo."
            )

        else:
            comparison_notes.append(
                "FC sin cambios respecto al ECG previo."
            )

    if axis_delta is not None:

        if abs(axis_delta) <= 10:
            comparison_notes.append(
                "Eje eléctrico sin cambios clínicamente relevantes."
            )

        elif axis_delta > 0:
            comparison_notes.append(
                f"Eje eléctrico aumentó {axis_delta:g}°."
            )

        else:
            comparison_notes.append(
                f"Eje eléctrico disminuyó {abs(axis_delta):g}°."
            )

    if aiao_delta is not None:

        if abs(aiao_delta) < 0.05:
            comparison_notes.append(
                "AI/Ao sin variación significativa."
            )

        elif aiao_delta > 0:
            comparison_notes.append(
                f"AI/Ao aumentó {aiao_delta:g}."
            )

        else:
            comparison_notes.append(
                f"AI/Ao disminuyó {abs(aiao_delta):g}."
            )

    if fs_delta is not None:

        if fs_delta > 0:
            comparison_notes.append(
                f"FS aumentó {fs_delta:g}%."
            )

        elif fs_delta < 0:
            comparison_notes.append(
                f"FS disminuyó {abs(fs_delta):g}%."
            )

        else:
            comparison_notes.append(
                "FS sin cambios."
            )

    if not comparison_notes:

        comparison_notes.append(
            "Todavía no hay controles suficientes para comparar la evolución."
        )

    cardio_ai = build_cardio_ai(
        last_ecg,
        previous_ecg,
        last_eco,
        previous_eco,
        last_rx,
        previous_rx
    )

    return templates.TemplateResponse(
        "patient_cardiology.html",
        {
            "request": request,
            "patient": patient,

            "cardiology_events": cardiology_events,

            "ecg_events": ecg_events,
            "eco_events": eco_events,
            "rx_events": rx_events,

            "last_ecg": last_ecg,
            "previous_ecg": previous_ecg,

            "last_eco": last_eco,
            "previous_eco": previous_eco,

            "last_rx": last_rx,
            "previous_rx": previous_rx,

            "fc_delta": fc_delta,
            "axis_delta": axis_delta,
            "aiao_delta": aiao_delta,
            "fs_delta": fs_delta,

            "cardio_status": cardio_status,

            "active_treatment": active_treatment,

            "next_control": next_control,

            "fc_trend": fc_trend,
            "aiao_trend": aiao_trend,
            "fs_trend": fs_trend,

            "comparison_notes": comparison_notes,

            "cardio_score": cardio_ai["score"],
            "cardio_score_label": cardio_ai["label"],
            "cardio_ai_conclusion": cardio_ai["conclusion"],
            "cardio_ai_sections": cardio_ai["sections"],
            "cardio_ai_alerts": cardio_ai["alerts"],
            "cardio_ai_recommendations": cardio_ai["recommendations"],

            "today": argentina_now().date()
        }
    )
@app.get('/patients/{patient_id}/cardiology/print', response_class=HTMLResponse)
def patient_cardiology_print(
    request: Request,
    patient_id: int,
    event_id: Optional[int] = None,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(
            status_code=404,
            detail="Paciente no encontrado"
        )

    ecg_query = db.query(ClinicalEvent).filter(
        ClinicalEvent.patient_id == patient.id,
        ClinicalEvent.event_type == "ECG"
    )

    if event_id is not None:
        ecg = ecg_query.filter(
            ClinicalEvent.id == event_id
        ).first()
    else:
        ecg = (
            ecg_query
            .order_by(ClinicalEvent.event_date.desc())
            .first()
        )

    if not ecg:
        raise HTTPException(
            status_code=404,
            detail="No se encontró el electrocardiograma solicitado"
        )

    return templates.TemplateResponse(
        "patient_cardiology_print.html",
        {
            "request": request,
            "patient": patient,
            "ecg": ecg,
            "today": argentina_now().date()
        }
    )    
@app.get('/patients/{patient_id}/edit', response_class=HTMLResponse)
def patient_edit_form(request: Request, patient_id: int, db: Session = Depends(get_db), user: User = Depends(require_user)):
    patient = db.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404)
    owners = db.query(Owner).order_by(Owner.name).limit(500).all()
    return templates.TemplateResponse(
        'patient_edit.html',
        {'request': request, 'patient': patient, 'owners': owners}
    )


@app.post('/patients/{patient_id}/edit')
def patient_edit_save(
    patient_id: int,
    name: str = Form(...),
    owner_id: int = Form(...),
    species: str = Form(''),
    breed: str = Form(''),
    sex: str = Form(''),
    birth_date: str = Form(''),
    weight: str = Form(''),
    alerts: str = Form(''),
    notes: str = Form(''),
    owner_name: str = Form(''),
    owner_phone: str = Form(''),
    owner_whatsapp: str = Form(''),
    owner_email: str = Form(''),
    owner_address: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404)

    patient.name = name
    patient.owner_id = owner_id
    patient.species = species
    patient.breed = breed
    patient.sex = sex
    patient.birth_date = datetime.strptime(birth_date, '%Y-%m-%d').date() if birth_date.strip() else None
    patient.weight = float(weight.replace(',', '.')) if weight.strip() else None
    patient.alerts = alerts
    patient.notes = notes
    owner = db.get(Owner, owner_id)

    if owner:
        owner.name = owner_name
        owner.phone = owner_phone
        owner.whatsapp = owner_whatsapp
        owner.email = owner_email
        owner.address = owner_address

    db.commit()
    return RedirectResponse(f'/patients/{patient.id}', status_code=303)
@app.get('/patients/{patient_id}/events/{event_id}', response_class=HTMLResponse)
def event_detail(
    request: Request,
    patient_id: int,
    event_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    event = db.get(ClinicalEvent, event_id)
    previous_ecg = None

    if event and event.event_type == "ECG":
        previous_ecg = (
            db.query(ClinicalEvent)
            .filter(
                ClinicalEvent.patient_id == patient_id,
                ClinicalEvent.event_type == "ECG",
                ClinicalEvent.id != event.id,
                ClinicalEvent.event_date < event.event_date
            )
            .order_by(ClinicalEvent.event_date.desc())
            .first()
        )
    if not patient or not event:
        raise HTTPException(status_code=404)

    return templates.TemplateResponse(
        'event_detail.html',
{
    'request': request,
    'patient': patient,
    'event': event,
    'previous_ecg': previous_ecg,
    'clinical_ai': None
}
    )
@app.post('/patients/{patient_id}/events/{event_id}/ai')
def event_ai_analyze(
    patient_id: int,
    event_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    event = db.get(ClinicalEvent, event_id)

    if not patient or not event or event.patient_id != patient.id:
        raise HTTPException(status_code=404)

    result = ai_clinical_summary(event)

    return JSONResponse(result)
@app.post('/patients/{patient_id}/events/{event_id}/ai/save')
def event_ai_save_to_event(
    patient_id: int,
    event_id: int,
    ai_text: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    event = db.get(ClinicalEvent, event_id)

    if not patient or not event or event.patient_id != patient.id:
        raise HTTPException(status_code=404)

    if not ai_text.strip():
        return JSONResponse({
            "ok": False,
            "message": "No hay análisis IA para guardar."
        })

    separator = "\n\n────────────────────────\n\n"
    block = "🧠 ANÁLISIS IA\n\n" + ai_text.strip()

    current_description = event.description or ""

    if "🧠 ANÁLISIS IA" not in current_description:
        event.description = current_description + separator + block
    else:
        event.description = current_description + "\n\n" + block

    db.commit()

    return JSONResponse({
        "ok": True,
        "message": "Análisis IA agregado al evento."
    })
@app.post('/patients/{patient_id}/events/{event_id}/delete')
def event_delete(
    patient_id: int,
    event_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    event = db.get(ClinicalEvent, event_id)

    if not patient or not event or event.patient_id != patient.id:
        raise HTTPException(status_code=404)

    db.delete(event)
    db.commit()

    return RedirectResponse(f'/patients/{patient.id}', status_code=303)
@app.post('/patients/{patient_id}/events')
def event_create(
    request: Request,
    patient_id: int,
    event_type: str = Form(...),
    title: str = Form(''),
    description: str = Form(''),
    anamnesis: str = Form(''),
    physical_exam: str = Form(''),
    diagnosis: str = Form(''),
    treatment: str = Form(''),
    vaccine_name: str = Form(''),
    vaccine_lot: str = Form(''),
    vaccine_expiration: str = Form(''),
    next_vaccine_date: str = Form(''),
    dewormer_product: str = Form(''),
    dewormer_drug: str = Form(''),
    dewormer_dose: str = Form(''),
    next_deworming_date: str = Form(''),
    reminder_date: str = Form(''),
    reminder_type: str = Form(''),
    reminder_vaccine_name: str = Form(''),
    reminder_dewormer_name: str = Form(''),
    reminder_notes: str = Form(''),
    event_date: str = Form(''),
    weight: str = Form(''),
    temperature: str = Form(''),
    heart_rate: str = Form(''),
    respiratory_rate: str = Form(''),
    mucous_membranes: str = Form(''),
    crt: str = Form(''),
    hydration: str = Form(''),
    ecg_hr: str = Form(''),
    ecg_rhythm: str = Form(''),

    ecg_p: str = Form(''),
    ecg_pr: str = Form(''),

    ecg_qrs: str = Form(''),
    ecg_st: str = Form(''),

    ecg_t: str = Form(''),
    ecg_qt: str = Form(''),

    ecg_axis: str = Form(''),
    ecg_interpretation: str = Form(''),
    eco_aiao: str = Form(''),
    eco_fs: str = Form(''),
    eco_acvim: str = Form(''),

    eco_diagnosis: str = Form(''),
    eco_treatment: str = Form(''),

    # ECG ampliado
    ecg_p_mv: str = Form(''),
    ecg_qrs_mv: str = Form(''),
    ecg_t_mv: str = Form(''),
    ecg_qtc: str = Form(''),
    ecg_polarity: str = Form(''),
    ecg_arrhythmia: str = Form(''),
    ecg_conduction: str = Form(''),
    ecg_notes: str = Form(''),

    # Ecocardiografía ampliada
    eco_epss: str = Form(''),
    eco_lvidd: str = Form(''),
    eco_lvids: str = Form(''),
    eco_ivsd: str = Form(''),
    eco_ivss: str = Form(''),
    eco_lvpwd: str = Form(''),
    eco_lvpws: str = Form(''),
    eco_fe: str = Form(''),
    eco_la_size: str = Form(''),
    eco_lv_size: str = Form(''),
    eco_rv_size: str = Form(''),
    eco_ra_size: str = Form(''),
    eco_mitral: str = Form(''),
    eco_tricuspid: str = Form(''),
    eco_aortic: str = Form(''),
    eco_pulmonary: str = Form(''),
    eco_pulmonary_htn: str = Form(''),
    eco_pericardium: str = Form(''),
    eco_doppler: str = Form(''),
    eco_observations: str = Form(''),

    # Radiografía cardiológica
    rx_vhs: str = Form(''),
    rx_vlas: str = Form(''),
    rx_heart_size: str = Form(''),
    rx_left_atrium: str = Form(''),
    rx_left_heart: str = Form(''),
    rx_right_heart: str = Form(''),
    rx_pulmonary_vessels: str = Form(''),
    rx_lung_pattern: str = Form(''),
    rx_edema: str = Form(''),
    rx_congestion: str = Form(''),
    rx_trachea: str = Form(''),
    rx_observations: str = Form(''),

    attachments: list[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)
    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    quick_event_type = request.session.pop('quick_event_type', '')

    quick_map = {
        'rx': 'Radiografía',
        'ecg': 'ECG',
        'eco': 'Ecografía',
        'lab': 'Laboratorio',
        'vacuna': 'Vacuna'
        }

    if quick_event_type in quick_map:
        event_type = quick_map[quick_event_type]
    def to_float(value):
        try:
            return float(value.replace(',', '.')) if value and value.strip() else None
        except ValueError:
            return None

    def to_int(value):
        try:
            return int(float(value.replace(',', '.'))) if value and value.strip() else None
        except ValueError:
            return None 



    rd = None
    
    def parse_date_field(value):
        if value and value.strip():
            try:
                return datetime.strptime(value.strip(), '%Y-%m-%d').date()
            except ValueError:
                return None
        return None
    
    extra_vaccine = False
    extra_deworming = False
    
    saved_vaccine_name = vaccine_name
    saved_vaccine_lot = vaccine_lot
    saved_vaccine_expiration = vaccine_expiration
    saved_next_vaccine_date = next_vaccine_date
    
    saved_dewormer_product = dewormer_product
    saved_dewormer_drug = dewormer_drug
    saved_dewormer_dose = dewormer_dose
    saved_next_deworming_date = next_deworming_date
    
    if event_type == 'Vacuna':
        rd = parse_date_field(next_vaccine_date)
    
        dewormer_product = ''
        dewormer_drug = ''
        dewormer_dose = ''
        next_deworming_date = ''
    
    elif event_type == 'Desparasitación':
        rd = parse_date_field(next_deworming_date)
    
        vaccine_name = ''
        vaccine_lot = ''
        vaccine_expiration = ''
        next_vaccine_date = ''
    
    else:
        extra_vaccine = any([
            vaccine_name.strip(),
            vaccine_lot.strip(),
            vaccine_expiration.strip(),
            next_vaccine_date.strip()
        ])
    
        extra_deworming = any([
            dewormer_product.strip(),
            dewormer_drug.strip(),
            dewormer_dose.strip(),
            next_deworming_date.strip()
        ])
    
        vaccine_name = ''
        vaccine_lot = ''
        vaccine_expiration = ''
        next_vaccine_date = ''
    
        dewormer_product = ''
        dewormer_drug = ''
        dewormer_dose = ''
        next_deworming_date = ''
    
        rd = parse_date_field(reminder_date)
    managed_reminder_date = None

    if reminder_type and reminder_date and reminder_date.strip():
        try:
            managed_reminder_date = datetime.strptime(reminder_date.strip(), '%Y-%m-%d').date()
            rd = None
        except ValueError:
            managed_reminder_date = None    
    event_created_at = argentina_now()
    event_date = event_date or ""
    
    if event_date.strip():
        try:
            selected_date = datetime.strptime(event_date.strip(), "%Y-%m-%d").date()
            current_time = argentina_now().time()
            event_created_at = datetime.combine(selected_date, current_time)
        except ValueError:
            pass

    event = ClinicalEvent(
        patient_id=patient_id,
        event_date=event_created_at,
        event_type=event_type,
        title=title or event_type,
        description=description or '',
        anamnesis=anamnesis or '',
        physical_exam=physical_exam or '',
        diagnosis=diagnosis or '',
        treatment=treatment or '',
        vaccine_name=vaccine_name or '',
        vaccine_lot=vaccine_lot or '',
        vaccine_expiration=vaccine_expiration or '',
        next_vaccine_date=next_vaccine_date or '',
        dewormer_product=dewormer_product or '',
        dewormer_drug=dewormer_drug or '',
        dewormer_dose=dewormer_dose or '',
        next_deworming_date=next_deworming_date or '',
        reminder_date=rd,
        weight=to_float(weight),
        temperature=to_float(temperature),
        heart_rate=to_int(heart_rate),
        respiratory_rate=to_int(respiratory_rate),
        mucous_membranes=mucous_membranes or '',
        crt=crt or '',
        hydration=hydration or '',
        ecg_hr=ecg_hr or '',
        ecg_rhythm=ecg_rhythm or '',
        ecg_p=ecg_p or '',
        ecg_pr=ecg_pr or '',
        ecg_qrs=ecg_qrs or '',
        ecg_st=ecg_st or '',
        ecg_t=ecg_t or '',
        ecg_qt=ecg_qt or '',
        ecg_axis=ecg_axis or '',
        ecg_interpretation=ecg_interpretation or '',
        eco_aiao=eco_aiao or '',
        eco_fs=eco_fs or '',
        eco_acvim=eco_acvim or '',
        eco_diagnosis=eco_diagnosis or '',
        eco_treatment=eco_treatment or '',

        ecg_p_mv=ecg_p_mv or '',
        ecg_qrs_mv=ecg_qrs_mv or '',
        ecg_t_mv=ecg_t_mv or '',
        ecg_qtc=ecg_qtc or '',
        ecg_polarity=ecg_polarity or '',
        ecg_arrhythmia=ecg_arrhythmia or '',
        ecg_conduction=ecg_conduction or '',
        ecg_notes=ecg_notes or '',

        eco_epss=eco_epss or '',
        eco_lvidd=eco_lvidd or '',
        eco_lvids=eco_lvids or '',
        eco_ivsd=eco_ivsd or '',
        eco_ivss=eco_ivss or '',
        eco_lvpwd=eco_lvpwd or '',
        eco_lvpws=eco_lvpws or '',
        eco_fe=eco_fe or '',
        eco_la_size=eco_la_size or '',
        eco_lv_size=eco_lv_size or '',
        eco_rv_size=eco_rv_size or '',
        eco_ra_size=eco_ra_size or '',
        eco_mitral=eco_mitral or '',
        eco_tricuspid=eco_tricuspid or '',
        eco_aortic=eco_aortic or '',
        eco_pulmonary=eco_pulmonary or '',
        eco_pulmonary_htn=eco_pulmonary_htn or '',
        eco_pericardium=eco_pericardium or '',
        eco_doppler=eco_doppler or '',
        eco_observations=eco_observations or '',

        rx_vhs=rx_vhs or '',
        rx_vlas=rx_vlas or '',
        rx_heart_size=rx_heart_size or '',
        rx_left_atrium=rx_left_atrium or '',
        rx_left_heart=rx_left_heart or '',
        rx_right_heart=rx_right_heart or '',
        rx_pulmonary_vessels=rx_pulmonary_vessels or '',
        rx_lung_pattern=rx_lung_pattern or '',
        rx_edema=rx_edema or '',
        rx_congestion=rx_congestion or '',
        rx_trachea=rx_trachea or '',
        rx_observations=rx_observations or '',

        created_by=user.username
    )

    db.add(event)
    db.commit()
    db.refresh(event)
    if extra_vaccine:
        vaccine_event = ClinicalEvent(
            patient_id=patient_id,
            event_date=event_created_at,
            event_type='Vacuna',
            title=saved_vaccine_name or 'Vacuna',
            description='Vacuna aplicada dentro de la misma visita clínica.',
            vaccine_name=saved_vaccine_name or '',
            vaccine_lot=saved_vaccine_lot or '',
            vaccine_expiration=saved_vaccine_expiration or '',
            next_vaccine_date=saved_next_vaccine_date or '',
            reminder_date=parse_date_field(saved_next_vaccine_date),
            created_by=user.username
        )
    
        db.add(vaccine_event)
    
    if extra_deworming:
        deworming_event = ClinicalEvent(
            patient_id=patient_id,
            event_date=event_created_at,
            event_type='Desparasitación',
            title=saved_dewormer_product or 'Desparasitación',
            description='Desparasitación aplicada dentro de la misma visita clínica.',
            dewormer_product=saved_dewormer_product or '',
            dewormer_drug=saved_dewormer_drug or '',
            dewormer_dose=saved_dewormer_dose or '',
            next_deworming_date=saved_next_deworming_date or '',
            reminder_date=parse_date_field(saved_next_deworming_date),
            created_by=user.username
        )
    
        db.add(deworming_event)
    
    if extra_vaccine or extra_deworming:
        db.commit()    
    close_managed_due_events(db, patient, event, user)
    db.commit()
    if managed_reminder_date:
        reminder_title = 'Recordatorio pendiente'
        reminder_event_type = 'Control'

        if reminder_type == 'Vacuna':
            reminder_event_type = 'Vacuna'
            reminder_title = f'Vacuna pendiente: {reminder_vaccine_name or "Vacuna"}'

        elif reminder_type == 'Desparasitación':
            reminder_event_type = 'Desparasitación'
            reminder_title = f'Desparasitación pendiente: {reminder_dewormer_name or "Desparasitación"}'

        elif reminder_type == 'Control clínico':
            reminder_event_type = 'Control'
            reminder_title = 'Control clínico pendiente'

        reminder_description = (
            f'{DUE_ACTIVE_MARKER}\n'
            f'Tipo de recordatorio: {reminder_type}\n'
            f'Vacuna: {reminder_vaccine_name or "-"}\n'
            f'Desparasitación: {reminder_dewormer_name or "-"}\n'
            f'Notas: {reminder_notes or "-"}'
        )

        pending_event = ClinicalEvent(
            patient_id=patient.id,
            event_date=argentina_now(),
            event_type=reminder_event_type,
            title=reminder_title,
            description=reminder_description,
            reminder_date=managed_reminder_date,
            created_by=user.username
        )

        db.add(pending_event)
        db.commit()    
    if weight and str(weight).strip():
        try:
            patient.weight = float(str(weight).replace(',', '.'))
            db.commit()
        except ValueError:
            pass
    active_waiting_entries = (
        db.query(WaitingListEntry)
        .filter(WaitingListEntry.patient_id == patient.id)
        .filter(WaitingListEntry.status.in_(['Esperando', 'En consulta']))
        .all()
    )

    for waiting_entry in active_waiting_entries:
        waiting_entry.status = 'Finalizado'
        waiting_entry.finished_at = argentina_now()

    if active_waiting_entries:
        db.commit()
    for file in attachments:
        if not file or not file.filename:
            continue

        if supabase is None:
            raise HTTPException(
                status_code=500,
                detail="Supabase Storage no configurado. Faltan SUPABASE_URL o SUPABASE_KEY en Render."
            )

        original_name = os.path.basename(file.filename)
        safe_name = safe_storage_filename(original_name)
        unique_name = f"{uuid.uuid4().hex}_{safe_name}"

        content = file.file.read()

        if not content:
            continue

        content_type = file.content_type or mimetypes.guess_type(original_name)[0] or "application/octet-stream"
        storage_path = f"patient_{patient_id}/event_{event.id}/{unique_name}"

        try:
            supabase.storage.from_("adjuntos").upload(
                path=storage_path,
                file=content,
                file_options={
                    "content-type": content_type,
                    "upsert": "false"
                }
            )

            public_url = supabase.storage.from_("adjuntos").get_public_url(storage_path)

        except Exception as e:
            import traceback
        
            print("=" * 80)
            print("ERROR SUBIENDO A SUPABASE")
            traceback.print_exc()
            print("=" * 80)
        
            db.rollback()
        
            raise HTTPException(
                status_code=500,
                detail=str(e)
            )

        attachment = EventAttachment(
            event_id=event.id,
            filename=original_name,
            file_path=public_url
        )

        db.add(attachment)

    db.commit()

    return RedirectResponse(
        url=f"/patients/{patient_id}",
        status_code=303
    )
    

@app.get('/products', response_class=HTMLResponse)
def products_page(
    request: Request,
    q: str = '',
    rubro: str = '',
    page: int = 1,
    alert: str = '',
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    query = db.query(Product).filter(Product.active == True)

    today = argentina_now().date()
    soon = today + timedelta(days=60)

    if q:
        like = f"%{q}%"
        query = query.filter(
            or_(
                Product.name.ilike(like),
                Product.code.ilike(like),
                Product.barcode.ilike(like),
                Product.provider.ilike(like),
                Product.manufacturer.ilike(like),
            )
        )

    if rubro:
        query = query.filter(Product.rubro == rubro)

    if alert == 'critical':
        query = query.filter(
            Product.stock != None,
            Product.min_stock != None,
            Product.min_stock > 0,
            Product.stock <= Product.min_stock
        )

    elif alert == 'expired':
        query = query.filter(
            Product.expiration_date != None,
            Product.expiration_date < today
        )

    elif alert == 'expires_soon':
        query = query.filter(
            Product.expiration_date != None,
            Product.expiration_date >= today,
            Product.expiration_date <= soon
        )

    total_products = db.query(Product).filter(Product.active == True).count()

    low_stock = db.query(Product).filter(
        Product.active == True,
        Product.stock != None,
        Product.min_stock != None,
        Product.min_stock > 0,
        Product.stock <= Product.min_stock
    ).count()

    expired_or_soon = db.query(Product).filter(
        Product.active == True,
        Product.expiration_date != None,
        Product.expiration_date <= soon
    ).count()

    per_page = 50

    if page < 1:
        page = 1

    total_filtered = query.count()
    total_pages = max(1, (total_filtered + per_page - 1) // per_page)

    if page > total_pages:
        page = total_pages

    products = (
        query
        .order_by(Product.name)
        .offset((page - 1) * per_page)
        .limit(per_page)
        .all()
    )

    for p in products:
        p.row_color = ''

        if p.expiration_date:
            if p.expiration_date < today:
                p.row_color = '#ffd6d6'
            elif (p.expiration_date - today).days <= 60:
                p.row_color = '#fff6cc'

        if (
            p.stock is not None and
            p.min_stock is not None and
            p.min_stock > 0 and
            p.stock <= p.min_stock
        ):
            p.row_color = '#ffe5e5'

    total_cost = 0
    total_sale = 0

    for p in db.query(Product).filter(Product.active == True).all():
        stock = p.stock or 0
        total_cost += stock * (p.cost_price or 0)
        total_sale += stock * (p.sale_price or 0)

    avg_margin = 0
    if total_cost > 0:
        avg_margin = ((total_sale - total_cost) / total_cost) * 100

    rubros = [
        r[0] for r in db.query(Product.rubro)
        .filter(Product.rubro != '')
        .distinct()
        .order_by(Product.rubro)
        .all()
    ]

    return templates.TemplateResponse(
        'products.html',
        {
            'request': request,
            'products': products,
            'q': q,
            'rubro': rubro,
            'alert': alert,
            'rubros': rubros,
            'total_products': total_products,
            'low_stock': low_stock,
            'expired_or_soon': expired_or_soon,
            'total_cost': total_cost,
            'total_sale': total_sale,
            'avg_margin': avg_margin,
            'today': today,
            'page': page,
            'total_pages': total_pages,
            'total_filtered': total_filtered,
            'per_page': per_page
        }
    )
@app.get('/products/export/csv')
def export_products_csv(
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    products = (
        db.query(Product)
        .filter(Product.active == True)
        .order_by(Product.name)
        .all()
    )

    output = StringIO()
    writer = csv.writer(output, delimiter=';')

    writer.writerow([
        'Nombre', 'Rubro', 'Tipo', 'Código', 'Código de barras',
        'Costo', 'Precio venta', 'Margen %', 'Stock', 'Stock mínimo',
        'Vencimiento', 'Proveedor', 'Laboratorio', 'Observaciones'
    ])

    for p in products:
        writer.writerow([
            p.name or '',
            p.rubro or '',
            p.tipo or '',
            p.code or '',
            p.barcode or '',
            p.cost_price if p.cost_price is not None else '',
            p.sale_price if p.sale_price is not None else '',
            p.margin_percent if p.margin_percent is not None else '',
            p.stock if p.stock is not None else '',
            p.min_stock if p.min_stock is not None else '',
            p.expiration_date.strftime('%d/%m/%Y') if p.expiration_date else '',
            p.provider or '',
            p.manufacturer or '',
            p.notes or ''
        ])

    filename = f"productos_{argentina_now().strftime('%Y-%m-%d')}.csv"

    return Response(
        content=output.getvalue().encode('utf-8-sig'),
        media_type='text/csv; charset=utf-8',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
    )


@app.get('/products/export/excel')
def export_products_excel(
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    products = (
        db.query(Product)
        .filter(Product.active == True)
        .order_by(Product.name)
        .all()
    )

    wb = Workbook()
    ws = wb.active
    ws.title = "Productos"

    headers = [
        'Nombre', 'Rubro', 'Tipo', 'Código', 'Código de barras',
        'Costo', 'Precio venta', 'Margen %', 'Stock', 'Stock mínimo',
        'Vencimiento', 'Proveedor', 'Laboratorio', 'Observaciones'
    ]

    ws.append(headers)

    for p in products:
        ws.append([
            p.name or '',
            p.rubro or '',
            p.tipo or '',
            p.code or '',
            p.barcode or '',
            p.cost_price if p.cost_price is not None else '',
            p.sale_price if p.sale_price is not None else '',
            p.margin_percent if p.margin_percent is not None else '',
            p.stock if p.stock is not None else '',
            p.min_stock if p.min_stock is not None else '',
            p.expiration_date.strftime('%d/%m/%Y') if p.expiration_date else '',
            p.provider or '',
            p.manufacturer or '',
            p.notes or ''
        ])

    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)

    filename = f"productos_{argentina_now().strftime('%Y-%m-%d')}.xlsx"

    return Response(
        content=stream.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
    )    
@app.post('/products/invoice/preview', response_class=HTMLResponse)
async def products_invoice_preview(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    filename = (file.filename or "").lower()
    if not filename.endswith(".pdf"):
        return RedirectResponse('/products?invoice_error=Solo+se+admiten+facturas+PDF', status_code=303)

    content = await file.read()
    if not content:
        return RedirectResponse('/products?invoice_error=El+archivo+está+vacío', status_code=303)

    try:
        payload = invoice_parse_pdf(content)
    except Exception as exc:
        message = str(exc).replace(' ', '+')
        return RedirectResponse(f'/products?invoice_error={message}', status_code=303)

    if payload.get("provider_cuit") and payload.get("invoice_number"):
        existing_invoice = db.execute(
            text("""
                SELECT id
                FROM purchase_invoices
                WHERE provider_cuit = :provider_cuit
                  AND invoice_number = :invoice_number
                LIMIT 1
            """),
            {
                "provider_cuit": payload["provider_cuit"],
                "invoice_number": payload["invoice_number"],
            }
        ).first()
        if existing_invoice:
            return RedirectResponse('/products?invoice_error=Esta+factura+ya+fue+importada', status_code=303)

    products = (
        db.query(Product)
        .filter(Product.active == True)
        .order_by(Product.name)
        .all()
    )

    for index, item in enumerate(payload["items"]):
        matched, score, match_type = invoice_find_match(
            db,
            payload.get("provider_cuit", ""),
            item["description"],
            products,
        )
        item["index"] = index
        item["matched_product_id"] = matched.id if matched else None
        item["matched_product_name"] = matched.name if matched else ""
        item["confidence"] = round(score * 100, 1)
        item["match_type"] = match_type

        if matched:
            current_cost = matched.cost_price or 0
            current_sale = matched.sale_price or 0
            margin = matched.margin_percent
            if margin is None and current_cost > 0:
                margin = ((current_sale - current_cost) / current_cost) * 100
            if margin is None or margin < -100 or margin > 1000:
                margin = 0

            item["current_cost"] = current_cost
            item["current_sale"] = current_sale
            item["current_margin"] = margin
            item["new_sale"] = round(item["unit_cost_vat"] * (1 + margin / 100), 2)
            item["status"] = "safe" if match_type in {"alias", "regla"} or score >= 0.86 else "review"
        else:
            item["current_cost"] = 0
            item["current_sale"] = 0
            item["current_margin"] = 0
            item["new_sale"] = round(item["unit_cost_vat"], 2)
            item["status"] = "new"

    token = invoice_save_preview(payload)

    return templates.TemplateResponse(
        'product_invoice_preview.html',
        {
            'request': request,
            'invoice': payload,
            'products': products,
            'token': token,
        }
    )


@app.post('/products/invoice/confirm')
async def products_invoice_confirm(
    request: Request,
    token: str = Form(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    payload = invoice_load_preview(token)
    form = await request.form()

    provider_cuit = payload.get("provider_cuit", "")
    invoice_number = payload.get("invoice_number", "")

    if provider_cuit and invoice_number:
        already_imported = db.execute(
            text("""
                SELECT id FROM purchase_invoices
                WHERE provider_cuit = :provider_cuit
                  AND invoice_number = :invoice_number
                LIMIT 1
            """),
            {"provider_cuit": provider_cuit, "invoice_number": invoice_number}
        ).first()
        if already_imported:
            return RedirectResponse('/products?invoice_error=Esta+factura+ya+fue+importada', status_code=303)

    updated = 0
    created = 0
    skipped = 0

    try:
        for item in payload.get("items", []):
            index = item["index"]
            if form.get(f"include_{index}") != "1":
                skipped += 1
                continue

            selected_id = (form.get(f"product_id_{index}") or "").strip()
            create_new = form.get(f"create_new_{index}") == "1" or not selected_id
            quantity = float(item.get("quantity") or 0)
            new_cost = round(float(item.get("unit_cost_vat") or 0), 4)

            if quantity <= 0 or new_cost <= 0:
                raise ValueError(f"Datos inválidos en {item.get('description', 'un producto')}")

            if create_new:
                new_name = (form.get(f"new_name_{index}") or item["description"]).strip()
                rubro = (form.get(f"new_rubro_{index}") or "").strip()
                tipo = (form.get(f"new_tipo_{index}") or "").strip()

                product = Product(
                    name=new_name,
                    rubro=rubro,
                    tipo=tipo,
                    cost_price=new_cost,
                    sale_price=new_cost,
                    margin_percent=0,
                    stock=quantity,
                    min_stock=0,
                    provider=payload.get("provider_name", ""),
                    manufacturer="",
                    notes=(
                        f"Creado desde factura {invoice_number}. "
                        "Producto nuevo con margen 0%; revisar precio de venta."
                    )
                )
                db.add(product)
                db.flush()
                created += 1
            else:
                product = db.get(Product, int(selected_id))
                if not product or not product.active:
                    raise ValueError(f"Producto seleccionado inválido para {item['description']}")

                old_cost = product.cost_price or 0
                old_sale = product.sale_price or 0
                margin = product.margin_percent

                if margin is None and old_cost > 0:
                    margin = ((old_sale - old_cost) / old_cost) * 100
                if margin is None or margin < -100 or margin > 1000:
                    margin = 0

                db.execute(
                    text("""
                        INSERT INTO product_cost_history (
                            product_id, provider_cuit, provider_name, invoice_number,
                            invoice_date, old_cost, new_cost, old_sale, new_sale,
                            margin_percent, quantity_added, created_by
                        ) VALUES (
                            :product_id, :provider_cuit, :provider_name, :invoice_number,
                            :invoice_date, :old_cost, :new_cost, :old_sale, :new_sale,
                            :margin_percent, :quantity_added, :created_by
                        )
                    """),
                    {
                        "product_id": product.id,
                        "provider_cuit": provider_cuit,
                        "provider_name": payload.get("provider_name", ""),
                        "invoice_number": invoice_number,
                        "invoice_date": payload.get("invoice_date", ""),
                        "old_cost": old_cost,
                        "new_cost": new_cost,
                        "old_sale": old_sale,
                        "new_sale": round(new_cost * (1 + margin / 100), 2),
                        "margin_percent": margin,
                        "quantity_added": quantity,
                        "created_by": user.username,
                    }
                )

                product.cost_price = new_cost
                product.sale_price = round(new_cost * (1 + margin / 100), 2)
                product.margin_percent = margin
                product.stock = (product.stock or 0) + quantity
                if payload.get("provider_name"):
                    product.provider = payload["provider_name"]
                updated += 1

            db.execute(
                text("""
                    INSERT INTO product_invoice_aliases (
                        provider_cuit, provider_name, source_description,
                        source_normalized, product_id, updated_at
                    ) VALUES (
                        :provider_cuit, :provider_name, :source_description,
                        :source_normalized, :product_id, CURRENT_TIMESTAMP
                    )
                    ON CONFLICT (provider_cuit, source_normalized)
                    DO UPDATE SET
                        product_id = EXCLUDED.product_id,
                        source_description = EXCLUDED.source_description,
                        provider_name = EXCLUDED.provider_name,
                        updated_at = CURRENT_TIMESTAMP
                """),
                {
                    "provider_cuit": provider_cuit,
                    "provider_name": payload.get("provider_name", ""),
                    "source_description": item["description"],
                    "source_normalized": invoice_normalize(item["description"]),
                    "product_id": product.id,
                }
            )

        db.execute(
            text("""
                INSERT INTO purchase_invoices (
                    provider_cuit, provider_name, invoice_number,
                    invoice_date, vat_percent, items_count,
                    updated_products, created_products, skipped_items, created_by
                ) VALUES (
                    :provider_cuit, :provider_name, :invoice_number,
                    :invoice_date, :vat_percent, :items_count,
                    :updated_products, :created_products, :skipped_items, :created_by
                )
            """),
            {
                "provider_cuit": provider_cuit,
                "provider_name": payload.get("provider_name", ""),
                "invoice_number": invoice_number,
                "invoice_date": payload.get("invoice_date", ""),
                "vat_percent": payload.get("vat_percent", 21),
                "items_count": len(payload.get("items", [])),
                "updated_products": updated,
                "created_products": created,
                "skipped_items": skipped,
                "created_by": user.username,
            }
        )

        db.commit()
        try:
            os.remove(invoice_preview_path(token))
        except OSError:
            pass

    except Exception as exc:
        db.rollback()
        message = str(exc).replace(' ', '+')
        return RedirectResponse(f'/products?invoice_error={message}', status_code=303)

    return RedirectResponse(
        f'/products?invoice_ok=Actualizados:+{updated}+|+Nuevos:+{created}+|+Omitidos:+{skipped}',
        status_code=303
    )

@app.post('/products/import')
async def import_products(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    content = await file.read()

    wb = load_workbook(BytesIO(content), data_only=True)
    ws = wb.active

    headers = []
    for cell in ws[1]:
        headers.append(str(cell.value).strip() if cell.value else "")

    def get_value(row, *column_names):
        normalized = {
            str(h).strip().lower(): i
            for i, h in enumerate(headers)
            if h
        }
    
        for name in column_names:
            idx = normalized.get(str(name).strip().lower())
            if idx is not None and idx < len(row):
                return row[idx]
    
        return None

    def to_float(value):
        try:
            if value is None:
                return None
            value = str(value).strip()
            if value == "" or value == "---":
                return None
            return float(value.replace(",", "."))
        except ValueError:
            return None

    def to_date(value):
        try:
            if value is None:
                return None
            if hasattr(value, "date"):
                return value.date()
            value = str(value).strip()
            if value == "" or value.lower() == "sin asignar":
                return None
            return datetime.strptime(value, "%d/%m/%Y").date()
        except ValueError:
            return None

    imported = 0

    for row in ws.iter_rows(min_row=2, values_only=True):
        name = str(get_value(row, "Nombre") or "").strip()
        if not name:
            continue

        cost = to_float(
            get_value(row, "Costo", "Precio costo")
        )
        
        sale = to_float(
            get_value(row, "Precio", "Precio venta")
        )
        
        margin = to_float(
            get_value(row, "Margen (%)", "Margen %")
        )

        if margin is None and cost and sale:
            margin = ((sale - cost) / cost) * 100

        code = str(get_value(row, "Código") or "").strip()
        barcode = str(
            get_value(row, "Código de Barras", "Código de barras")
            or ""
        ).strip()
        
        existing = None
        
        if code:
            existing = db.query(Product).filter(Product.code == code).first()
        
        if existing is None and barcode:
            existing = db.query(Product).filter(Product.barcode == barcode).first()
        
        if existing is None:
            existing = db.query(Product).filter(Product.name == name).first()
        
        rubro_value = str(get_value(row, "Rubro") or "").strip()
        tipo_value = str(get_value(row, "Tipo") or "").strip()
        provider_value = str(get_value(row, "Proveedor", "Proveedores") or "").strip()
        manufacturer_value = str(get_value(row, "Laboratorio", "Elaborador", "Manufacturer") or "").strip()
        notes_value = str(get_value(row, "Observaciones", "Notas", "Nota") or "").strip()

        if existing:
            existing.rubro = rubro_value
            existing.tipo = tipo_value
            existing.name = name
            existing.code = code
            existing.barcode = barcode
            existing.sale_price = sale
            existing.cost_price = cost
            existing.margin_percent = margin
            existing.stock = to_float(get_value(row, "Stock"))
            existing.min_stock = to_float(get_value(row, "Stock Min", "Stock mínimo"))
            existing.expiration_date = to_date(get_value(row, "Vencimiento"))
            existing.manufacturer = manufacturer_value
            existing.provider = provider_value
            existing.notes = notes_value
            existing.active = True
        else:
            product = Product(
                rubro=rubro_value,
                tipo=tipo_value,
                name=name,
                code=code,
                barcode=barcode,
                sale_price=sale,
                cost_price=cost,
                margin_percent=margin,
                stock=to_float(get_value(row, "Stock")),
                min_stock=to_float(get_value(row, "Stock Min", "Stock mínimo")),
                expiration_date=to_date(get_value(row, "Vencimiento")),
                manufacturer=manufacturer_value,
                provider=provider_value,
                notes=notes_value
            )

            db.add(product)
        imported += 1

    db.commit()

    return RedirectResponse(
        url=f"/products?imported={imported}",
        status_code=303
    )
@app.post('/products')
def product_create(
    name: str = Form(''),
    rubro: str = Form(''),
    tipo: str = Form(''),
    code: str = Form(''),
    barcode: str = Form(''),
    cost_price: str = Form(''),
    sale_price: str = Form(''),
    stock: str = Form(''),
    min_stock: str = Form(''),
    expiration_date: str = Form(''),
    provider: str = Form(''),
    manufacturer: str = Form(''),
    notes: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    def to_float(value):
        try:
            return float(value.replace(',', '.')) if value and value.strip() else None
        except ValueError:
            return None

    exp = None
    if expiration_date and expiration_date.strip():
        try:
            exp = datetime.strptime(expiration_date.strip(), "%Y-%m-%d").date()
        except ValueError:
            exp = None

    cost = to_float(cost_price)
    sale = to_float(sale_price)
    margin = None
    if cost and sale:
        margin = ((sale - cost) / cost) * 100

    product = Product(
        name=name or '',
        rubro=rubro or '',
        tipo=tipo or '',
        code=code or '',
        barcode=barcode or '',
        cost_price=cost,
        sale_price=sale,
        margin_percent=margin,
        stock=to_float(stock),
        min_stock=to_float(min_stock),
        expiration_date=exp,
        provider=provider or '',
        manufacturer=manufacturer or '',
        notes=notes or ''
    )

    db.add(product)
    db.commit()

    return RedirectResponse('/products', status_code=303)
@app.get('/products/{product_id}/edit', response_class=HTMLResponse)
def product_edit_page(
    request: Request,
    product_id: int,
    page: int = 1,
    q: str = '',
    rubro: str = '',
    alert: str = '',
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    product = db.get(Product, product_id)

    if not product:
        raise HTTPException(
            status_code=404,
            detail="Producto no encontrado"
        )

    if page < 1:
        page = 1

    return templates.TemplateResponse(
        'product_edit.html',
        {
            'request': request,
            'product': product,
            'page': page,
            'q': q or '',
            'rubro': rubro or '',
            'alert': alert or ''
        }
    )


@app.post('/products/{product_id}/edit')
def product_edit_save(
    product_id: int,
    name: str = Form(''),
    rubro: str = Form(''),
    tipo: str = Form(''),
    code: str = Form(''),
    barcode: str = Form(''),
    cost_price: str = Form(''),
    sale_price: str = Form(''),
    margin_percent: str = Form(''),
    stock: str = Form(''),
    min_stock: str = Form(''),
    provider: str = Form(''),
    manufacturer: str = Form(''),
    notes: str = Form(''),

    page: int = Form(1),
    return_q: str = Form(''),
    return_rubro: str = Form(''),
    return_alert: str = Form(''),

    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    product = db.get(Product, product_id)

    if not product:
        raise HTTPException(
            status_code=404,
            detail="Producto no encontrado"
        )

    def to_float(value):
        try:
            return (
                float(str(value).replace(',', '.'))
                if value is not None and str(value).strip()
                else None
            )
        except (ValueError, TypeError):
            return None

    product.name = name or ''
    product.rubro = rubro or ''
    product.tipo = tipo or ''
    product.code = code or ''
    product.barcode = barcode or ''
    product.cost_price = to_float(cost_price)
    product.sale_price = to_float(sale_price)
    product.margin_percent = to_float(margin_percent)
    product.stock = to_float(stock)
    product.min_stock = to_float(min_stock)
    product.provider = provider or ''
    product.manufacturer = manufacturer or ''
    product.notes = notes or ''

    db.commit()

    if page < 1:
        page = 1

    return_params = {
        'page': page,
        'q': return_q or '',
        'rubro': return_rubro or ''
    }

    if return_alert:
        return_params['alert'] = return_alert

    return_url = '/products?' + urlencode(return_params)

    return RedirectResponse(
        url=return_url,
        status_code=303
    )
@app.post('/products/{product_id}/delete')
def product_delete(
    product_id: int,
    page: int = Form(1),
    return_q: str = Form(''),
    return_rubro: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    product = db.get(Product, product_id)

    if not product:
        raise HTTPException(
            status_code=404,
            detail="Producto no encontrado"
        )

    # No se borra físicamente de la base.
    # Se desactiva para conservar ventas, presupuestos,
    # facturas y movimientos históricos.
    product.active = False

    db.commit()

    if page < 1:
        page = 1

    return_params = {
        'page': page,
        'q': return_q or '',
        'rubro': return_rubro or '',
        'deleted': 'Producto eliminado correctamente'
    }

    return_url = '/products?' + urlencode(return_params)

    return RedirectResponse(
        url=return_url,
        status_code=303
    )



@app.post('/products/{product_id}/adjust')
def product_adjust_stock(
    product_id: int,
    qty: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    product = db.get(Product, product_id)

    if not product:
        raise HTTPException(
            status_code=404,
            detail="Producto no encontrado"
        )

    try:
        amount = float(str(qty).replace(',', '.'))
    except ValueError:
        amount = 0

    current_stock = product.stock or 0
    product.stock = current_stock + amount

    db.commit()

    return RedirectResponse(
        url='/products',
        status_code=303
    )
# ===== VENTAS =====

@app.get('/sales', response_class=HTMLResponse)
def sales_page(
    request: Request,
    sale_date: str = '',
    period: str = 'today',
    date_from: str = '',
    date_to: str = '',
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    products = (
        db.query(Product)
        .filter(Product.active == True)
        .order_by(Product.name)
        .all()
    )

    patients = db.query(Patient).order_by(Patient.name).all()
    owners = db.query(Owner).order_by(Owner.name).all()
    patient_owner_map = {}

    for patient in patients:
        patient_owner_map[str(patient.id)] = (
            patient.owner_id if patient.owner_id else ''
        )

    today = argentina_now().date()

    valid_periods = {
        'today',
        'yesterday',
        '7days',
        'month',
        'last_month',
        'all',
        'custom'
    }

    selected_period = period if period in valid_periods else 'today'
    selected_date = today
    selected_date_from = date_from
    selected_date_to = date_to

    # Mantiene compatibilidad con el filtro anterior por una sola fecha.
    if sale_date:
        try:
            selected_date = datetime.strptime(
                sale_date,
                '%Y-%m-%d'
            ).date()

            selected_period = 'custom'
            selected_date_from = sale_date
            selected_date_to = sale_date

        except ValueError:
            selected_date = today

    if selected_period == 'yesterday':
        start_date = today - timedelta(days=1)
        end_date = start_date

    elif selected_period == '7days':
        start_date = today - timedelta(days=6)
        end_date = today

    elif selected_period == 'month':
        start_date = today.replace(day=1)
        end_date = today

    elif selected_period == 'last_month':
        first_this_month = today.replace(day=1)
        end_date = first_this_month - timedelta(days=1)
        start_date = end_date.replace(day=1)

    elif selected_period == 'all':
        start_date = None
        end_date = None

    elif selected_period == 'custom':
        try:
            start_date = datetime.strptime(
                selected_date_from,
                '%Y-%m-%d'
            ).date()

            end_date = datetime.strptime(
                selected_date_to,
                '%Y-%m-%d'
            ).date()

            if start_date > end_date:
                start_date, end_date = end_date, start_date
                selected_date_from = start_date.isoformat()
                selected_date_to = end_date.isoformat()

        except ValueError:
            selected_period = 'today'
            start_date = today
            end_date = today
            selected_date_from = ''
            selected_date_to = ''

    else:
        selected_period = 'today'
        start_date = today
        end_date = today

    sales_query = (
        db.query(Sale)
        .filter(Sale.status != 'quote')
        .filter(Sale.status != 'cancelled')
    )

    if start_date is not None and end_date is not None:
        filter_start = datetime.combine(
            start_date,
            datetime.min.time()
        )

        filter_end = datetime.combine(
            end_date,
            datetime.max.time()
        )

        sales_query = sales_query.filter(
            Sale.date >= filter_start,
            Sale.date <= filter_end
        )

    sales = (
        sales_query
        .order_by(Sale.date.desc())
        .all()
    )

    sales_result_count = len(sales)

    for sale in sales:
        sale.items_count = (
            db.query(SaleItem)
            .filter(SaleItem.sale_id == sale.id)
            .count()
        )

        sale.patient_name = ''
        sale.owner_name = ''
        sale.items_detail = []

        sale.total_paid = sum(
            payment.amount or 0
            for payment in (
                db.query(SalePayment)
                .filter(SalePayment.sale_id == sale.id)
                .all()
            )
            if payment.method != 'Cuenta corriente'
        )

        sale.balance_due = (sale.total or 0) - sale.total_paid

        items = (
            db.query(SaleItem)
            .filter(SaleItem.sale_id == sale.id)
            .all()
        )

        for item in items:
            product = db.get(Product, item.product_id)

            sale.items_detail.append({
                'product_name': (
                    product.name
                    if product
                    else 'Producto eliminado'
                ),
                'quantity': item.quantity,
                'unit_price': item.unit_price,
                'subtotal': item.subtotal
            })

        if sale.patient_id:
            patient = db.get(Patient, sale.patient_id)

            if patient:
                sale.patient_name = patient.name

        if sale.owner_id:
            owner = db.get(Owner, sale.owner_id)

            if owner:
                sale.owner_name = owner.name

    month_start = today.replace(day=1)

    today_sales_total = 0
    month_sales_total = 0
    today_products_count = 0
    month_profit_total = 0
    month_cost_total = 0
    month_sales_count = 0
    ticket_average = 0

    all_sales = db.query(Sale).all()

    for sale in all_sales:
        current_sale_date = sale.date.date() if sale.date else None

        if current_sale_date == today:
            today_sales_total += sale.total or 0

        if current_sale_date and current_sale_date >= month_start:
            month_sales_total += sale.total or 0
            month_cost_total += sale.cost_total or 0
            month_profit_total += sale.profit_amount or 0
            month_sales_count += 1

    if month_sales_count > 0:
        ticket_average = month_sales_total / month_sales_count

    today_items = (
        db.query(SaleItem)
        .join(Sale, SaleItem.sale_id == Sale.id)
        .filter(
            Sale.date >= datetime.combine(
                today,
                datetime.min.time()
            )
        )
        .all()
    )

    for item in today_items:
        today_products_count += item.quantity or 0

    month_items = (
        db.query(SaleItem)
        .join(Sale, SaleItem.sale_id == Sale.id)
        .filter(
            Sale.date >= datetime.combine(
                month_start,
                datetime.min.time()
            )
        )
        .all()
    )

    product_stats = {}

    for item in month_items:
        product = db.get(Product, item.product_id)

        if not product:
            continue

        name = product.name

        if name not in product_stats:
            product_stats[name] = 0

        product_stats[name] += item.quantity or 0

    top_product_name = '-'
    top_product_qty = 0

    top_products_sold = sorted(
        product_stats.items(),
        key=lambda item: item[1],
        reverse=True
    )[:10]

    if product_stats:
        top_product_name = max(
            product_stats,
            key=product_stats.get
        )

        top_product_qty = product_stats[top_product_name]

    quote_sales = (
        db.query(Sale)
        .filter(Sale.status == 'quote')
        .order_by(Sale.date.desc())
        .all()
    )

    for quote in quote_sales:
        quote.items_count = (
            db.query(SaleItem)
            .filter(SaleItem.sale_id == quote.id)
            .count()
        )

        quote.patient_name = ''
        quote.owner_name = ''
        quote.items_detail = []

        items = (
            db.query(SaleItem)
            .filter(SaleItem.sale_id == quote.id)
            .all()
        )

        for item in items:
            product = db.get(Product, item.product_id)

            quote.items_detail.append({
                'product_name': (
                    product.name
                    if product
                    else 'Producto eliminado'
                ),
                'quantity': item.quantity,
                'unit_price': item.unit_price,
                'subtotal': item.subtotal
            })

        if quote.patient_id:
            patient = db.get(Patient, quote.patient_id)

            if patient:
                quote.patient_name = patient.name

        if quote.owner_id:
            owner = db.get(Owner, quote.owner_id)

            if owner:
                quote.owner_name = owner.name

    return templates.TemplateResponse(
        'sales_v2.html',
        {
            'request': request,
            'products': products,
            'patients': patients,
            'patient_owner_map': patient_owner_map,
            'owners': owners,
            'sales': sales,
            'quote_sales': quote_sales,
            'selected_date': selected_date,
            'selected_period': selected_period,
            'selected_date_from': selected_date_from,
            'selected_date_to': selected_date_to,
            'sales_result_count': sales_result_count,
            'today_sales_total': today_sales_total,
            'month_sales_total': month_sales_total,
            'today_products_count': today_products_count,
            'month_cost_total': month_cost_total,
            'month_profit_total': month_profit_total,
            'month_sales_count': month_sales_count,
            'ticket_average': ticket_average,
            'top_product_name': top_product_name,
            'top_product_qty': top_product_qty,
            'top_products_sold': top_products_sold
        }
    )
@app.post('/sales')
def sales_create(
    patient_id: str = Form(''),
    owner_id: str = Form(''),
    product_id: list[str] = Form([]),
    quantity: list[str] = Form([]),
    unit_price: list[str] = Form([]),
    notes: str = Form(''),

    discount_percent: str = Form('0'),
    discount_amount: str = Form('0'),
    credit_surcharge_percent: str = Form('0'),
    credit_surcharge_amount: str = Form('0'),

    pay_efectivo: str = Form('0'),
    pay_debito: str = Form('0'),
    pay_credito: str = Form('0'),
    pay_transferencia: str = Form('0'),
    pay_cuenta_corriente: str = Form('0'),
    save_as_quote: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    def to_float(value):
        try:
            if value is None or not str(value).strip():
                return 0.0

            return float(str(value).replace(',', '.'))

        except (TypeError, ValueError):
            return 0.0

    has_product = any(
        pid and str(pid).strip()
        for pid in product_id
    )

    if not has_product:
        return RedirectResponse(
            url='/sales',
            status_code=303
        )

    is_quote = bool(save_as_quote)

    sale = Sale(
        status='quote' if is_quote else 'pending',
        total=0,
        payment_method='Mixto',
        patient_id=int(patient_id) if patient_id else None,
        owner_id=int(owner_id) if owner_id else None,
        notes=notes or ''
    )

    db.add(sale)
    db.flush()

    subtotal_total = 0.0
    cost_total = 0.0

    for pid, qty_raw, price_raw in zip(
        product_id,
        quantity,
        unit_price
    ):
        if not pid:
            continue

        product = db.get(Product, int(pid))

        if not product:
            continue

        qty = to_float(qty_raw)
        price = to_float(price_raw)

        if qty <= 0:
            continue

        if price < 0:
            price = 0

        item_subtotal = round(qty * price, 2)
        product_cost = product.cost_price or 0

        subtotal_total += item_subtotal
        cost_total += qty * product_cost

        item = SaleItem(
            sale_id=sale.id,
            product_id=product.id,
            quantity=qty,
            unit_price=price,
            subtotal=item_subtotal
        )

        db.add(item)

        if not is_quote:
            current_stock = product.stock or 0
            product.stock = current_stock - qty

    subtotal_total = round(subtotal_total, 2)
    cost_total = round(cost_total, 2)

    applied_discount_percent = max(
        0.0,
        min(to_float(discount_percent), 100.0)
    )

    calculated_discount_amount = round(
        subtotal_total * applied_discount_percent / 100,
        2
    )

    calculated_discount_amount = min(
        calculated_discount_amount,
        subtotal_total
    )

    total_after_discount = round(
        subtotal_total - calculated_discount_amount,
        2
    )

    credit_payment = round(
        max(0.0, to_float(pay_credito)),
        2
    )
    
    applied_surcharge_percent = max(
        0.0,
        min(
            to_float(credit_surcharge_percent),
            100.0
        )
    )
    
    calculated_surcharge_amount = 0.0
    
    if (
        credit_payment > 0
        and applied_surcharge_percent > 0
    ):
        calculated_surcharge_amount = round(
            credit_payment
            * applied_surcharge_percent
            / (
                100
                + applied_surcharge_percent
            ),
            2
        )
    
    final_total = round(
        total_after_discount
        + calculated_surcharge_amount,
        2
    )

    sale.discount_percent = applied_discount_percent
    sale.discount_amount = calculated_discount_amount
    sale.credit_surcharge_percent = applied_surcharge_percent
    sale.credit_surcharge_amount = calculated_surcharge_amount

    sale.total = final_total
    sale.cost_total = cost_total
    sale.profit_amount = round(final_total - cost_total, 2)
    sale.margin_percent = (
        round(
            ((final_total - cost_total) / final_total) * 100,
            2
        )
        if final_total > 0
        else 0
    )

    payments_to_create = [
        ('Efectivo', to_float(pay_efectivo)),
        ('Débito', to_float(pay_debito)),
        ('Crédito', credit_payment),
        ('Transferencia', to_float(pay_transferencia)),
        ('Cuenta corriente', to_float(pay_cuenta_corriente)),
    ]

    real_payment_total = sum(
        amount
        for method, amount in payments_to_create
        if method != 'Cuenta corriente' and amount > 0
    )

    has_account_debt = any(
        method == 'Cuenta corriente' and amount > 0
        for method, amount in payments_to_create
    )

    if (
        final_total > 0
        and real_payment_total <= 0
        and not has_account_debt
        and not is_quote
    ):
        payments_to_create = [
            ('Efectivo', final_total)
        ]

    if not is_quote:
        real_payment_total = 0.0
        registered_methods = []

        for method, amount in payments_to_create:
            if amount <= 0:
                continue

            payment = SalePayment(
                sale_id=sale.id,
                method=method,
                amount=round(amount, 2)
            )

            db.add(payment)
            registered_methods.append(method)

            if method != 'Cuenta corriente':
                real_payment_total += amount

        real_payment_total = round(real_payment_total, 2)

        if real_payment_total >= final_total - 0.01:
            sale.status = 'paid'
        else:
            sale.status = 'pending'

        real_methods = [
            method
            for method in registered_methods
            if method != 'Cuenta corriente'
        ]

        if sale.status == 'paid':
            if len(real_methods) == 1:
                sale.payment_method = real_methods[0]
            elif len(real_methods) > 1:
                sale.payment_method = 'Mixto'
            else:
                sale.payment_method = 'Efectivo'
        else:
            sale.payment_method = 'Cuenta corriente'

    else:
        sale.status = 'quote'
        sale.payment_method = 'Presupuesto'

    db.commit()

    return RedirectResponse(
        url='/sales',
        status_code=303
    )
@app.post('/sales/{sale_id}/convert')
def sales_convert(
    sale_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    sale = db.get(Sale, sale_id)

    if not sale:
        raise HTTPException(
            status_code=404,
            detail='Venta no encontrada'
        )

    if sale.status != 'quote':
        return RedirectResponse(
            url=f'/sales/{sale_id}',
            status_code=303
        )

    items = (
        db.query(SaleItem)
        .filter(SaleItem.sale_id == sale.id)
        .all()
    )

    for item in items:
        product = db.get(Product, item.product_id)

        if product:
            product.stock = (
                product.stock or 0
            ) - (
                item.quantity or 0
            )

    real_paid = sum(
        payment.amount or 0
        for payment in (
            db.query(SalePayment)
            .filter(SalePayment.sale_id == sale.id)
            .all()
        )
        if not (
            'cuenta' in (payment.method or '').strip().lower()
            and
            'corriente' in (payment.method or '').strip().lower()
        )
    )

    if real_paid >= (sale.total or 0) - 0.01:
        sale.status = 'paid'
    else:
        sale.status = 'pending'

    sale.payment_method = (
        'Mixto'
        if real_paid > 0
        else 'Cuenta corriente'
    )

    db.commit()

    return RedirectResponse(
        url=f'/sales/{sale_id}',
        status_code=303
    )
@app.post('/sales/{sale_id}/cancel')
def sales_cancel(
    sale_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    sale = db.get(Sale, sale_id)

    if not sale:
        raise HTTPException(
            status_code=404,
            detail='Venta no encontrada'
        )

    if sale.status == 'cancelled':
        return RedirectResponse(
            url='/sales',
            status_code=303
        )

    items = (
        db.query(SaleItem)
        .filter(SaleItem.sale_id == sale.id)
        .all()
    )

    for item in items:
        product = db.get(Product, item.product_id)
        if product:
            current_stock = product.stock or 0
            product.stock = current_stock + (item.quantity or 0)

    sale.status = 'cancelled'

    db.commit()

    return RedirectResponse(
        url='/sales',
        status_code=303
    )
@app.get('/sales/{sale_id}', response_class=HTMLResponse)
def sales_detail(
    sale_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    sale = db.get(Sale, sale_id)

    if not sale:
        raise HTTPException(status_code=404, detail='Venta no encontrada')

    patient = db.get(Patient, sale.patient_id) if sale.patient_id else None
    owner = db.get(Owner, sale.owner_id) if sale.owner_id else None

    items = (
        db.query(SaleItem)
        .filter(SaleItem.sale_id == sale.id)
        .all()
    )

    item_details = []

    for item in items:
        product = db.get(Product, item.product_id)

        item_details.append({
            'product_name': product.name if product else 'Producto eliminado',
            'quantity': item.quantity,
            'unit_price': item.unit_price,
            'subtotal': item.subtotal
        })
    products = (
        db.query(Product)
        .filter(Product.active == True)
        .order_by(Product.name)
        .all()
    )
    payments = (
        db.query(SalePayment)
        .filter(SalePayment.sale_id == sale.id)
        .all()
    )
    
    total_paid = 0
    
    for p in payments:
        method_clean = (p.method or '').strip().lower()
    
        if 'cuenta' in method_clean and 'corriente' in method_clean:
            continue
    
        total_paid += p.amount or 0
    
    balance_due = (sale.total or 0) - total_paid
    return templates.TemplateResponse(
        'sale_detail.html',
    {
        'request': request,
        'sale': sale,
        'patient': patient,
        'owner': owner,
        'items': item_details,
        'products': products,
        'payments': payments,
        'balance_due': balance_due
    }
    )
@app.post('/sales/{sale_id}/add-item')
def sale_add_item(
    sale_id: int,
    product_id: str = Form(''),
    quantity: str = Form('1'),
    unit_price: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    def to_float(value):
        try:
            if value is None or not str(value).strip():
                return 0.0

            return float(str(value).replace(',', '.'))

        except (TypeError, ValueError):
            return 0.0

    sale = db.get(Sale, sale_id)

    if not sale:
        raise HTTPException(
            status_code=404,
            detail='Venta no encontrada'
        )

    if sale.status == 'cancelled':
        return RedirectResponse(
            url=f'/sales/{sale.id}',
            status_code=303
        )

    if not product_id:
        return RedirectResponse(
            url=f'/sales/{sale.id}',
            status_code=303
        )

    product = db.get(Product, int(product_id))

    if not product:
        raise HTTPException(
            status_code=404,
            detail='Producto no encontrado'
        )

    qty = to_float(quantity)
    price = to_float(unit_price)

    if qty <= 0:
        return RedirectResponse(
            url=f'/sales/{sale.id}',
            status_code=303
        )

    if price <= 0:
        price = product.sale_price or 0

    item_subtotal = round(qty * price, 2)

    item = SaleItem(
        sale_id=sale.id,
        product_id=product.id,
        quantity=qty,
        unit_price=price,
        subtotal=item_subtotal
    )

    db.add(item)

    if sale.status != 'quote' and product.stock is not None:
        product.stock = (product.stock or 0) - qty

    db.flush()

    all_items = (
        db.query(SaleItem)
        .filter(SaleItem.sale_id == sale.id)
        .all()
    )

    subtotal_total = 0.0
    cost_total = 0.0

    for current_item in all_items:
        current_product = db.get(
            Product,
            current_item.product_id
        )

        current_subtotal = (
            current_item.subtotal
            if current_item.subtotal is not None
            else
            (current_item.quantity or 0)
            * (current_item.unit_price or 0)
        )

        subtotal_total += current_subtotal or 0

        if current_product:
            cost_total += (
                current_item.quantity or 0
            ) * (
                current_product.cost_price or 0
            )

    subtotal_total = round(subtotal_total, 2)
    cost_total = round(cost_total, 2)

    discount_percent = max(
        0.0,
        min(sale.discount_percent or 0, 100.0)
    )

    discount_amount = round(
        subtotal_total * discount_percent / 100,
        2
    )

    total_after_discount = round(
        subtotal_total - discount_amount,
        2
    )

    payments = (
        db.query(SalePayment)
        .filter(SalePayment.sale_id == sale.id)
        .all()
    )

    credit_payment_total = sum(
        payment.amount or 0
        for payment in payments
        if (
            payment.method or ''
        ).strip().lower() in {
            'crédito',
            'credito'
        }
    )
    
    surcharge_percent = max(
        0.0,
        min(
            sale.credit_surcharge_percent or 0,
            100.0
        )
    )
    
    surcharge_amount = 0.0
    
    if (
        credit_payment_total > 0
        and surcharge_percent > 0
    ):
        surcharge_amount = round(
            credit_payment_total
            * surcharge_percent
            / (
                100
                + surcharge_percent
            ),
            2
        )
   
    

    final_total = round(
        total_after_discount + surcharge_amount,
        2
    )

    sale.discount_amount = discount_amount
    sale.credit_surcharge_amount = surcharge_amount
    sale.total = final_total
    sale.cost_total = cost_total
    sale.profit_amount = round(
        final_total - cost_total,
        2
    )
    sale.margin_percent = (
        round(
            ((final_total - cost_total) / final_total) * 100,
            2
        )
        if final_total > 0
        else 0
    )

    real_paid = sum(
        payment.amount or 0
        for payment in payments
        if not (
            'cuenta' in (payment.method or '').strip().lower()
            and
            'corriente' in (payment.method or '').strip().lower()
        )
    )

    if sale.status != 'quote':
        if real_paid >= final_total - 0.01:
            sale.status = 'paid'
        else:
            sale.status = 'pending'

    db.commit()

    return RedirectResponse(
        url=f'/sales/{sale.id}',
        status_code=303
    )
@app.post('/sales/{sale_id}/confirm')
def sale_confirm(
    sale_id: int,
    payment_method: str = Form('Efectivo'),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    sale = db.get(Sale, sale_id)

    if not sale:
        raise HTTPException(status_code=404, detail='Venta no encontrada')
    sale.payment_method = payment_method or 'Efectivo'
    sale.status = 'paid'

    db.commit()

    return RedirectResponse(
        url=f'/sales/{sale.id}',
        status_code=303
    )
@app.post('/sales/{sale_id}/add-payment')
def sale_add_payment(
    sale_id: int,
    method: str = Form('Efectivo'),
    amount: str = Form('0'),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    def to_float(value):
        try:
            if value is None or not str(value).strip():
                return 0.0

            return float(str(value).replace(',', '.'))

        except (TypeError, ValueError):
            return 0.0

    sale = db.get(Sale, sale_id)

    if not sale:
        raise HTTPException(
            status_code=404,
            detail='Venta no encontrada'
        )

    if sale.status == 'cancelled':
        return RedirectResponse(
            url=f'/sales/{sale.id}',
            status_code=303
        )

    payment_amount = round(
        to_float(amount),
        2
    )

    if payment_amount <= 0:
        return RedirectResponse(
            url=f'/sales/{sale.id}',
            status_code=303
        )

    payment_method = (
        method or 'Efectivo'
    ).strip()

    payment = SalePayment(
        sale_id=sale.id,
        method=payment_method,
        amount=payment_amount
    )

    db.add(payment)
    db.flush()

    payments = (
        db.query(SalePayment)
        .filter(SalePayment.sale_id == sale.id)
        .all()
    )

    real_paid = 0.0
    real_methods = []

    for current_payment in payments:
        current_method = (
            current_payment.method or ''
        ).strip()

        current_method_clean = current_method.lower()

        is_account_payment = (
            'cuenta' in current_method_clean
            and
            'corriente' in current_method_clean
        )

        if is_account_payment:
            continue

        real_paid += current_payment.amount or 0

        if current_method:
            real_methods.append(current_method)

    real_paid = round(real_paid, 2)
    balance_due = round(
        (sale.total or 0) - real_paid,
        2
    )

    if balance_due <= 0.01:
        sale.status = 'paid'

        unique_methods = list(dict.fromkeys(real_methods))

        if len(unique_methods) == 1:
            sale.payment_method = unique_methods[0]
        elif len(unique_methods) > 1:
            sale.payment_method = 'Mixto'
        else:
            sale.payment_method = payment_method

    else:
        sale.status = 'pending'
        sale.payment_method = 'Cuenta corriente'

    db.commit()

    return RedirectResponse(
        url=f'/sales/{sale.id}',
        status_code=303
    )
@app.get('/migration', response_class=HTMLResponse)
def migration(request: Request, user: User = Depends(require_user)):
    return templates.TemplateResponse(
        'migration.html',
        {
            'request': request,
            'result': None
        }
    )


@app.post('/migration/agenda-pendientes', response_class=HTMLResponse)
async def migration_agenda_pendientes(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    import csv
    import io

    def clean_text(value):
        if value is None:
            return ''
        value = str(value).strip()
        try:
            value = value.encode('latin1').decode('utf-8')
        except Exception:
            pass
        return value

    def pick(row, names):
        normalized = {clean_text(k).lower(): clean_text(v) for k, v in row.items()}
        for name in names:
            key = name.lower()
            if key in normalized:
                return normalized[key]
        return ''

    def parse_date(value):
        if value is None:
            return None
        if hasattr(value, 'date'):
            return value.date()

        value = clean_text(value)

        for fmt in ['%d/%m/%Y', '%Y-%m-%d', '%Y-%m-%d %H:%M:%S', '%d-%m-%Y', '%d/%m/%y']:
            try:
                return datetime.strptime(value, fmt).date()
            except Exception:
                pass

        return None

    content = await file.read()
    filename = (file.filename or '').lower()

    rows = []

    if filename.endswith('.xlsx'):
        wb = load_workbook(BytesIO(content), data_only=True)
        ws = wb.active
        headers = [clean_text(c.value) for c in ws[1]]

        for row in ws.iter_rows(min_row=2, values_only=True):
            rows.append(dict(zip(headers, row)))
    else:
        text_content = content.decode('utf-8-sig', errors='replace')
        sample = text_content[:1000]
        delimiter = ';' if sample.count(';') >= sample.count(',') else ','
        reader = csv.DictReader(io.StringIO(text_content), delimiter=delimiter)
        rows = list(reader)

    imported = 0
    skipped = 0
    created_owners = 0
    created_patients = 0

    for row in rows:
        owner_name = pick(row, ['Propietario / Cliente', 'Propietario', 'Cliente', 'Dueño', 'Responsable'])
        patient_name = pick(row, ['Paciente / Mascota', 'Paciente', 'Mascota', 'Animal'])
        phone = pick(row, ['Teléfono / WhatsApp', 'Teléfono', 'Telefono', 'WhatsApp', 'Whatsapp', 'Celular'])
        date_value = pick(row, ['Fecha', 'Fecha pendiente', 'Fecha agenda', 'Fecha turno'])
        service = pick(row, ['Motivo / Servicio', 'Motivo', 'Servicio', 'Detalle'])
        notes = pick(row, ['Notas', 'Observaciones', 'Comentario'])

        reminder_date = parse_date(date_value)

        if not reminder_date:
            skipped += 1
            continue

        if not owner_name:
            owner_name = 'Sin propietario'

        if not patient_name:
            patient_name = 'Sin paciente'

        owner = db.query(Owner).filter(Owner.name.ilike(owner_name)).first()

        if not owner:
            owner = Owner(
                name=owner_name,
                phone=phone,
                whatsapp=phone
            )
            db.add(owner)
            db.flush()
            created_owners += 1
        else:
            if phone and not owner.phone:
                owner.phone = phone
            if phone and not owner.whatsapp:
                owner.whatsapp = phone

        patient = (
            db.query(Patient)
            .filter(Patient.name.ilike(patient_name))
            .filter(Patient.owner_id == owner.id)
            .first()
        )

        if not patient:
            patient = Patient(
                name=patient_name,
                owner_id=owner.id
            )
            db.add(patient)
            db.flush()
            created_patients += 1

        service_text = f'{service} {notes}'.lower()

        if 'vacun' in service_text:
            event_type = 'Vacuna'
            title = service or 'Vacuna pendiente'
        elif 'despar' in service_text:
            event_type = 'Desparasitación'
            title = service or 'Desparasitación pendiente'
        elif 'control' in service_text:
            event_type = 'Control'
            title = service or 'Control pendiente'
        else:
            event_type = 'Consulta clínica'
            title = service or 'Pendiente importado MyVete'

        existing = (
            db.query(ClinicalEvent)
            .filter(ClinicalEvent.patient_id == patient.id)
            .filter(ClinicalEvent.reminder_date == reminder_date)
            .filter(ClinicalEvent.event_type == event_type)
            .filter(ClinicalEvent.title == title)
            .first()
        )

        if existing:
            skipped += 1
            continue

        description = (
            'Importado desde pendientes MyVete\n'
            f'Fecha: {reminder_date.strftime("%d/%m/%Y")}\n'
            f'Propietario: {owner.name}\n'
            f'Paciente: {patient.name}\n'
            f'Motivo/servicio: {service}\n'
            f'Notas: {notes}'
        )

        event = ClinicalEvent(
            patient_id=patient.id,
            event_date=datetime.combine(reminder_date, datetime.min.time()),
            event_type=event_type,
            title=title,
            description=description,
            reminder_date=reminder_date,
            created_by=user.username
        )

        db.add(event)
        imported += 1

    db.commit()

    result = {
        'type': 'agenda_pendientes',
        'imported': imported,
        'skipped': skipped,
        'created_owners': created_owners,
        'created_patients': created_patients
    }

    return templates.TemplateResponse(
        'migration.html',
        {
            'request': request,
            'result': result
        }
    )

@app.post('/migration/clientes-pacientes', response_class=HTMLResponse)
async def migration_clientes_pacientes(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    import csv
    import io

    def clean_text(value):
        if value is None:
            return ''
        value = str(value).strip()
        try:
            value = value.encode('latin1').decode('utf-8')
        except Exception:
            pass
        return value

    def pick(row, names):
        normalized = {clean_text(k).lower(): clean_text(v) for k, v in row.items()}
        for name in names:
            if name.lower() in normalized:
                return normalized[name.lower()]
        return ''

    content = await file.read()
    filename = (file.filename or '').lower()

    rows = []

    if filename.endswith('.xlsx'):
        wb = load_workbook(BytesIO(content), data_only=True)
        ws = wb.active
        headers = [clean_text(c.value) for c in ws[1]]

        for row in ws.iter_rows(min_row=2, values_only=True):
            rows.append(dict(zip(headers, row)))
    else:
        text_content = content.decode('utf-8-sig', errors='replace')
        sample = text_content[:1000]
        delimiter = ';' if sample.count(';') > sample.count(',') else ','
        reader = csv.DictReader(io.StringIO(text_content), delimiter=delimiter)
        rows = list(reader)

    created_owners = 0
    created_patients = 0
    updated_owners = 0
    skipped = 0

    for row in rows:
        owner_name = pick(row, ['Propietario', 'Cliente', 'Dueño', 'Responsable', 'Nombre propietario', 'Apellido y Nombre'])
        phone = pick(row, ['Teléfono', 'Telefono', 'Celular', 'WhatsApp', 'Whatsapp'])
        email = pick(row, ['Email', 'Mail', 'Correo'])
        address = pick(row, ['Dirección', 'Direccion', 'Domicilio'])
        patient_name = pick(row, ['Paciente', 'Mascota', 'Animal', 'Nombre paciente', 'Nombre mascota'])
        species = pick(row, ['Especie'])
        breed = pick(row, ['Raza'])
        sex = pick(row, ['Sexo'])
        color = pick(row, ['Color'])
        notes = pick(row, ['Notas', 'Observaciones', 'Comentario'])

        if not owner_name and not patient_name:
            skipped += 1
            continue

        if not owner_name:
            owner_name = 'Sin propietario'

        owner = (
            db.query(Owner)
            .filter(Owner.name.ilike(owner_name))
            .first()
        )

        if not owner:
            owner = Owner(
                name=owner_name,
                phone=phone,
                whatsapp=phone,
                email=email,
                address=address,
                notes=notes
            )
            db.add(owner)
            db.flush()
            created_owners += 1
        else:
            if phone and not owner.phone:
                owner.phone = phone
            if phone and not owner.whatsapp:
                owner.whatsapp = phone
            if email and not owner.email:
                owner.email = email
            if address and not owner.address:
                owner.address = address
            updated_owners += 1

        if patient_name:
            patient = (
                db.query(Patient)
                .filter(Patient.name.ilike(patient_name))
                .filter(Patient.owner_id == owner.id)
                .first()
            )

            if not patient:
                patient = Patient(
                    name=patient_name,
                    owner_id=owner.id,
                    species=species,
                    breed=breed,
                    sex=sex,
                    color=color,
                    notes=notes
                )
                db.add(patient)
                created_patients += 1

    db.commit()

    result = {
        'type': 'clientes_pacientes',
        'created_owners': created_owners,
        'updated_owners': updated_owners,
        'created_patients': created_patients,
        'skipped': skipped
    }
@app.post('/migration/visitas', response_class=HTMLResponse)
async def migration_visitas(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    import csv
    import io

    def clean_text(value):
        if value is None:
            return ''
        value = str(value).strip()
        try:
            value = value.encode('latin1').decode('utf-8')
        except Exception:
            pass
        return value

    def pick(row, names):
        normalized = {clean_text(k).lower(): clean_text(v) for k, v in row.items()}
        for name in names:
            key = name.lower()
            if key in normalized:
                return normalized[key]
        return ''

    def parse_datetime(value):
        value = clean_text(value)
        for fmt in [
            '%d/%m/%Y %H:%M',
            '%d/%m/%Y',
            '%Y-%m-%d %H:%M',
            '%Y-%m-%d'
        ]:
            try:
                return datetime.strptime(value, fmt)
            except Exception:
                pass
        return None

    content = await file.read()
    filename = (file.filename or '').lower()

    rows = []

    if filename.endswith('.xlsx'):
        wb = load_workbook(BytesIO(content), data_only=True)
        ws = wb.active
        headers = [clean_text(c.value) for c in ws[1]]

        for row in ws.iter_rows(min_row=2, values_only=True):
            rows.append(dict(zip(headers, row)))
    else:
        text_content = content.decode('utf-8-sig', errors='replace')
        sample = text_content[:1000]
        delimiter = ';' if sample.count(';') >= sample.count(',') else ','
        reader = csv.DictReader(io.StringIO(text_content), delimiter=delimiter)
        rows = list(reader)

    imported = 0
    skipped = 0
    created_owners = 0
    created_patients = 0

    for row in rows:
        fecha = pick(row, ['Fecha'])
        owner_name = pick(row, ['Cliente', 'Propietario', 'Dueño'])
        patient_name = pick(row, ['Paciente', 'Mascota', 'Animal'])
        email = pick(row, ['Email', 'Mail'])
        usuario = pick(row, ['Usuario'])

        event_dt = parse_datetime(fecha)

        if not event_dt:
            skipped += 1
            continue

        if not owner_name:
            owner_name = 'Sin propietario'

        if not patient_name:
            patient_name = 'Sin paciente'

        owner = (
            db.query(Owner)
            .filter(Owner.name.ilike(owner_name))
            .first()
        )

        if not owner:
            owner = Owner(
                name=owner_name,
                email=email
            )
            db.add(owner)
            db.flush()
            created_owners += 1
        else:
            if email and not owner.email:
                owner.email = email

        patient = (
            db.query(Patient)
            .filter(Patient.name.ilike(patient_name))
            .filter(Patient.owner_id == owner.id)
            .first()
        )

        if not patient:
            patient = Patient(
                name=patient_name,
                owner_id=owner.id
            )
            db.add(patient)
            db.flush()
            created_patients += 1

        existing = (
            db.query(ClinicalEvent)
            .filter(ClinicalEvent.patient_id == patient.id)
            .filter(ClinicalEvent.event_date == event_dt)
            .filter(ClinicalEvent.title == 'Visita importada MyVete')
            .first()
        )

        if existing:
            skipped += 1
            continue

        description = (
            'Importado desde Reporte de visitas MyVete\n'
            f'Fecha original: {event_dt.strftime("%d/%m/%Y %H:%M")}\n'
            f'Cliente: {owner.name}\n'
            f'Paciente: {patient.name}\n'
            f'Usuario original: {usuario}'
        )

        event = ClinicalEvent(
            patient_id=patient.id,
            event_date=event_dt,
            event_type='Consulta clínica',
            title='Visita importada MyVete',
            description=description,
            created_by=user.username
        )

        db.add(event)
        imported += 1

    db.commit()

    result = {
        'type': 'visitas',
        'imported': imported,
        'created_owners': created_owners,
        'created_patients': created_patients,
        'skipped': skipped
    }

    return templates.TemplateResponse(
        'migration.html',
        {
            'request': request,
            'result': result
        }
    )
    return templates.TemplateResponse(
        'migration.html',
        {
            'request': request,
            'result': result
        }
    )
# ==========================================================
# IMPORTADOR INTELIGENTE MYVETE
# No importa productos, stock, precios ni inventario
# ==========================================================

@app.get('/migration/myvete', response_class=HTMLResponse)
def migration_myvete_form(
    request: Request,
    user: User = Depends(require_user)
):
    html = """
    <!doctype html>
    <html>
    <head>
        <title>Importar MyVete</title>
        <style>
            body { font-family: Arial; background:#fff7fb; margin:0; padding:35px; color:#1f2937; }
            .box { max-width:760px; margin:auto; background:white; border:1px solid #f3c6d8; border-radius:22px; padding:28px; box-shadow:0 8px 24px rgba(0,0,0,.06); }
            h1 { margin-top:0; }
            .btn { background:#d85b9a; color:white; border:none; border-radius:14px; padding:13px 18px; font-weight:bold; cursor:pointer; }
            .muted { color:#6b7280; }
            .warn { background:#fff1f2; border:1px solid #fecdd3; padding:14px; border-radius:14px; margin:18px 0; }
            input[type=file] { padding:12px; border:1px solid #f3c6d8; border-radius:12px; width:100%; margin:16px 0; }
            a { color:#d85b9a; font-weight:bold; text-decoration:none; }
        </style>
    </head>
    <body>
        <div class="box">
            <h1>📥 Importador Inteligente MyVete</h1>
            <p class="muted">Subí el ZIP o los CSV exportados de MyVete.</p>

            <div class="warn">
                Este importador <b>NO importa productos, stock, precios, ventas ni inventario</b>.
                Solo importa propietarios, pacientes, historia clínica, vacunas, desparasitaciones,
                exámenes, cirugías, patologías y agenda.
            </div>

            <form action="/migration/myvete" method="post" enctype="multipart/form-data">
                <input type="file" name="file" accept=".zip,.csv" required>
                <button class="btn" type="submit">📥 Importar base MyVete</button>
            </form>

            <p style="margin-top:22px;">
                <a href="/migration">← Volver a migración</a>
            </p>
        </div>
    </body>
    </html>
    """
    return HTMLResponse(html)


@app.post('/migration/myvete', response_class=HTMLResponse)
async def migration_myvete_import(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    import csv
    import io
    import zipfile

    def clean(value):
        if value is None:
            return ''
        value = str(value).strip()
        if value.lower() in ['nan', 'none', 'null']:
            return ''
        return value

    def normalize_name(value):
        return clean(value).lower().replace(' ', '').replace('-', '').replace('_', '')

    def read_csv_bytes(content):
        text = None
        for enc in ['utf-8-sig', 'latin1']:
            try:
                text = content.decode(enc)
                break
            except Exception:
                pass

        if text is None:
            text = content.decode('utf-8', errors='replace')

        sample = text[:2000]
        delimiter = ';' if sample.count(';') >= sample.count(',') else ','

        reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
        return [
            {clean(k): clean(v) for k, v in row.items()}
            for row in reader
        ]

    def parse_date(value):
        value = clean(value)

        if not value:
            return None

        for fmt in [
            '%Y-%m-%d %H:%M:%S',
            '%Y-%m-%d',
            '%d/%m/%Y %H:%M',
            '%d/%m/%Y',
            '%d-%m-%Y',
            '%d/%m/%y'
        ]:
            try:
                dt = datetime.strptime(value, fmt)
                return dt.date()
            except Exception:
                pass

        return None

    def parse_datetime(value):
        value = clean(value)

        if not value:
            return None

        for fmt in [
            '%Y-%m-%d %H:%M:%S',
            '%Y-%m-%d %H:%M',
            '%Y-%m-%d',
            '%d/%m/%Y %H:%M',
            '%d/%m/%Y',
            '%d-%m-%Y',
            '%d/%m/%y'
        ]:
            try:
                dt = datetime.strptime(value, fmt)
                if fmt in ['%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y', '%d/%m/%y']:
                    return datetime.combine(dt.date(), datetime.min.time())
                return dt
            except Exception:
                pass

        return None

    def to_float(value):
        try:
            value = clean(value).replace(',', '.')
            return float(value) if value else None
        except Exception:
            return None

    def pick(row, *names):
        normalized = {
            normalize_name(k): clean(v)
            for k, v in row.items()
        }

        for name in names:
            key = normalize_name(name)
            if key in normalized:
                return normalized[key]

        return ''

    def marker_exists(model, field, marker):
        return (
            db.query(model)
            .filter(field.ilike(f'%{marker}%'))
            .first()
        )

    content = await file.read()
    filename = clean(file.filename).lower()

    raw_files = {}

    if filename.endswith('.zip'):
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            for name in z.namelist():
                low = name.lower()
                if not low.endswith('.csv'):
                    continue
                if (
                    'product' in low
                    or 'stock' in low
                    or 'venta' in low
                    or 'sale' in low
                    or 'precio' in low
                    or 'price' in low
                    or 'inventario' in low
                    or 'inventory' in low
                    or 'caja' in low
                    or 'cash' in low
                ):
                    continue
                raw_files[low] = z.read(name)
    else:
        low = filename
        blocked_names = [
            'product', 'stock', 'venta', 'sale', 'precio', 'price',
            'inventario', 'inventory', 'caja', 'cash'
        ]
        
        if not any(blocked in low for blocked in blocked_names):
            raw_files[low] = content

    detected = {
        'clientes': [],
        'pacientes': [],
        'hc': [],
        'vacunas': [],
        'antiparasitarios': [],
        'attr': [],
        'agenda': [],
        'examenes': [],
        'cirugias': [],
        'patologias': []
    }

    for name, data in raw_files.items():
        rows = read_csv_bytes(data)

        if 'clientes' in name:
            detected['clientes'] = rows
        elif 'pacientes' in name:
            detected['pacientes'] = rows
        elif 'hc-vacunas' in name:
            detected['vacunas'] = rows
        elif 'hc-antiparasitarios' in name:
            detected['antiparasitarios'] = rows
        elif 'hc-attr' in name:
            detected['attr'] = rows
        elif 'hc-agenda' in name:
            detected['agenda'] = rows
        elif 'hc-examenes' in name:
            detected['examenes'] = rows
        elif 'hc-cirugias' in name:
            detected['cirugias'] = rows
        elif 'hc-patologias' in name:
            detected['patologias'] = rows
        elif name.endswith('hc.csv') or 'exp-hc.csv' in name:
            detected['hc'] = rows

    stats = {
        'owners_created': 0,
        'owners_skipped': 0,
        'patients_created': 0,
        'patients_skipped': 0,
        'events_created': 0,
        'events_skipped': 0,
        'appointments_created': 0,
        'appointments_skipped': 0,
        'pathologies_added': 0,
        'ignored_products_stock': True
    }

    owner_map = {}

    for row in detected['clientes']:
        myvete_id = pick(row, 'id')
        marker = f'MYVETE_CLIENT_ID:{myvete_id}'

        existing = marker_exists(Owner, Owner.notes, marker) if myvete_id else None

        first_name = pick(row, 'nombres')
        last_name = pick(row, 'apellido')
        owner_name = clean(f'{first_name} {last_name}').strip()

        if not owner_name:
            owner_name = pick(row, 'nombre', 'cliente', 'propietario') or 'Sin propietario'

        phone = pick(row, 'telMovil', 'telefono', 'teléfono', 'celular')
        email = pick(row, 'email')
        address = pick(row, 'direccion', 'dirección', 'domicilio')
        comment = pick(row, 'comentario')

        if not existing:
            existing = (
                db.query(Owner)
                .filter(Owner.name.ilike(owner_name))
                .first()
            )

        if existing:
            if phone and not existing.phone:
                existing.phone = phone
            if phone and not existing.whatsapp:
                existing.whatsapp = phone
            if email and not existing.email:
                existing.email = email
            if address and not existing.address:
                existing.address = address

            if marker and marker not in (existing.notes or ''):
                existing.notes = ((existing.notes or '') + f'\n{marker}').strip()

            stats['owners_skipped'] += 1
            owner = existing
        else:
            owner = Owner(
                name=owner_name,
                phone=phone,
                whatsapp=phone,
                email=email,
                address=address,
                notes=f'{marker}\nImportado desde MyVete\n{comment}'.strip()
            )
            db.add(owner)
            db.flush()
            stats['owners_created'] += 1

        if myvete_id:
            owner_map[myvete_id] = owner.id

    patient_map = {}

    for row in detected['pacientes']:
        myvete_id = pick(row, 'id')
        owner_mvv_id = pick(row, 'idCliente')
        owner_id = owner_map.get(owner_mvv_id)

        marker = f'MYVETE_PATIENT_ID:{myvete_id}'
        existing = marker_exists(Patient, Patient.notes, marker) if myvete_id else None

        patient_name = pick(row, 'nombre') or 'Sin nombre'
        species = pick(row, 'especie')
        breed = pick(row, 'raza')
        sex = pick(row, 'sexo')
        color = pick(row, 'color')
        birth = parse_date(pick(row, 'fechaNacimiento'))
        comment = pick(row, 'comentario')
       

        if not owner_id:
            owner = Owner(
                name='Sin propietario MyVete',
                notes='Creado automáticamente durante importación MyVete'
            )
            db.add(owner)
            db.flush()
            owner_id = owner.id

        if not existing:
            existing = (
                db.query(Patient)
                .filter(Patient.name.ilike(patient_name))
                .filter(Patient.owner_id == owner_id)
                .first()
            )

        if existing:
            if marker and marker not in (existing.notes or ''):
                existing.notes = ((existing.notes or '') + f'\n{marker}').strip()
            stats['patients_skipped'] += 1
            patient = existing
        else:
            patient = Patient(
                name=patient_name,
                owner_id=owner_id,
                species=species,
                breed=breed,
                sex=sex,
                birth_date=birth,
                color=color,
 
                notes=f'{marker}\nImportado desde MyVete\n{comment}'.strip()
            )
            db.add(patient)
            db.flush()
            stats['patients_created'] += 1

        if myvete_id:
            patient_map[myvete_id] = patient.id

    vaccines_by_hc = {}
    for row in detected['vacunas']:
        idhc = pick(row, 'idHC')
        vaccines_by_hc.setdefault(idhc, []).append(row)

    antip_by_hc = {}
    for row in detected['antiparasitarios']:
        idhc = pick(row, 'idHC')
        antip_by_hc.setdefault(idhc, []).append(row)

    attr_by_hc = {}
    for row in detected['attr']:
        idhc = pick(row, 'idHC')
        attr_by_hc.setdefault(idhc, []).append(row)

    exams_by_hc = {}
    for row in detected['examenes']:
        idhc = pick(row, 'idHC')
        exams_by_hc.setdefault(idhc, []).append(row)

    surgeries_by_hc = {}
    for row in detected['cirugias']:
        idhc = pick(row, 'idHC')
        surgeries_by_hc.setdefault(idhc, []).append(row)

    pathologies_by_hc = {}
    for row in detected['patologias']:
        idhc = pick(row, 'idHC')
        pathologies_by_hc.setdefault(idhc, []).append(row)

    def event_type_from_exam(name):
        text = clean(name).lower()
        if 'rx' in text or 'radio' in text:
            return 'Radiografía'
        if 'eco' in text:
            return 'Ecografía'
        if 'ecg' in text or 'electro' in text:
            return 'ECG'
        if 'lab' in text or 'sangre' in text or 'orina' in text:
            return 'Laboratorio'
        return 'Otro procedimiento'

    for row in detected['hc']:
        hc_id = pick(row, 'id')
        patient_mvv_id = pick(row, 'idPaciente')
        patient_id = patient_map.get(patient_mvv_id)

        if not patient_id:
            stats['events_skipped'] += 1
            continue

        event_dt = parse_datetime(pick(row, 'fechaCreacion')) or argentina_now()
        description_original = pick(row, 'descripcion')
        usuario = pick(row, 'usuario')

        marker = f'MYVETE_HC_ID:{hc_id}'

        existing = marker_exists(ClinicalEvent, ClinicalEvent.description, marker) if hc_id else None

        if existing:
            stats['events_skipped'] += 1
            continue

        weight = None
        for attr in attr_by_hc.get(hc_id, []):
            if 'peso' in pick(attr, 'nombreParam', 'descParam').lower():
                weight = to_float(pick(attr, 'valor'))

        sections = []

        sections.append(f'Importado desde MyVete\n{marker}')

        if description_original:
            sections.append(f'🩺 Registro original\n{description_original}')

        if weight is not None:
            sections.append(f'⚖️ Peso registrado\n{weight} kg')

        if vaccines_by_hc.get(hc_id):
            lines = []
            for v in vaccines_by_hc[hc_id]:
                lines.append(f"- {pick(v, 'vacuna') or 'Vacuna'} {pick(v, 'comentario')}")
            sections.append('💉 Vacunas asociadas\n' + '\n'.join(lines))

        if antip_by_hc.get(hc_id):
            lines = []
            for a in antip_by_hc[hc_id]:
                lines.append(f"- {pick(a, 'nombreATP') or 'Antiparasitario'} {pick(a, 'comentario')}")
            sections.append('🪱 Antiparasitarios asociados\n' + '\n'.join(lines))

        if exams_by_hc.get(hc_id):
            lines = []
            for ex in exams_by_hc[hc_id]:
                lines.append(f"- {pick(ex, 'nombreExamen') or 'Examen'} {pick(ex, 'comentarioHCExamen')}")
            sections.append('📎 Estudios / exámenes\n' + '\n'.join(lines))

        if surgeries_by_hc.get(hc_id):
            lines = []
            for sx in surgeries_by_hc[hc_id]:
                lines.append(f"- {pick(sx, 'nombreCirugia') or 'Cirugía'} {pick(sx, 'comentarioCirugia')}")
            sections.append('🔪 Cirugías asociadas\n' + '\n'.join(lines))

        if pathologies_by_hc.get(hc_id):
            lines = []
            for pa in pathologies_by_hc[hc_id]:
                lines.append(f"- {pick(pa, 'patologia') or 'Patología'} {pick(pa, 'comentario')}")
            sections.append('⚠️ Patologías asociadas\n' + '\n'.join(lines))

            patient = db.get(Patient, patient_id)
            if patient:
                existing_alerts = patient.alerts or ''
                for pa in pathologies_by_hc[hc_id]:
                    pathology_text = pick(pa, 'patologia')
                    if pathology_text and pathology_text not in existing_alerts:
                        existing_alerts = (existing_alerts + f'\n{pathology_text}').strip()
                        stats['pathologies_added'] += 1
                patient.alerts = existing_alerts

        if usuario:
            sections.append(f'👤 Usuario MyVete\n{usuario}')

        has_only_preventive = (
            not description_original
            and (
                vaccines_by_hc.get(hc_id)
                or antip_by_hc.get(hc_id)
            )
            and not exams_by_hc.get(hc_id)
            and not surgeries_by_hc.get(hc_id)
            and not pathologies_by_hc.get(hc_id)
        )

        if not has_only_preventive:
            event = ClinicalEvent(
                patient_id=patient_id,
                event_date=event_dt,
                event_type='Consulta clínica',
                title='Consulta importada MyVete',
                description='\n\n'.join(sections),
                weight=weight,
                created_by=user.username
            )
            db.add(event)
            stats['events_created'] += 1

        if weight is not None:
            patient = db.get(Patient, patient_id)
            if patient:
                patient.weight = weight

        for v in vaccines_by_hc.get(hc_id, []):
            vaccine_name = pick(v, 'vacuna') or 'Vacuna'
            vaccine_marker = f'MYVETE_HC_ID:{hc_id}:VACUNA:{vaccine_name}'
            if marker_exists(ClinicalEvent, ClinicalEvent.description, vaccine_marker):
                continue

            db.add(ClinicalEvent(
                patient_id=patient_id,
                event_date=event_dt,
                event_type='Vacuna',
                title=vaccine_name,
                description=f'💉 Vacuna importada desde MyVete\n{vaccine_marker}\nComentario: {pick(v, "comentario") or "-"}',
                vaccine_name=vaccine_name,
                created_by=user.username
            ))
            stats['events_created'] += 1

        for a in antip_by_hc.get(hc_id, []):
            antip_name = pick(a, 'nombreATP') or 'Antiparasitario'
            antip_marker = f'MYVETE_HC_ID:{hc_id}:ANTIP:{antip_name}'
            if marker_exists(ClinicalEvent, ClinicalEvent.description, antip_marker):
                continue

            db.add(ClinicalEvent(
                patient_id=patient_id,
                event_date=event_dt,
                event_type='Desparasitación',
                title=antip_name,
                description=f'🪱 Desparasitación importada desde MyVete\n{antip_marker}\nComentario: {pick(a, "comentario") or "-"}',
                dewormer_product=antip_name,
                created_by=user.username
            ))
            stats['events_created'] += 1

        for sx in surgeries_by_hc.get(hc_id, []):
            sx_name = pick(sx, 'nombreCirugia') or 'Cirugía'
            sx_marker = f'MYVETE_HC_ID:{hc_id}:CIRUGIA:{sx_name}'
            if marker_exists(ClinicalEvent, ClinicalEvent.description, sx_marker):
                continue

            db.add(ClinicalEvent(
                patient_id=patient_id,
                event_date=event_dt,
                event_type='Cirugía',
                title=sx_name,
                description=f'🔪 Cirugía importada desde MyVete\n{sx_marker}\nComentario: {pick(sx, "comentarioCirugia") or "-"}',
                created_by=user.username
            ))
            stats['events_created'] += 1

        for ex in exams_by_hc.get(hc_id, []):
            ex_name = pick(ex, 'nombreExamen') or 'Examen'
            ex_marker = f'MYVETE_HC_ID:{hc_id}:EXAMEN:{ex_name}'
            if marker_exists(ClinicalEvent, ClinicalEvent.description, ex_marker):
                continue

            db.add(ClinicalEvent(
                patient_id=patient_id,
                event_date=event_dt,
                event_type=event_type_from_exam(ex_name),
                title=ex_name,
                description=f'📎 Examen importado desde MyVete\n{ex_marker}\nComentario: {pick(ex, "comentarioHCExamen") or "-"}',
                created_by=user.username
            ))
            stats['events_created'] += 1

    for row in detected['agenda']:
        patient_mvv_id = pick(row, 'idPaciente')
        patient_id = patient_map.get(patient_mvv_id)

        start_dt = parse_datetime(pick(row, 'fechaInicioTurno'))
        end_dt = parse_datetime(pick(row, 'fechaFinTurno'))

        if not start_dt:
            stats['appointments_skipped'] += 1
            continue

        event_name = pick(row, 'evento')
        service = (
            pick(row, 'nombreServicio')
            or pick(row, 'nombreVacuna')
            or pick(row, 'nombreATP')
            or pick(row, 'nombreExamen')
            or event_name
            or 'Turno importado MyVete'
        )

        comment = pick(row, 'comentario')
        user_mvv = pick(row, 'usuario')

        existing = (
            db.query(Appointment)
            .filter(Appointment.patient_id == patient_id)
            .filter(Appointment.appointment_date == datetime.combine(start_dt.date(), datetime.min.time()))
            .filter(Appointment.service == service)
            .first()
        )

        if existing:
            stats['appointments_skipped'] += 1
            continue

        appointment = Appointment(
            appointment_date=datetime.combine(start_dt.date(), datetime.min.time()),
            start_time=start_dt.strftime('%H:%M') if start_dt else '',
            end_time=end_dt.strftime('%H:%M') if end_dt else '',
            patient_id=patient_id,
            service=service,
            title=event_name or service,
            notes=f'Importado desde MyVete\nComentario: {comment or "-"}\nUsuario: {user_mvv or "-"}',
            status='Pendiente'
        )

        db.add(appointment)
        stats['appointments_created'] += 1

    db.commit()

    html = f"""
    <!doctype html>
    <html>
    <head>
        <title>Resultado importación MyVete</title>
        <style>
            body {{ font-family: Arial; background:#fff7fb; margin:0; padding:35px; color:#1f2937; }}
            .box {{ max-width:820px; margin:auto; background:white; border:1px solid #f3c6d8; border-radius:22px; padding:28px; box-shadow:0 8px 24px rgba(0,0,0,.06); }}
            h1 {{ margin-top:0; }}
            .grid {{ display:grid; grid-template-columns:1fr 1fr; gap:12px; }}
            .item {{ background:#fff7fb; border:1px solid #f3c6d8; border-radius:14px; padding:14px; }}
            .num {{ font-size:28px; font-weight:bold; }}
            a {{ color:#d85b9a; font-weight:bold; text-decoration:none; }}
        </style>
    </head>
    <body>
        <div class="box">
            <h1>✅ Importación MyVete finalizada</h1>

            <div class="grid">
                <div class="item"><div class="num">{stats['owners_created']}</div>Propietarios creados</div>
                <div class="item"><div class="num">{stats['owners_skipped']}</div>Propietarios existentes/actualizados</div>

                <div class="item"><div class="num">{stats['patients_created']}</div>Pacientes creados</div>
                <div class="item"><div class="num">{stats['patients_skipped']}</div>Pacientes existentes/actualizados</div>

                <div class="item"><div class="num">{stats['events_created']}</div>Eventos clínicos creados</div>
                <div class="item"><div class="num">{stats['events_skipped']}</div>Eventos duplicados/salteados</div>

                <div class="item"><div class="num">{stats['appointments_created']}</div>Turnos importados</div>
                <div class="item"><div class="num">{stats['appointments_skipped']}</div>Turnos salteados</div>

                <div class="item"><div class="num">{stats['pathologies_added']}</div>Patologías agregadas a alertas</div>
                <div class="item"><div class="num">0</div>Productos / stock importados</div>
            </div>

            <p style="margin-top:24px;">
                <a href="/migration/myvete">Importar otra base</a> ·
                <a href="/search">Buscar pacientes</a> ·
                <a href="/">Volver al inicio</a>
            </p>
        </div>
    </body>
    </html>
    """

    return HTMLResponse(html)
@app.post('/attachment/{attachment_id}/delete')
def delete_attachment(
    attachment_id: int,
    db: Session = Depends(get_db)
):
    attachment = db.query(EventAttachment).filter(
        EventAttachment.id == attachment_id
    ).first()

    if not attachment:
        raise HTTPException(status_code=404, detail="Adjunto no encontrado")

    patient_id = attachment.event.patient_id

    try:
        path_parts = attachment.file_path.split("/storage/v1/object/public/adjuntos/")
        if len(path_parts) > 1:
            storage_path = path_parts[1]
            supabase.storage.from_("adjuntos").remove([storage_path])
    except Exception:
        pass

    db.delete(attachment)
    db.commit()

    return RedirectResponse(f"/patients/{patient_id}", status_code=303)
@app.get('/pendientes', response_class=HTMLResponse)
def pendientes(
    request: Request,
    show_all: int = 0,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    today = argentina_now().date()

    from calendar import monthrange

    last_day = monthrange(today.year, today.month)[1]
    end_of_month = today.replace(day=last_day)

    # ============================================================
    # TURNOS MYVETE: SOLO HOY Y PRÓXIMOS DEL MES ACTUAL
    # ============================================================

    month_start = datetime.combine(
        today,
        datetime.min.time()
    )

    month_end = datetime.combine(
        end_of_month,
        datetime.max.time()
    )

    # Solamente se buscan turnos importados desde hoy hasta
    # el último día del mes actual.
    myvete_appointments = (
        db.query(Appointment)
        .filter(Appointment.status == 'Pendiente')
        .filter(
            Appointment.notes.ilike(
                '%Importado desde MyVete%'
            )
        )
        .filter(Appointment.patient_id != None)
        .filter(
            Appointment.appointment_date >= month_start
        )
        .filter(
            Appointment.appointment_date <= month_end
        )
        .all()
    )

    # Se recuperan todos los pendientes MyVete ya creados
    # en una sola consulta, evitando una consulta por turno.
    existing_myvete_events = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.description.ilike(
                '%[MYVETE_APPOINTMENT_ID:%'
            )
        )
        .filter(
            ClinicalEvent.reminder_date >= today
        )
        .filter(
            ClinicalEvent.reminder_date <= end_of_month
        )
        .all()
    )

    existing_appointment_ids = set()

    for existing_event in existing_myvete_events:
        existing_description = (
            existing_event.description or ''
        )

        marker_start = (
            '[MYVETE_APPOINTMENT_ID:'
        )

        if marker_start not in existing_description:
            continue

        try:
            imported_id = (
                existing_description
                .split(marker_start, 1)[1]
                .split(']', 1)[0]
            )

            existing_appointment_ids.add(
                int(imported_id)
            )
        except (ValueError, IndexError):
            continue

    new_myvete_events = []

    for appointment in myvete_appointments:
        if appointment.id in existing_appointment_ids:
            continue

        appointment_day = appointment.appointment_date

        if isinstance(appointment_day, datetime):
            reminder_day = appointment_day.date()
        else:
            reminder_day = appointment_day

        if not reminder_day:
            continue

        service_text = (
            appointment.service
            or appointment.title
            or 'Turno pendiente'
        )

        service_lower = service_text.lower()

        if 'vacun' in service_lower:
            event_type = 'Vacuna'
        elif 'despar' in service_lower:
            event_type = 'Desparasitación'
        elif 'control' in service_lower:
            event_type = 'Control'
        else:
            event_type = 'Consulta clínica'

        description_parts = [
            DUE_ACTIVE_MARKER,
            (
                f'[MYVETE_APPOINTMENT_ID:'
                f'{appointment.id}]'
            ),
            'Importado desde agenda MyVete'
        ]

        if appointment.start_time:
            description_parts.append(
                f'Horario: {appointment.start_time}'
            )

        if appointment.notes:
            description_parts.append(
                str(appointment.notes)
            )

        new_myvete_events.append(
            ClinicalEvent(
                patient_id=appointment.patient_id,
                event_date=datetime.combine(
                    reminder_day,
                    datetime.min.time()
                ),
                event_type=event_type,
                title=service_text,
                description='\n'.join(
                    description_parts
                ),
                reminder_date=reminder_day,
                created_by=user.username
            )
        )

    if new_myvete_events:
        db.add_all(new_myvete_events)
        db.commit()
    all_events = (
        db.query(ClinicalEvent)
        .filter(
            or_(
                ClinicalEvent.reminder_date != None,
                ClinicalEvent.description.ilike(
                    f'%{DUE_ACTIVE_MARKER}%'
                )
            )
        )
        .filter(
            ~ClinicalEvent.description.ilike(
                f'%{DUE_CLOSED_MARKER}%'
            )
        )
        .filter(
            or_(
                # Los pendientes normales de Aromos Cloud
                # conservan el funcionamiento actual.
                ~ClinicalEvent.description.ilike(
                    '%[MYVETE_APPOINTMENT_ID:%'
                ),

                # Los importados de MyVete solamente aparecen
                # desde hoy hasta el final del mes actual.
                and_(
                    ClinicalEvent.description.ilike(
                        '%[MYVETE_APPOINTMENT_ID:%'
                    ),
                    ClinicalEvent.reminder_date >= today,
                    ClinicalEvent.reminder_date <= end_of_month
                )
            )
        )
        .order_by(
            ClinicalEvent.reminder_date.asc().nullslast(),
            ClinicalEvent.id.asc()
        )
        .all()
    )

    overdue_events = []
    today_events = []
    upcoming_events = []
    notified_events = []

    for event in all_events:
        description = event.description or ''

        is_notified = (
            WP_SENT_MARKER in description
            and DUE_ACTIVE_MARKER in description
        )

        if is_notified:
            notified_events.append(
                extract_due_metadata(event)
            )
            continue

        if not event.reminder_date:
            continue

        if event.reminder_date < today:
            overdue_events.append(event)

        elif event.reminder_date == today:
            today_events.append(event)

        elif show_all or event.reminder_date <= end_of_month:
            upcoming_events.append(event)

    eventos = (
        overdue_events
        + today_events
        + upcoming_events
    )

    stats = {
        'total': len(eventos) + len(notified_events),
        'overdue': len(overdue_events),
        'today': len(today_events),
        'upcoming': len(upcoming_events),
        'notified': len(notified_events)
    }

    return templates.TemplateResponse(
        'pendientes.html',
        {
            'request': request,
            'eventos': eventos,
            'overdue_events': overdue_events,
            'today_events': today_events,
            'upcoming_events': upcoming_events,
            'notified_events': notified_events,
            'stats': stats,
            'today': today,
            'pending_count': stats['total'],
            'show_all': show_all
        }
    )


@app.get('/patients/{patient_id}/quick-pendiente')
def quick_pendiente(
    patient_id: int,
    type: str,
    days: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(
            status_code=404,
            detail='Paciente no encontrado'
        )

    reminder_date = (
        argentina_now().date()
        + timedelta(days=days)
    )

    existing_event = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .filter(ClinicalEvent.event_type == type)
        .filter(ClinicalEvent.title == type)
        .filter(
            ClinicalEvent.reminder_date == reminder_date
        )
        .first()
    )

    if existing_event:
        return RedirectResponse(
            f'/patients/{patient.id}/v2',
            status_code=303
        )

    event = ClinicalEvent(
        patient_id=patient.id,
        event_date=argentina_now(),
        event_type=type,
        title=type,
        description=DUE_ACTIVE_MARKER,
        reminder_date=reminder_date,
        created_by=user.username
    )

    db.add(event)
    db.commit()

    return RedirectResponse(
        f'/patients/{patient.id}/v2',
        status_code=303
    )


@app.post('/events/{event_id}/done')
def event_done(
    event_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    event = db.get(ClinicalEvent, event_id)

    if not event:
        raise HTTPException(
            status_code=404,
            detail='Evento no encontrado'
        )

    mark_due_as_whatsapp_sent(event, user)

    db.commit()

    return RedirectResponse(
        '/pendientes',
        status_code=303
    )

@app.post('/events/{event_id}/complete-pending')
def complete_pending_event(
    event_id: int,
    return_to: str = Form('pendientes'),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    event = db.get(ClinicalEvent, event_id)

    if not event:
        raise HTTPException(
            status_code=404,
            detail='Evento no encontrado'
        )

    description = event.description or ''

    if DUE_CLOSED_MARKER not in description:
        if DUE_ACTIVE_MARKER in description:
            description = description.replace(
                DUE_ACTIVE_MARKER,
                DUE_CLOSED_MARKER
            )
        else:
            description = (
                description.rstrip()
                + '\n\n'
                + DUE_CLOSED_MARKER
            ).strip()

        description += (
            '\n\n'
            '✅ Pendiente marcado manualmente como realizado.\n'
            f'Fecha real: {argentina_now().strftime("%d/%m/%Y %H:%M")}\n'
            f'Registrado por: {user.username}.'
        )

        event.description = description
        event.reminder_date = None
        db.commit()

    if return_to == 'patient':
        destination = f'/patients/{event.patient_id}/v2'
    else:
        destination = '/pendientes'

    return RedirectResponse(
        destination,
        status_code=303
    )
@app.post('/events/{event_id}/delete')
def delete_event(
    event_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    event = db.get(ClinicalEvent, event_id)

    if not event:
        raise HTTPException(
            status_code=404,
            detail='Evento no encontrado'
        )

    patient_id = event.patient_id

    db.delete(event)
    db.commit()

    return RedirectResponse(
        f'/patients/{patient_id}/v2',
        status_code=303
    )


@app.get('/events/{event_id}/whatsapp')
def event_whatsapp(
    event_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    event = db.get(ClinicalEvent, event_id)

    if not event:
        raise HTTPException(
            status_code=404,
            detail='Evento no encontrado'
        )

    patient = event.patient
    owner = patient.owner if patient else None

    if not patient or not owner:
        return RedirectResponse(
            '/pendientes',
            status_code=303
        )

    number = (
        owner.whatsapp
        or owner.phone
        or ''
    ).strip()

    if not number:
        return RedirectResponse(
            f'/patients/{patient.id}/v2',
            status_code=303
        )

    number = ''.join(
        character
        for character in number
        if character.isdigit()
    )

    if number.startswith('0'):
        number = number[1:]

    if not number.startswith('54'):
        number = '549' + number

    title_text = event.title or event.event_type or 'recordatorio'

    if (
        event.event_type == 'Vacuna'
        or 'vacun' in title_text.lower()
    ):
        message = (
            f'Buenos días *{owner.name}*.\n\n'
            f'Te recordamos que *{patient.name}* debe recibir '
            f'su *{title_text}*.\n\n'
            f'¡Los esperamos!'
        )

    elif (
        event.event_type == 'Desparasitación'
        or 'despar' in title_text.lower()
        or 'antiparas' in title_text.lower()
        or 'aprax' in title_text.lower()
        or 'pipeta' in title_text.lower()
    ):
        message = (
            f'Buenos días *{owner.name}*.\n\n'
            f'Te recordamos que *{patient.name}* debe '
            f'desparasitarse.\n\n'
            f'Pueden traerlo a la clínica, retirar la '
            f'medicación o solicitar envío a domicilio.\n\n'
            f'¿Qué preferís?'
        )

    else:
        message = (
            f'Hola *{owner.name}*.\n\n'
            f'Te recordamos que *{patient.name}* tiene '
            f'pendiente: *{title_text}*.\n\n'
            f'Clínica Veterinaria Los Aromos.'
        )

    mark_due_as_whatsapp_sent(event, user)

    db.commit()

    import urllib.parse

    url = (
        f'https://wa.me/{number}'
        f'?text={urllib.parse.quote(message)}'
    )

    return RedirectResponse(
        url,
        status_code=303
    )
@app.get('/patients/{patient_id}/history', response_class=HTMLResponse)
def patient_history(
    request: Request,
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404)

    events = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .order_by(ClinicalEvent.event_date.desc())
        .all()
    )

    return templates.TemplateResponse(
    'patient_history.html',
    {
        'request': request,
        'patient': patient,
        'events': events,
        'timedelta': timedelta
    }
)
@app.get('/patients/{patient_id}/print', response_class=HTMLResponse)
def patient_print(
    request: Request,
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")

    events = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .order_by(ClinicalEvent.event_date.asc(), ClinicalEvent.id.asc())
        .all()
    )

    return templates.TemplateResponse(
        'patient_print.html',
        {
            'request': request,
            'patient': patient,
            'events': events
        }
    )
def calculate_anesthesia_doses(weight: float):
    try:
        weight = float(weight)
    except:
        weight = 0

    return {
        "acepromacina_mg": round(weight * 0.03, 2),
        "acepromacina_ml": round((weight * 0.03) / 1, 2),
        "dexmedetomidina_mcg_baja": round(weight * 5, 2),
        "dexmedetomidina_mcg_alta": round(weight * 10, 2),
        "dexmedetomidina_ml_baja": round((weight * 5 / 1000) / 0.5, 2),
        "dexmedetomidina_ml_alta": round((weight * 10 / 1000) / 0.5, 2),
        "ketamina_premed_mg": round(weight * 3, 2),
        "ketamina_premed_ml": round((weight * 3) / 50, 2),
        "midazolam_mg": round(weight * 0.2, 2),
        "midazolam_ml": round((weight * 0.2) / 5, 2),

        "nalbufina_mg": round(weight * 0.5, 2),
        "nalbufina_ml": round((weight * 0.5) / 10, 2),

        "butorfanol_mg": round(weight * 0.3, 2),
        "butorfanol_ml": round((weight * 0.3) / 10, 2),

        "propofol_ind_mg": round(weight * 3, 2),
        "propofol_ind_ml": round((weight * 3) / 10, 2),

        "ketamina_ind_mg": round(weight * 5, 2),
        "ketamina_ind_ml": round((weight * 5) / 50, 2),

        "midazolam_ind_mg": round(weight * 0.2, 2),
        "midazolam_ind_ml": round((weight * 0.2) / 5, 2),

        "meloxicam_mg": round(weight * 0.2, 2),
        "meloxicam_ml": round((weight * 0.2) / 5, 2),

        "fentanilo_bolo_mcg": round(weight * 2, 2),
        "fentanilo_bolo_ml": round((weight * 2) / 50, 2),

        "flk_velocidad_ml_h": round(weight * 3, 2),
"flk_macro_gtt_min": round((weight * 3 * 20) / 60, 1),
"flk_micro_gtt_min": round((weight * 3 * 60) / 60, 1),

"flk_fentanilo_mcg_h": round(weight * 5, 2),
"flk_lidocaina_mg_h": round(weight * 0.6, 2),
"flk_ketamina_mg_h": round(weight * 0.6, 2),

        # FLK:
        # Fentanilo 5 mcg/kg/h
        # Lidocaína 0.6 mg/kg/h
        # Ketamina 0.6 mg/kg/h
        # Velocidad final 3 ml/kg/h
        #
        # Al mantener estas dosis y esta velocidad, la concentración
        # requerida en la bolsa es independiente del peso del paciente.

        "flk_fentanilo_ml_100": round(
            (100 * 5 / 3) / 50,
            2
        ),
        "flk_lidocaina_ml_100": round(
            (100 * 0.6 / 3) / 20,
            2
        ),
        "flk_ketamina_ml_100": round(
            (100 * 0.6 / 3) / 50,
            2
        ),
        "flk_total_drogas_ml_100": round(
            ((100 * 5 / 3) / 50)
            + ((100 * 0.6 / 3) / 20)
            + ((100 * 0.6 / 3) / 50),
            2
        ),
        "flk_ssf_ml_100": round(
            100
            - (
                ((100 * 5 / 3) / 50)
                + ((100 * 0.6 / 3) / 20)
                + ((100 * 0.6 / 3) / 50)
            ),
            2
        ),

        "flk_fentanilo_ml_200": round(
            (200 * 5 / 3) / 50,
            2
        ),
        "flk_lidocaina_ml_200": round(
            (200 * 0.6 / 3) / 20,
            2
        ),
        "flk_ketamina_ml_200": round(
            (200 * 0.6 / 3) / 50,
            2
        ),
        "flk_total_drogas_ml_200": round(
            ((200 * 5 / 3) / 50)
            + ((200 * 0.6 / 3) / 20)
            + ((200 * 0.6 / 3) / 50),
            2
        ),
        "flk_ssf_ml_200": round(
            200
            - (
                ((200 * 5 / 3) / 50)
                + ((200 * 0.6 / 3) / 20)
                + ((200 * 0.6 / 3) / 50)
            ),
            2
        ),

        "flk_fentanilo_ml_250": round(
            (250 * 5 / 3) / 50,
            2
        ),
        "flk_lidocaina_ml_250": round(
            (250 * 0.6 / 3) / 20,
            2
        ),
        "flk_ketamina_ml_250": round(
            (250 * 0.6 / 3) / 50,
            2
        ),
        "flk_total_drogas_ml_250": round(
            ((250 * 5 / 3) / 50)
            + ((250 * 0.6 / 3) / 20)
            + ((250 * 0.6 / 3) / 50),
            2
        ),
        "flk_ssf_ml_250": round(
            250
            - (
                ((250 * 5 / 3) / 50)
                + ((250 * 0.6 / 3) / 20)
                + ((250 * 0.6 / 3) / 50)
            ),
            2
        ),
"rl_ml_h": round(weight * 5, 2),
"rl_macro_gtt_min": round((weight * 5 * 20) / 60, 1),
"rl_micro_gtt_min": round((weight * 5 * 60) / 60, 1),

"dobutamina_mcg_min_baja": round(weight * 5, 2),
"dobutamina_mcg_min_alta": round(weight * 10, 2),
"dobutamina_ml_h_baja": round(((weight * 5 * 60) / 1000) / 12.5, 2),
"dobutamina_ml_h_alta": round(((weight * 10 * 60) / 1000) / 12.5, 2),
"dobutamina_macro_gtt_min_baja": round(((((weight * 5 * 60) / 1000) / 12.5) * 20) / 60, 1),
"dobutamina_macro_gtt_min_alta": round(((((weight * 10 * 60) / 1000) / 12.5) * 20) / 60, 1),
"dobutamina_micro_gtt_min_baja": round(((((weight * 5 * 60) / 1000) / 12.5) * 60) / 60, 1),
"dobutamina_micro_gtt_min_alta": round(((((weight * 10 * 60) / 1000) / 12.5) * 60) / 60, 1),
"dobutamina_macro_seg_gota_baja": round(
    60 / ((((weight * 5 * 60) / 1000) / 12.5) * 20 / 60),
    0
) if weight > 0 else 0,

"dobutamina_macro_seg_gota_alta": round(
    60 / ((((weight * 10 * 60) / 1000) / 12.5) * 20 / 60),
    0
) if weight > 0 else 0,

"dobutamina_micro_seg_gota_baja": round(
    60 / ((((weight * 5 * 60) / 1000) / 12.5) * 60 / 60),
    0
) if weight > 0 else 0,

"dobutamina_micro_seg_gota_alta": round(
    60 / ((((weight * 10 * 60) / 1000) / 12.5) * 60 / 60),
    0
) if weight > 0 else 0,

"atropina_mg": round(weight * 0.04, 2),
"atropina_ml": round((weight * 0.04) / 0.5, 2),

"adrenalina_mg_baja": round(weight * 0.01, 2),
"adrenalina_mg_alta": round(weight * 0.02, 2),
"adrenalina_ml_baja": round((weight * 0.01) / 1, 2),
"adrenalina_ml_alta": round((weight * 0.02) / 1, 2),

"diazepam_mg": round(weight * 0.5, 2),
"diazepam_ml": round((weight * 0.5) / 5, 2),

"dexametasona_mg": round(weight * 0.2, 2),
"dexametasona_ml": round((weight * 0.2) / 4, 2)

}
@app.get('/patients/{patient_id}/anesthesia', response_class=HTMLResponse)
def patient_anesthesia(
    request: Request,
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")

    weight = patient.weight or 0

    try:
        weight = float(weight)
    except:
        weight = 0

    doses = calculate_anesthesia_doses(weight)

    return templates.TemplateResponse(
        'anesthesia.html',
        {
            'request': request,
            'patient': patient,
            'doses': doses,
            'user': user
        }
    )
@app.post('/patients/{patient_id}/anesthesia/recalculate', response_class=HTMLResponse)
async def patient_anesthesia_recalculate(
    request: Request,
    patient_id: int,
    weight: float = Form(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")

    doses = calculate_anesthesia_doses(weight)

    return templates.TemplateResponse(
        'anesthesia.html',
        {
            'request': request,
            'patient': patient,
            'doses': doses,
            'user': user,
            'recalculated': True,
            'recalculated_weight': weight
        }
    )
@app.get('/patients/{patient_id}/anesthesia/print', response_class=HTMLResponse)
async def patient_anesthesia_print(
    request: Request,
    patient_id: int,
    event_id: int | None = None,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    last_anesthesia = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .filter(ClinicalEvent.event_type == 'Anestesia')
        .order_by(ClinicalEvent.event_date.desc())
        .first()
    )
    selected_anesthesia = last_anesthesia

    if event_id:
        selected_anesthesia = db.get(ClinicalEvent, event_id)

        if (
            not selected_anesthesia
            or selected_anesthesia.patient_id != patient.id
            or selected_anesthesia.event_type != 'Anestesia'
        ):
            raise HTTPException(status_code=404, detail="Protocolo anestésico no encontrado")
    weight = (
        selected_anesthesia.weight
        if selected_anesthesia and selected_anesthesia.weight
        else patient.weight or 0
    )

    try:
        weight = float(weight)
    except:
        weight = 0

    doses = calculate_anesthesia_doses(weight)
    return templates.TemplateResponse(
        'anesthesia_print.html',
        {
            'request': request,
            'patient': patient,
            'doses': doses,
            'last_anesthesia': selected_anesthesia,
            'user': user
        }
    )    
@app.post('/patients/{patient_id}/anesthesia')
async def save_patient_anesthesia(
    request: Request,
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")

    form = await request.form()

    weight = form.get('weight', '')

    if form.get('update_patient_weight') and weight:
        patient.weight = float(weight)
    procedure = form.get('procedure', '')
    asa = form.get('asa', '')
    estimated_duration = form.get('estimated_duration', '')
    drip_set = form.get('drip_set', '')
    fluid_rate = form.get('fluid_rate', '')
    
    premedication = []

    if form.get('premed_acepromazine'):
        premedication.append('Acepromacina')

    if form.get('premed_dexmedetomidine'):
        premedication.append('Dexmedetomidina')

    if form.get('premed_ketamine'):
        premedication.append('Ketamina IM')

    if form.get('premed_midazolam'):
        premedication.append('Midazolam')

    if form.get('premed_nalbuphine'):
        premedication.append('Nalbufina')

    if form.get('premed_butorphanol'):
        premedication.append('Butorfanol')


    induction = []

    if form.get('ind_propofol'):
        induction.append('Propofol')

    if form.get('ind_ketamine'):
        induction.append('Ketamina')

    if form.get('ind_midazolam'):
        induction.append('Midazolam') 
    maintenance = []

    if form.get('maint_isoflurane'):
        maintenance.append('Isofluorano')

    if form.get('maint_propofol'):
        maintenance.append('Propofol CRI')

    if form.get('maint_ketamine'):
        maintenance.append('Ketamina CRI')

    if form.get('maint_flk'):
        maintenance.append('FLK')


    intraop_analgesia = []

    if form.get('intra_fentanyl'):
        intraop_analgesia.append('Fentanilo')

    if form.get('intra_lidocaine'):
        intraop_analgesia.append('Lidocaína')

    if form.get('intra_ketamine'):
        intraop_analgesia.append('Ketamina')

    if form.get('intra_dobutamine'):
        intraop_analgesia.append('Dobutamina')
    flk_detail = []

    if form.get('maint_flk') or form.get('intra_fentanyl') or form.get('intra_lidocaine') or form.get('intra_ketamine'):
        flk_detail.append('FLK en 100 ml SSF')
        flk_detail.append('Fentanilo + Lidocaína + Ketamina')
        flk_detail.append('Velocidad según cálculo automático')
    postop_analgesia = []
    if form.get('meloxicam'):
        postop_analgesia.append('Meloxicam')
    if form.get('dipyrone'):
        postop_analgesia.append('Dipirona')
    if form.get('tramadol'):
        postop_analgesia.append('Tramadol')
    if form.get('gabapentin'):
        postop_analgesia.append('Gabapentina')
    if form.get('pregabalin'):
        postop_analgesia.append('Pregabalina')
    if form.get('nalbuphine'):
        postop_analgesia.append('Nalbufina')
    if form.get('butorphanol'):
        postop_analgesia.append('Butorfanol')
    if form.get('dexamethasone'):
        postop_analgesia.append('Dexametasona')

    antibiotics = []
    if form.get('cefazolin'):
        antibiotics.append('Cefazolina')
    if form.get('amoxiclav'):
        antibiotics.append('Amoxicilina clavulánico')
    if form.get('cephalexin'):
        antibiotics.append('Cefalexina')
    if form.get('enrofloxacin'):
        antibiotics.append('Enrofloxacina')
    if form.get('metronidazole'):
        antibiotics.append('Metronidazol')
    if form.get('clindamycin'):
        antibiotics.append('Clindamicina')

    description = f"""
Protocolo anestésico

Procedimiento: {procedure}
Peso: {weight} kg
ASA: {asa}
Duración estimada: {estimated_duration}
Premedicación:
{', '.join(premedication) if premedication else 'No especificada'}

Inducción:
{', '.join(induction) if induction else 'No especificada'}
Mantenimiento:
{', '.join(maintenance) if maintenance else 'No especificado'}

Analgesia intraoperatoria / soporte:
{', '.join(intraop_analgesia) if intraop_analgesia else 'No especificada'}
Detalle FLK:
{chr(10).join(flk_detail) if flk_detail else 'No especificado'}
Analgesia posoperatoria:
{', '.join(postop_analgesia) if postop_analgesia else 'No especificada'}

Observaciones analgesia:
{form.get('postop_notes', '')}

Antibioticoterapia:
{', '.join(antibiotics) if antibiotics else 'No especificada'}

Observaciones antibioticoterapia:
{form.get('antibiotic_notes', '')}

Fluidoterapia:
Equipo: {drip_set}
Velocidad: {fluid_rate} ml/kg/h
"""

    event = ClinicalEvent(
        patient_id=patient.id,
        event_type='Anestesia',
        description=description,
        event_date=argentina_now()
    )

    db.add(event)
    db.commit()

    return RedirectResponse(
            url=f"/patients/{patient.id}",
            status_code=303
    )
@app.get('/stats', response_class=HTMLResponse)
def stats_page(
    request: Request,
    period: str = 'month',
    month: str = '',
    compare_month: str = '',
    week_start: str = '',
    compare_week_start: str = '',
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    from collections import defaultdict
    import calendar
    import unicodedata

    today = argentina_now().date()

    def normalize(value):
        text_value = str(value or '').strip().lower()
        text_value = ''.join(
            char for char in unicodedata.normalize('NFD', text_value)
            if unicodedata.category(char) != 'Mn'
        )
        text_value = re.sub(r'[^a-z0-9]+', ' ', text_value)
        return re.sub(r'\s+', ' ', text_value).strip()

    def month_bounds(month_value):
        try:
            base = datetime.strptime(month_value, '%Y-%m').date()
        except Exception:
            base = today.replace(day=1)

        start = base.replace(day=1)
        last_day = calendar.monthrange(start.year, start.month)[1]
        return start, start.replace(day=last_day)

    def previous_month_value(month_value):
        start, _ = month_bounds(month_value)
        return (start - timedelta(days=1)).strftime('%Y-%m')

    if not month:
        month = today.strftime('%Y-%m')

    if not compare_month:
        compare_month = previous_month_value(month)

    selected_start, selected_end = month_bounds(month)
    compare_month_start, compare_month_end = month_bounds(compare_month)

    if period == 'today':
        start_date = today
        end_date = today
        compare_start = today - timedelta(days=1)
        compare_end = compare_start
        period_label = today.strftime('%d/%m/%Y')
        compare_label = compare_start.strftime('%d/%m/%Y')

    elif period == 'week':
        current_week_start = today - timedelta(days=today.weekday())

        try:
            requested_week_start = datetime.strptime(
                week_start,
                '%Y-%m-%d'
            ).date()

            start_date = (
                requested_week_start
                - timedelta(days=requested_week_start.weekday())
            )
        except Exception:
            start_date = current_week_start

        natural_end_date = start_date + timedelta(days=6)
        end_date = (
            today
            if start_date <= today <= natural_end_date
            else natural_end_date
        )

        try:
            requested_compare_start = datetime.strptime(
                compare_week_start,
                '%Y-%m-%d'
            ).date()

            compare_start = (
                requested_compare_start
                - timedelta(days=requested_compare_start.weekday())
            )
        except Exception:
            compare_start = start_date - timedelta(days=7)

        compared_days = (end_date - start_date).days
        compare_natural_end = compare_start + timedelta(days=6)
        compare_end = min(
            compare_start + timedelta(days=compared_days),
            compare_natural_end
        )

        week_start = start_date.strftime('%Y-%m-%d')
        compare_week_start = compare_start.strftime('%Y-%m-%d')

        period_label = (
            f'{start_date.strftime("%d/%m/%Y")} '
            f'al {end_date.strftime("%d/%m/%Y")}'
        )

        compare_label = (
            f'{compare_start.strftime("%d/%m/%Y")} '
            f'al {compare_end.strftime("%d/%m/%Y")}'
        )

    else:
        period = 'month'
        start_date = selected_start
        end_date = (
            today
            if selected_start <= today <= selected_end
            else selected_end
        )
        compare_start = compare_month_start
        compared_days = (end_date - start_date).days
        compare_end = min(
            compare_start + timedelta(days=compared_days),
            compare_month_end
        )

        if end_date < selected_end:
            period_label = (
                f'{start_date.strftime("%d/%m/%Y")} '
                f'al {end_date.strftime("%d/%m/%Y")}'
            )
            compare_label = (
                f'{compare_start.strftime("%d/%m/%Y")} '
                f'al {compare_end.strftime("%d/%m/%Y")}'
            )
        else:
            period_label = selected_start.strftime('%m/%Y')
            compare_label = compare_start.strftime('%m/%Y')

    category_meta = {
        'clinical': {
            'label': 'Servicios clínicos',
            'icon': '🩺'
        },
        'pharmacy': {
            'label': 'Farmacia',
            'icon': '💊'
        },
        'food': {
            'label': 'Alimentos balanceados',
            'icon': '🥣'
        },
        'petshop': {
            'label': 'Pet Shop',
            'icon': '🐾'
        },
        'other': {
            'label': 'Otros',
            'icon': '📦'
        }
    }

    def classify_product(product, rubro_value=''):
        rubro_text = normalize(
            rubro_value or getattr(product, 'rubro', '')
        )
        tipo_text = normalize(getattr(product, 'tipo', ''))
        name_text = normalize(getattr(product, 'name', ''))
        full_text = f'{rubro_text} {tipo_text} {name_text}'.strip()

        if (
            'servicio clinico' in rubro_text
            or 'servicios clinicos' in rubro_text
            or rubro_text in {
                'servicio',
                'servicios',
                'clinica',
                'clinico'
            }
        ):
            return 'clinical'

        if (
            'alimento balanceado' in rubro_text
            or 'alimentos balanceados' in rubro_text
            or 'balanceado' in rubro_text
            or 'alimento' in rubro_text
        ):
            return 'food'

        if (
            'petshop' in rubro_text
            or 'pet shop' in rubro_text
            or 'accesorio' in rubro_text
            or 'accesorios' in rubro_text
        ):
            return 'petshop'

        if (
            'farmacia' in rubro_text
            or 'medicamento' in rubro_text
            or 'medicina' in rubro_text
        ):
            return 'pharmacy'

        clinical_words = [
            'consulta',
            'control',
            'ecg',
            'radiografia',
            'ecografia',
            'cirugia',
            'internacion',
            'anestesia',
            'curacion',
            'laboratorio',
            'vacunacion',
            'servicio'
        ]

        food_words = [
            'balanceado',
            'alimento',
            'racion',
            'dieta'
        ]

        petshop_words = [
            'collar',
            'correa',
            'pretal',
            'juguete',
            'cama',
            'comedero',
            'bebedero',
            'transportadora',
            'arena',
            'piedras sanitarias',
            'shampoo',
            'cepillo',
            'ropa'
        ]

        pharmacy_words = [
            'farmacia',
            'medicamento',
            'comprimido',
            'inyectable',
            'ampolla',
            'antibiotico',
            'analgesico',
            'antiinflamatorio',
            'pipeta',
            'antiparasitario',
            'vacuna',
            'jarabe',
            'suspension'
        ]

        if any(word in full_text for word in clinical_words):
            return 'clinical'

        if any(word in full_text for word in food_words):
            return 'food'

        if any(word in full_text for word in petshop_words):
            return 'petshop'

        if any(word in full_text for word in pharmacy_words):
            return 'pharmacy'

        return 'other'

    def build_period_data(range_start, range_end):
        start_dt = datetime.combine(
            range_start,
            datetime.min.time()
        )

        end_dt = datetime.combine(
            range_end,
            datetime.max.time()
        )

        sales = (
            db.query(Sale)
            .filter(Sale.date >= start_dt)
            .filter(Sale.date <= end_dt)
            .filter(Sale.status != 'cancelled')
            .filter(Sale.status != 'quote')
            .all()
        )

        sale_ids = [sale.id for sale in sales]
        sales_by_id = {
            sale.id: sale
            for sale in sales
        }

        category_stats = {
            key: {
                'key': key,
                'label': meta['label'],
                'icon': meta['icon'],
                'sales': 0.0,
                'cost': 0.0,
                'profit': 0.0,
                'qty': 0.0,
                'percent': 0.0
            }
            for key, meta in category_meta.items()
        }

        product_stats = defaultdict(
            lambda: {
                'name': '',
                'category': 'other',
                'qty': 0.0,
                'sales': 0.0,
                'cost': 0.0,
                'profit': 0.0
            }
        )

        daily_stats = defaultdict(
            lambda: {
                'date': None,
                'sales': 0.0,
                'profit': 0.0,
                'count': 0,
                'ticket': 0.0
            }
        )

        sale_item_totals = defaultdict(float)
        items = []

        if sale_ids:
            items = (
                db.query(SaleItem)
                .filter(SaleItem.sale_id.in_(sale_ids))
                .all()
            )

            for item in items:
                sale_item_totals[item.sale_id] += (
                    item.subtotal or 0
                )

        for sale in sales:
            sale_day = (
                sale.date.date()
                if sale.date
                else range_start
            )

            daily_stats[sale_day]['date'] = sale_day
            daily_stats[sale_day]['sales'] += sale.total or 0
            daily_stats[sale_day]['count'] += 1

        for item in items:
            sale = sales_by_id.get(item.sale_id)
            product = db.get(Product, item.product_id)

            if not sale or not product:
                continue

            raw_subtotal = (
                item.subtotal
                or (
                    (item.quantity or 0)
                    * (item.unit_price or 0)
                )
            )

            item_total_for_sale = sale_item_totals.get(
                item.sale_id,
                0
            )

            allocation_factor = (
                (sale.total or 0) / item_total_for_sale
                if item_total_for_sale > 0
                else 1
            )

            allocated_sales = raw_subtotal * allocation_factor

            category = classify_product(product)
            qty = item.quantity or 0

            if category == 'clinical':
                cost = 0
                profit = allocated_sales
            else:
                cost = (
                    (product.cost_price or 0)
                    * qty
                )
                profit = allocated_sales - cost

            category_stats[category]['sales'] += allocated_sales
            category_stats[category]['cost'] += cost
            category_stats[category]['profit'] += profit
            category_stats[category]['qty'] += qty

            product_key = (
                category,
                product.id
            )

            product_stats[product_key]['name'] = (
                product.name or 'Sin nombre'
            )

            product_stats[product_key]['category'] = category
            product_stats[product_key]['qty'] += qty
            product_stats[product_key]['sales'] += allocated_sales
            product_stats[product_key]['cost'] += cost
            product_stats[product_key]['profit'] += profit

            sale_day = (
                sale.date.date()
                if sale.date
                else range_start
            )

            daily_stats[sale_day]['profit'] += profit

        sales_total = sum(
            sale.total or 0
            for sale in sales
        )

        sales_count = len(sales)

        profit_total = sum(
            data['profit']
            for data in category_stats.values()
        )

        cost_total = sum(
            data['cost']
            for data in category_stats.values()
        )

        ticket_average = (
            sales_total / sales_count
            if sales_count
            else 0
        )

        patients_seen = len({
            sale.patient_id
            for sale in sales
            if sale.patient_id
        })

        margin_percent = (
            profit_total / sales_total * 100
            if sales_total
            else 0
        )

        for data in category_stats.values():
            data['percent'] = (
                data['sales'] / sales_total * 100
                if sales_total
                else 0
            )

        daily_rows = []
        cursor = range_start

        while cursor <= range_end:
            data = daily_stats[cursor]
            data['date'] = cursor

            data['ticket'] = (
                data['sales'] / data['count']
                if data['count']
                else 0
            )

            daily_rows.append(dict(data))
            cursor += timedelta(days=1)

        days_with_sales = [
            row
            for row in daily_rows
            if row['count'] > 0
        ]

        best_day = (
            max(
                days_with_sales,
                key=lambda row: row['sales']
            )
            if days_with_sales
            else None
        )

        worst_day = (
            min(
                days_with_sales,
                key=lambda row: row['sales']
            )
            if days_with_sales
            else None
        )

        payments_by_method = {}

        if sale_ids:
            payments = (
                db.query(SalePayment)
                .filter(SalePayment.sale_id.in_(sale_ids))
                .all()
            )

            payment_stats = defaultdict(
                lambda: {
                    'count': 0,
                    'total': 0.0,
                    'percent': 0.0
                }
            )

            for payment in payments:
                method = payment.method or 'Sin método'
                payment_stats[method]['count'] += 1
                payment_stats[method]['total'] += (
                    payment.amount or 0
                )

            payment_total = sum(
                data['total']
                for data in payment_stats.values()
            )

            for data in payment_stats.values():
                data['percent'] = (
                    data['total'] / payment_total * 100
                    if payment_total
                    else 0
                )

            payments_by_method = dict(
                sorted(
                    payment_stats.items(),
                    key=lambda row: row[1]['total'],
                    reverse=True
                )
            )

        top_by_category = {}

        for category_key in category_meta:
            rows = [
                data
                for data in product_stats.values()
                if data['category'] == category_key
            ]

            top_by_category[category_key] = sorted(
                rows,
                key=lambda row: row['sales'],
                reverse=True
            )[:10]

        return {
            'sales': sales,
            'sales_total': sales_total,
            'sales_count': sales_count,
            'cost_total': cost_total,
            'profit_total': profit_total,
            'margin_percent': margin_percent,
            'ticket_average': ticket_average,
            'patients_seen': patients_seen,
            'category_stats': list(category_stats.values()),
            'category_stats_map': category_stats,
            'payments_by_method': payments_by_method,
            'daily_rows': daily_rows,
            'best_day': best_day,
            'worst_day': worst_day,
            'top_by_category': top_by_category
        }

    current = build_period_data(
        start_date,
        end_date
    )

    comparison = build_period_data(
        compare_start,
        compare_end
    )

    def diff(current_value, previous_value):
        amount = (
            (current_value or 0)
            - (previous_value or 0)
        )

        has_base = previous_value not in (None, 0)

        percent = (
            amount / previous_value * 100
            if has_base
            else None
        )

        return {
            'amount': amount,
            'percent': percent,
            'up': amount >= 0,
            'has_base': has_base
        }

    comparison_stats = {
        'sales_total': diff(
            current['sales_total'],
            comparison['sales_total']
        ),
        'profit_total': diff(
            current['profit_total'],
            comparison['profit_total']
        ),
        'sales_count': diff(
            current['sales_count'],
            comparison['sales_count']
        ),
        'ticket_average': diff(
            current['ticket_average'],
            comparison['ticket_average']
        ),
        'patients_seen': diff(
            current['patients_seen'],
            comparison['patients_seen']
        )
    }

    category_comparison = {}

    for key in category_meta:
        category_comparison[key] = {
            'sales': diff(
                current['category_stats_map'][key]['sales'],
                comparison['category_stats_map'][key]['sales']
            ),
            'profit': diff(
                current['category_stats_map'][key]['profit'],
                comparison['category_stats_map'][key]['profit']
            )
        }

    comparison_has_data = comparison['sales_count'] > 0

    earliest_sale = (
        db.query(Sale)
        .filter(Sale.status != 'cancelled')
        .filter(Sale.status != 'quote')
        .filter(Sale.date.isnot(None))
        .order_by(Sale.date.asc())
        .first()
    )

    earliest_sale_date = (
        earliest_sale.date
        if earliest_sale
        else None
    )

    first_data_date = (
        earliest_sale_date.date()
        if earliest_sale_date
        else today
    )

    first_week_start = (
        first_data_date
        - timedelta(days=first_data_date.weekday())
    )

    available_weeks = []
    current_week_start = (
        today - timedelta(days=today.weekday())
    )

    weeks_available = (
        (current_week_start - first_week_start).days // 7
    ) + 1

    for i in range(max(1, weeks_available)):
        option_start = (
            current_week_start
            - timedelta(weeks=i)
        )

        option_end = (
            option_start
            + timedelta(days=6)
        )

        available_weeks.append({
            'value': option_start.strftime('%Y-%m-%d'),
            'label': (
                f'{option_start.strftime("%d/%m/%Y")} '
                f'al {option_end.strftime("%d/%m/%Y")}'
            )
        })

    available_months = []
    ref = today.replace(day=1)
    first_month = first_data_date.replace(day=1)

    months_available = (
        (ref.year - first_month.year) * 12
        + ref.month
        - first_month.month
        + 1
    )

    for i in range(max(1, months_available)):
        year = ref.year
        month_number = ref.month - i

        while month_number <= 0:
            month_number += 12
            year -= 1

        available_months.append({
            'value': f'{year}-{month_number:02d}',
            'label': f'{month_number:02d}/{year}'
        })

    return templates.TemplateResponse(
        'stats.html',
        {
            'request': request,
            'period': period,
            'period_label': period_label,
            'compare_label': compare_label,
            'start_date': start_date,
            'end_date': end_date,
            'month': month,
            'compare_month': compare_month,
            'week_start': week_start,
            'compare_week_start': compare_week_start,
            'available_weeks': available_weeks,
            'available_months': available_months,
            'current': current,
            'comparison': comparison,
            'comparison_has_data': comparison_has_data,
            'comparison_stats': comparison_stats,
            'category_comparison': category_comparison,
            'category_meta': category_meta
        }
    )
# ===== VADEMÉCUM =====

@app.get('/vademecum', response_class=HTMLResponse)
def vademecum_page(
    request: Request,
    q: str = "",
    category: str = "",
    species: str = "",
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    active_ingredients = db.execute(
            text("""
                SELECT DISTINCT a.*
                FROM vademecum_active_ingredients a
                LEFT JOIN vademecum_brands b
                    ON b.active_ingredient_id = a.id
                WHERE a.active = TRUE
                AND (
                    :q = ''
                    OR a.name ILIKE :like
                    OR a.category ILIKE :like
                    OR a.species ILIKE :like
                    OR a.indications ILIKE :like
                    OR a.contraindications ILIKE :like
                    OR a.observations ILIKE :like
                    OR b.brand_name ILIKE :like
                    OR b.laboratory ILIKE :like
                )
                AND (
                    :category = ''
                    OR a.category ILIKE :category_like
                )
                AND (
                    :species = ''
                    OR a.species ILIKE :species_like
                )
                ORDER BY a.name
            """),
            {
                "q": q,
                "like": f"%{q}%",
                "category": category,
                "category_like": f"%{category}%",
                "species": species,
                "species_like": f"%{species}%"
            }
        ).mappings().all()
    
    brands = db.execute(
            text("""
                SELECT
                    b.*,
                    a.name AS active_name
                FROM vademecum_brands b
                JOIN vademecum_active_ingredients a
                    ON a.id = b.active_ingredient_id
                WHERE b.active = TRUE
                ORDER BY a.name, b.brand_name
            """)
        ).mappings().all()
    
    return templates.TemplateResponse(
            "vademecum.html",
            {
                "request": request,
                "active_ingredients": active_ingredients,
                "brands": brands,
                "q": q,
                "category": category,
                "species": species
            }
        )
@app.get('/vademecum/api/search')
def vademecum_api_search(
    q: str = '',
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    like = f"%{q.strip()}%"

    rows = db.execute(
        text("""
            SELECT
                a.id,
                a.name,
                a.category,
                a.species,
                a.dog_dose,
                a.cat_dose,
                a.route,
                a.frequency,
                a.indications,
                a.contraindications,
                a.warnings,
                a.observations,
                b.brand_name,
                b.laboratory,
                b.presentation,
                b.concentration
            FROM vademecum_active_ingredients a
            LEFT JOIN vademecum_brands b
                ON b.active_ingredient_id = a.id
                AND b.active = TRUE
            WHERE a.active = TRUE
            AND (
                :q = ''
                OR a.name ILIKE :like
                OR a.category ILIKE :like
                OR a.species ILIKE :like
                OR a.indications ILIKE :like
                OR a.contraindications ILIKE :like
                OR a.observations ILIKE :like
                OR b.brand_name ILIKE :like
                OR b.laboratory ILIKE :like
            )
            ORDER BY a.name, b.brand_name
            LIMIT 80
        """),
        {
            "q": q.strip(),
            "like": like
        }
    ).mappings().all()

    results = {}

    for row in rows:
        active_id = row["id"]

        if active_id not in results:
            results[active_id] = {
                "id": active_id,
                "name": row["name"] or "",
                "category": row["category"] or "",
                "species": row["species"] or "",
                "dog_dose": row["dog_dose"] or "",
                "cat_dose": row["cat_dose"] or "",
                "route": row["route"] or "",
                "frequency": row["frequency"] or "",
                "indications": row["indications"] or "",
                "contraindications": row["contraindications"] or "",
                "warnings": row["warnings"] or "",
                "observations": row["observations"] or "",
                "brands": []
            }

        if row["brand_name"]:
            results[active_id]["brands"].append({
                "brand_name": row["brand_name"] or "",
                "laboratory": row["laboratory"] or "",
                "presentation": row["presentation"] or "",
                "concentration": row["concentration"] or ""
            })

    return JSONResponse(list(results.values())[:20])
def extract_numbers(text_value):
    text_value = str(text_value or '').replace(',', '.')
    return [float(x) for x in re.findall(r'\d+(?:\.\d+)?', text_value)]


def extract_mgkg_range(dose_text):
    text_value = str(dose_text or '').lower().replace(',', '.')

    if 'mg/kg' not in text_value and 'mg / kg' not in text_value:
        return None, None

    numbers = extract_numbers(text_value)

    if not numbers:
        return None, None

    if len(numbers) >= 2 and ('-' in text_value or ' a ' in text_value):
        return numbers[0], numbers[1]

    return numbers[0], numbers[0]


def extract_concentration(concentration_text, presentation_text=''):
    text_value = str(concentration_text or '').lower().replace(',', '.')
    presentation = str(presentation_text or '').lower()

    match = re.search(r'(\d+(?:\.\d+)?)\s*mg\s*/\s*ml', text_value)
    if match:
        return float(match.group(1)), 'mg/ml'

    match = re.search(r'(\d+(?:\.\d+)?)\s*mg', text_value)
    if match:
        value = float(match.group(1))

        if any(word in presentation for word in ['comp', 'comprim', 'tablet', 'tab']):
            return value, 'mg/comprimido'

        if any(word in text_value for word in ['comp', 'comprim', 'tablet', 'tab']):
            return value, 'mg/comprimido'

        return value, 'mg/unidad'

    return None, ''


@app.get('/hospitalizations/{hospitalization_id}/vademecum-dose')
def hospitalization_vademecum_dose(
    hospitalization_id: int,
    active_id: int,
    brand_index: int = 0,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)

    if not hospitalization:
        raise HTTPException(status_code=404, detail='Internación no encontrada')

    patient = hospitalization.patient
    weight = hospitalization.initial_weight or patient.weight or 0

    try:
        weight = float(weight)
    except Exception:
        weight = 0

    active = db.execute(
        text("""
            SELECT *
            FROM vademecum_active_ingredients
            WHERE id = :active_id
            AND active = TRUE
        """),
        {"active_id": active_id}
    ).mappings().first()

    if not active:
        raise HTTPException(status_code=404, detail='Principio activo no encontrado')

    brands = db.execute(
        text("""
            SELECT *
            FROM vademecum_brands
            WHERE active_ingredient_id = :active_id
            AND active = TRUE
            ORDER BY brand_name
        """),
        {"active_id": active_id}
    ).mappings().all()

    species_text = str(patient.species or '').lower()

    if 'fel' in species_text or 'gato' in species_text:
        dose_text = active.get('cat_dose') or active.get('dog_dose') or ''
    else:
        dose_text = active.get('dog_dose') or active.get('cat_dose') or ''

    mgkg_min, mgkg_max = extract_mgkg_range(dose_text)

    selected_brand = None
    if brands:
        if brand_index < 0:
            brand_index = 0
        if brand_index >= len(brands):
            brand_index = 0
        selected_brand = brands[brand_index]

    concentration_text = selected_brand.get('concentration') if selected_brand else ''
    presentation_text = selected_brand.get('presentation') if selected_brand else ''
    concentration_value, concentration_unit = extract_concentration(concentration_text, presentation_text)

    mg_total_min = round(weight * mgkg_min, 2) if mgkg_min else None
    mg_total_max = round(weight * mgkg_max, 2) if mgkg_max else None

    amount_text = ''

    if mg_total_min and concentration_value:
        amount_min = round(mg_total_min / concentration_value, 2)
        amount_max = round(mg_total_max / concentration_value, 2) if mg_total_max else amount_min

        if concentration_unit == 'mg/ml':
            if amount_min == amount_max:
                amount_text = f'{amount_min:g} ml'
            else:
                amount_text = f'{amount_min:g} a {amount_max:g} ml'

        elif concentration_unit == 'mg/comprimido':
            if amount_min == amount_max:
                amount_text = f'{amount_min:g} comprimido(s)'
            else:
                amount_text = f'{amount_min:g} a {amount_max:g} comprimido(s)'

        else:
            if amount_min == amount_max:
                amount_text = f'{amount_min:g} unidad(es)'
            else:
                amount_text = f'{amount_min:g} a {amount_max:g} unidad(es)'

    dose_result = ''

    if mg_total_min:
        if mg_total_min == mg_total_max:
            dose_result = f'{mg_total_min:g} mg'
        else:
            dose_result = f'{mg_total_min:g} a {mg_total_max:g} mg'

        if amount_text:
            dose_result += f' = {amount_text}'

    return JSONResponse({
        'ok': True,
        'patient_weight': weight,
        'active_id': active_id,
        'drug_name': active.get('name') or '',
        'brand_name': selected_brand.get('brand_name') if selected_brand else '',
        'dose_text': dose_text,
        'mgkg_min': mgkg_min,
        'mgkg_max': mgkg_max,
        'mg_total_min': mg_total_min,
        'mg_total_max': mg_total_max,
        'concentration': concentration_text or '',
        'presentation': presentation_text or '',
        'amount_text': amount_text,
        'dose_result': dose_result,
        'route': active.get('route') or '',
        'frequency': active.get('frequency') or '',
        'observations': active.get('observations') or '',
        'warnings': active.get('warnings') or '',
        'brands': [
            {
                'brand_name': b.get('brand_name') or '',
                'laboratory': b.get('laboratory') or '',
                'presentation': b.get('presentation') or '',
                'concentration': b.get('concentration') or ''
            }
            for b in brands
        ]
    })
@app.get('/vademecum/{active_id}', response_class=HTMLResponse)
def vademecum_detail(
    active_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    active = db.execute(
        text("""
            SELECT *
            FROM vademecum_active_ingredients
            WHERE id = :active_id AND active = TRUE
        """),
        {"active_id": active_id}
    ).mappings().first()

    if not active:
        raise HTTPException(status_code=404, detail="Principio activo no encontrado")

    brands = db.execute(
        text("""
            SELECT *
            FROM vademecum_brands
            WHERE active_ingredient_id = :active_id
            AND active = TRUE
            ORDER BY brand_name
        """),
        {"active_id": active_id}
    ).mappings().all()

    return templates.TemplateResponse(
        "vademecum_detail.html",
        {
            "request": request,
            "active": active,
            "brands": brands
        }
    )
@app.get("/vademecum/{active_id}/edit", response_class=HTMLResponse)
def edit_vademecum(
    active_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    active = db.execute(
        text("""
            SELECT *
            FROM vademecum_active_ingredients
            WHERE id = :id
        """),
        {"id": active_id}
    ).mappings().first()

    if not active:
        raise HTTPException(status_code=404, detail="Principio activo no encontrado")

    return templates.TemplateResponse(
        "vademecum_edit.html",
        {
            "request": request,
            "active": active
        }
    )


@app.post("/vademecum/{active_id}/edit")
def save_vademecum(
    active_id: int,
    name: str = Form(...),
    category: str = Form(""),
    species: str = Form(""),
    dog_dose: str = Form(""),
    cat_dose: str = Form(""),
    route: str = Form(""),
    frequency: str = Form(""),
    indications: str = Form(""),
    contraindications: str = Form(""),
    interactions: str = Form(""),
    warnings: str = Form(""),
    observations: str = Form(""),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    db.execute(
        text("""
            UPDATE vademecum_active_ingredients
            SET
                name=:name,
                category=:category,
                species=:species,
                dog_dose=:dog_dose,
                cat_dose=:cat_dose,
                route=:route,
                frequency=:frequency,
                indications=:indications,
                contraindications=:contraindications,
                interactions=:interactions,
                warnings=:warnings,
                observations=:observations
            WHERE id=:id
        """),
        {
            "id": active_id,
            "name": name,
            "category": category,
            "species": species,
            "dog_dose": dog_dose,
            "cat_dose": cat_dose,
            "route": route,
            "frequency": frequency,
            "indications": indications,
            "contraindications": contraindications,
            "interactions": interactions,
            "warnings": warnings,
            "observations": observations
        }
    )

    db.commit()

    return RedirectResponse(
        f"/vademecum/{active_id}",
        status_code=303
    )
@app.post("/vademecum/{active_id}/delete")
def delete_vademecum(
    active_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    db.execute(
        text("""
            UPDATE vademecum_active_ingredients
            SET active = FALSE
            WHERE id = :id
        """),
        {"id": active_id}
    )

    db.execute(
        text("""
            UPDATE vademecum_brands
            SET active = FALSE
            WHERE active_ingredient_id = :id
        """),
        {"id": active_id}
    )

    db.commit()

    return RedirectResponse(
        "/vademecum",
        status_code=303
    )
@app.post('/vademecum')
def vademecum_create(
    commercial_name: str = Form(''),
    active_ingredient: str = Form(''),
    category: str = Form(''),
    species: str = Form(''),
    dose: str = Form(''),
    route: str = Form(''),
    frequency: str = Form(''),
    indications: str = Form(''),
    contraindications: str = Form(''),
    observations: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    db.execute(
        text("""
            INSERT INTO vademecum_drugs (
                commercial_name,
                active_ingredient,
                category,
                species,
                dose,
                route,
                frequency,
                indications,
                contraindications,
                observations,
                active
            )
            VALUES (
                :commercial_name,
                :active_ingredient,
                :category,
                :species,
                :dose,
                :route,
                :frequency,
                :indications,
                :contraindications,
                :observations,
                TRUE
            )
        """),
        {
            "commercial_name": commercial_name,
            "active_ingredient": active_ingredient,
            "category": category,
            "species": species,
            "dose": dose,
            "route": route,
            "frequency": frequency,
            "indications": indications,
            "contraindications": contraindications,
            "observations": observations
        }
    )

    db.commit()

    return RedirectResponse('/vademecum', status_code=303)


@app.post('/vademecum/{drug_id}/delete')
def vademecum_delete(
    drug_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    db.execute(
        text("""
            UPDATE vademecum_drugs
            SET active = FALSE
            WHERE id = :drug_id
        """),
        {"drug_id": drug_id}
    )

    db.commit()

    return RedirectResponse('/vademecum', status_code=303)
@app.get('/vademecum/import/template')
def vademecum_import_template(
    user: User = Depends(require_user)
):
    csv_content = (
        "Principio activo,Nombre comercial,Laboratorio,Presentación,Concentración,Especie,Categoría,Vía,"
        "Dosis perro,Dosis gato,Frecuencia,Indicaciones,Contraindicaciones,Interacciones,Advertencias,Observaciones\n"
        "Meloxicam,Meloxivet,,Inyectable,4 mg/ml,Canino y felino,AINE,SC,"
        "0.2 mg/kg inicial; luego 0.1 mg/kg cada 24 h,0.1 mg/kg inicial; luego 0.05 mg/kg cada 24 h,"
        "Cada 24 h,Dolor e inflamación,Evitar en insuficiencia renal/deshidratación/úlcera GI,"
        "No combinar con corticoides u otros AINE,Usar con monitoreo renal,Ejemplo de carga para Aromos Cloud\n"
    )

    return Response(
        content=csv_content.encode("utf-8-sig"),
        media_type="text/csv; charset=utf-8",
        headers={
            "Content-Disposition": "attachment; filename=plantilla_vademecum_ampliada.csv"
        }
    )


@app.post("/vademecum/update-senasa")
def update_senasa(
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    try:
        summary = update_from_senasa(db, limit=None)
        print(summary)
        return RedirectResponse(
            url=(
                "/vademecum?"
                f"imported=1"
                f"&new_active={summary.get('new_active', 0)}"
                f"&updated_active={summary.get('updated_active', 0)}"
                f"&new_brands={summary.get('new_brands', 0)}"
                f"&updated_brands={summary.get('updated_brands', 0)}"
                f"&skipped={summary.get('skipped', 0)}"
                f"&errors={summary.get('details_error', 0) + len(summary.get('errors', []))}"
            ),
            status_code=303
        )

    except Exception:
        return RedirectResponse(
            url=f"/vademecum?imported=1&errors=1",
            status_code=303
        )


@app.post('/vademecum/import')
async def vademecum_import(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    content = await file.read()
    filename = (file.filename or "").lower()

    def clean(value):
        if value is None:
            return ""
        return str(value).strip()

    def normalize(value):
        return (
            clean(value)
            .lower()
            .replace("á", "a")
            .replace("é", "e")
            .replace("í", "i")
            .replace("ó", "o")
            .replace("ú", "u")
            .replace("ñ", "n")
            .replace(" ", "")
            .replace("_", "")
            .replace("-", "")
        )

    def pick(row, *names):
        normalized = {normalize(k): clean(v) for k, v in row.items()}
        for name in names:
            key = normalize(name)
            if key in normalized:
                return normalized[key]
        return ""

    rows = []

    if filename.endswith(".xlsx"):
        wb = load_workbook(BytesIO(content), data_only=True)
        ws = wb.active
        headers = [clean(c.value) for c in ws[1]]

        for row_values in ws.iter_rows(min_row=2, values_only=True):
            rows.append(dict(zip(headers, row_values)))

    else:
        text_content = content.decode("utf-8-sig", errors="replace")
        sample = text_content[:2000]
        delimiter = ";" if sample.count(";") >= sample.count(",") else ","
        reader = csv.DictReader(StringIO(text_content), delimiter=delimiter)
        rows = list(reader)

    new_active = 0
    updated_active = 0
    new_brands = 0
    updated_brands = 0
    skipped = 0
    errors = 0

    for row in rows:
        try:
            active_name = pick(
                row,
                "Principio activo",
                "Droga",
                "Activo",
                "Monodroga"
            )

            brand_name = pick(
                row,
                "Nombre comercial",
                "Comercial",
                "Marca",
                "Producto"
            )

            if not active_name and not brand_name:
                skipped += 1
                continue

            if not active_name:
                active_name = brand_name

            laboratory = pick(row, "Laboratorio", "Elaborador", "Manufacturer")
            presentation = pick(row, "Presentación", "Presentacion", "Forma farmacéutica", "Forma farmaceutica")
            concentration = pick(row, "Concentración", "Concentracion", "Concentration")
            species = pick(row, "Especie", "Species")
            category = pick(row, "Categoría", "Categoria", "Rubro")
            route = pick(row, "Vía", "Via", "Ruta", "Route")
            frequency = pick(row, "Frecuencia", "Frequency")

            dog_dose = pick(
                row,
                "Dosis perro",
                "Dosis canino",
                "Dosis perros",
                "Canino",
                "Dosis"
            )

            cat_dose = pick(
                row,
                "Dosis gato",
                "Dosis felino",
                "Dosis gatos",
                "Felino"
            )

            indications = pick(row, "Indicaciones", "Usos", "Uso")
            contraindications = pick(row, "Contraindicaciones", "Contraindicacion")
            interactions = pick(row, "Interacciones", "Interacciones importantes")
            warnings = pick(row, "Advertencias", "Precauciones", "Precaución", "Precaucion")
            observations = pick(row, "Observaciones", "Notas", "Comentarios")

            active = db.execute(
                text("""
                    SELECT *
                    FROM vademecum_active_ingredients
                    WHERE lower(name) = lower(:name)
                    LIMIT 1
                """),
                {"name": active_name}
            ).mappings().first()

            if active:
                db.execute(
                    text("""
                        UPDATE vademecum_active_ingredients
                        SET
                            category = :category,
                            species = :species,
                            dog_dose = :dog_dose,
                            cat_dose = :cat_dose,
                            route = :route,
                            frequency = :frequency,
                            indications = :indications,
                            contraindications = :contraindications,
                            interactions = :interactions,
                            warnings = :warnings,
                            observations = :observations,
                            active = TRUE
                        WHERE id = :id
                    """),
                    {
                        "id": active["id"],
                        "category": category or active["category"] or "",
                        "species": species or active["species"] or "",
                        "dog_dose": dog_dose or active["dog_dose"] or "",
                        "cat_dose": cat_dose or active["cat_dose"] or "",
                        "route": route or active["route"] or "",
                        "frequency": frequency or active["frequency"] or "",
                        "indications": indications or active["indications"] or "",
                        "contraindications": contraindications or active["contraindications"] or "",
                        "interactions": interactions or active["interactions"] or "",
                        "warnings": warnings or active["warnings"] or "",
                        "observations": observations or active["observations"] or "",
                    }
                )
                active_id = active["id"]
                updated_active += 1

            else:
                result = db.execute(
                    text("""
                        INSERT INTO vademecum_active_ingredients (
                            name, category, species, dog_dose, cat_dose,
                            route, frequency, indications, contraindications,
                            interactions, warnings, observations, active
                        )
                        VALUES (
                            :name, :category, :species, :dog_dose, :cat_dose,
                            :route, :frequency, :indications, :contraindications,
                            :interactions, :warnings, :observations, TRUE
                        )
                        RETURNING id
                    """),
                    {
                        "name": active_name,
                        "category": category,
                        "species": species,
                        "dog_dose": dog_dose,
                        "cat_dose": cat_dose,
                        "route": route,
                        "frequency": frequency,
                        "indications": indications,
                        "contraindications": contraindications,
                        "interactions": interactions,
                        "warnings": warnings,
                        "observations": observations,
                    }
                )

                active_id = result.scalar()
                new_active += 1

            if brand_name:
                brand = db.execute(
                    text("""
                        SELECT *
                        FROM vademecum_brands
                        WHERE active_ingredient_id = :active_id
                        AND lower(brand_name) = lower(:brand_name)
                        LIMIT 1
                    """),
                    {
                        "active_id": active_id,
                        "brand_name": brand_name
                    }
                ).mappings().first()

                if brand:
                    db.execute(
                        text("""
                            UPDATE vademecum_brands
                            SET
                                laboratory = :laboratory,
                                presentation = :presentation,
                                concentration = :concentration,
                                active = TRUE
                            WHERE id = :id
                        """),
                        {
                            "id": brand["id"],
                            "laboratory": laboratory or brand["laboratory"] or "",
                            "presentation": presentation or brand["presentation"] or "",
                            "concentration": concentration or brand["concentration"] or "",
                        }
                    )
                    updated_brands += 1

                else:
                    db.execute(
                        text("""
                            INSERT INTO vademecum_brands (
                                active_ingredient_id,
                                brand_name,
                                laboratory,
                                presentation,
                                concentration,
                                active
                            )
                            VALUES (
                                :active_id,
                                :brand_name,
                                :laboratory,
                                :presentation,
                                :concentration,
                                TRUE
                            )
                        """),
                        {
                            "active_id": active_id,
                            "brand_name": brand_name,
                            "laboratory": laboratory,
                            "presentation": presentation,
                            "concentration": concentration,
                        }
                    )
                    new_brands += 1

            dose_text = ""
            if dog_dose and cat_dose:
                dose_text = f"Perro: {dog_dose} | Gato: {cat_dose}"
            elif dog_dose:
                dose_text = dog_dose
            elif cat_dose:
                dose_text = cat_dose

            old_drug = db.execute(
                text("""
                    SELECT id
                    FROM vademecum_drugs
                    WHERE lower(commercial_name) = lower(:commercial_name)
                    AND lower(active_ingredient) = lower(:active_ingredient)
                    LIMIT 1
                """),
                {
                    "commercial_name": brand_name or active_name,
                    "active_ingredient": active_name
                }
            ).mappings().first()

            if old_drug:
                db.execute(
                    text("""
                        UPDATE vademecum_drugs
                        SET
                            commercial_name = :commercial_name,
                            active_ingredient = :active_ingredient,
                            category = :category,
                            species = :species,
                            dose = :dose,
                            route = :route,
                            frequency = :frequency,
                            indications = :indications,
                            contraindications = :contraindications,
                            observations = :observations,
                            active = TRUE
                        WHERE id = :id
                    """),
                    {
                        "id": old_drug["id"],
                        "commercial_name": brand_name or active_name,
                        "active_ingredient": active_name,
                        "category": category,
                        "species": species,
                        "dose": dose_text,
                        "route": route,
                        "frequency": frequency,
                        "indications": indications,
                        "contraindications": contraindications,
                        "observations": observations,
                    }
                )
            else:
                db.execute(
                    text("""
                        INSERT INTO vademecum_drugs (
                            commercial_name,
                            active_ingredient,
                            category,
                            species,
                            dose,
                            route,
                            frequency,
                            indications,
                            contraindications,
                            observations,
                            active
                        )
                        VALUES (
                            :commercial_name,
                            :active_ingredient,
                            :category,
                            :species,
                            :dose,
                            :route,
                            :frequency,
                            :indications,
                            :contraindications,
                            :observations,
                            TRUE
                        )
                    """),
                    {
                        "commercial_name": brand_name or active_name,
                        "active_ingredient": active_name,
                        "category": category,
                        "species": species,
                        "dose": dose_text,
                        "route": route,
                        "frequency": frequency,
                        "indications": indications,
                        "contraindications": contraindications,
                        "observations": observations,
                    }
                )

        except Exception as e:
            print("ERROR IMPORTANDO VADEMECUM:", str(e))
            errors += 1

    db.commit()

    return RedirectResponse(
        url=(
            "/vademecum?"
            f"imported=1"
            f"&new_active={new_active}"
            f"&updated_active={updated_active}"
            f"&new_brands={new_brands}"
            f"&updated_brands={updated_brands}"
            f"&skipped={skipped}"
            f"&errors={errors}"
        ),
        status_code=303
    )
@app.post('/vademecum/clear')
def vademecum_clear(
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    db.execute(text("DELETE FROM vademecum_brands"))
    db.execute(text("DELETE FROM vademecum_active_ingredients"))
    db.execute(text("DELETE FROM vademecum_drugs"))

    db.commit()

    return RedirectResponse(
        "/vademecum?cleared=1",
        status_code=303
    )
@app.post('/vademecum/active-ingredient')
def vademecum_active_ingredient_create(
    name: str = Form(''),
    category: str = Form(''),
    species: str = Form(''),
    dog_dose: str = Form(''),
    cat_dose: str = Form(''),
    route: str = Form(''),
    frequency: str = Form(''),
    indications: str = Form(''),
    contraindications: str = Form(''),
    interactions: str = Form(''),
    warnings: str = Form(''),
    observations: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    db.execute(
        text("""
            INSERT INTO vademecum_active_ingredients (
                name, category, species, dog_dose, cat_dose, route,
                frequency, indications, contraindications, interactions,
                warnings, observations, active
            )
            VALUES (
                :name, :category, :species, :dog_dose, :cat_dose, :route,
                :frequency, :indications, :contraindications, :interactions,
                :warnings, :observations, TRUE
            )
        """),
        {
            "name": name,
            "category": category,
            "species": species,
            "dog_dose": dog_dose,
            "cat_dose": cat_dose,
            "route": route,
            "frequency": frequency,
            "indications": indications,
            "contraindications": contraindications,
            "interactions": interactions,
            "warnings": warnings,
            "observations": observations
        }
    )

    db.commit()

    return RedirectResponse('/vademecum', status_code=303)


@app.post('/vademecum/brand')
def vademecum_brand_create(
    active_ingredient_id: str = Form(''),
    brand_name: str = Form(''),
    laboratory: str = Form(''),
    presentation: str = Form(''),
    concentration: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    if not active_ingredient_id:
        return RedirectResponse('/vademecum', status_code=303)

    db.execute(
        text("""
            INSERT INTO vademecum_brands (
                active_ingredient_id, brand_name, laboratory,
                presentation, concentration, active
            )
            VALUES (
                :active_ingredient_id, :brand_name, :laboratory,
                :presentation, :concentration, TRUE
            )
        """),
        {
            "active_ingredient_id": int(active_ingredient_id),
            "brand_name": brand_name,
            "laboratory": laboratory,
            "presentation": presentation,
            "concentration": concentration
        }
    )

    db.commit()

    return RedirectResponse('/vademecum', status_code=303)
@app.get("/vademecum/brand/{brand_id}/edit", response_class=HTMLResponse)
def edit_vademecum_brand(
    brand_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    brand = db.execute(
        text("""
            SELECT *
            FROM vademecum_brands
            WHERE id = :brand_id
            AND active = TRUE
        """),
        {"brand_id": brand_id}
    ).mappings().first()

    if not brand:
        raise HTTPException(status_code=404, detail="Marca comercial no encontrada")

    return templates.TemplateResponse(
        "vademecum_brand_edit.html",
        {
            "request": request,
            "brand": brand
        }
    )


@app.post("/vademecum/brand/{brand_id}/edit")
def save_vademecum_brand(
    brand_id: int,
    brand_name: str = Form(""),
    laboratory: str = Form(""),
    presentation: str = Form(""),
    concentration: str = Form(""),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    brand = db.execute(
        text("""
            SELECT active_ingredient_id
            FROM vademecum_brands
            WHERE id = :brand_id
        """),
        {"brand_id": brand_id}
    ).mappings().first()

    if not brand:
        raise HTTPException(status_code=404, detail="Marca comercial no encontrada")

    db.execute(
        text("""
            UPDATE vademecum_brands
            SET
                brand_name = :brand_name,
                laboratory = :laboratory,
                presentation = :presentation,
                concentration = :concentration
            WHERE id = :brand_id
        """),
        {
            "brand_id": brand_id,
            "brand_name": brand_name,
            "laboratory": laboratory,
            "presentation": presentation,
            "concentration": concentration
        }
    )

    db.commit()

    return RedirectResponse(
        f"/vademecum/{brand.active_ingredient_id}",
        status_code=303
    )


@app.post("/vademecum/brand/{brand_id}/delete")
def delete_vademecum_brand(
    brand_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    brand = db.execute(
        text("""
            SELECT active_ingredient_id
            FROM vademecum_brands
            WHERE id = :brand_id
        """),
        {"brand_id": brand_id}
    ).mappings().first()

    if not brand:
        raise HTTPException(status_code=404, detail="Marca comercial no encontrada")

    db.execute(
        text("""
            UPDATE vademecum_brands
            SET active = FALSE
            WHERE id = :brand_id
        """),
        {"brand_id": brand_id}
    )

    db.commit()

    return RedirectResponse(
        f"/vademecum/{brand.active_ingredient_id}",
        status_code=303
    )
@app.get("/vademecum/{active_id}/print", response_class=HTMLResponse)
def vademecum_print(
    active_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    active = db.execute(
        text("""
            SELECT *
            FROM vademecum_active_ingredients
            WHERE id = :active_id
            AND active = TRUE
        """),
        {"active_id": active_id}
    ).mappings().first()

    if not active:
        raise HTTPException(status_code=404, detail="Principio activo no encontrado")

    brands = db.execute(
        text("""
            SELECT *
            FROM vademecum_brands
            WHERE active_ingredient_id = :active_id
            AND active = TRUE
            ORDER BY brand_name
        """),
        {"active_id": active_id}
    ).mappings().all()

    return templates.TemplateResponse(
        "vademecum_print.html",
        {
            "request": request,
            "active": active,
            "brands": brands
        }
    )
# ===== INTERNACIÓN =====

@app.get('/hospitalizations', response_class=HTMLResponse)
def hospitalizations_page(
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    active_hospitalizations = (
        db.query(Hospitalization)
        .filter(Hospitalization.status == 'Internado')
        .order_by(Hospitalization.admission_date.desc())
        .all()
    )

    closed_hospitalizations = (
        db.query(Hospitalization)
        .filter(Hospitalization.status != 'Internado')
        .order_by(Hospitalization.admission_date.desc())
        .limit(20)
        .all()
    )

    return templates.TemplateResponse(
        'hospitalizations.html',
        {
            'request': request,
            'active_hospitalizations': active_hospitalizations,
            'closed_hospitalizations': closed_hospitalizations,
            'today': argentina_now().date()
        }
    )


@app.get('/patients/{patient_id}/hospitalize', response_class=HTMLResponse)
def hospitalization_form(
    request: Request,
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404, detail='Paciente no encontrado')

    active_hospitalization = (
        db.query(Hospitalization)
        .filter(
            Hospitalization.patient_id == patient.id,
            Hospitalization.status == 'Internado'
        )
        .first()
    )

    if active_hospitalization:
        return RedirectResponse(
            f'/hospitalizations/{active_hospitalization.id}',
            status_code=303
        )

    return templates.TemplateResponse(
        'hospitalization_form.html',
        {
            'request': request,
            'patient': patient,
            'today': argentina_now().date()
        }
    )


@app.post('/patients/{patient_id}/hospitalize')
def hospitalization_create(
    patient_id: int,
    cage: str = Form(''),
    responsible_vet: str = Form(''),
    reason: str = Form(''),
    diagnosis: str = Form(''),
    treatment_plan: str = Form(''),
    notes: str = Form(''),
    initial_weight: str = Form(''),
    initial_temperature: str = Form(''),
    initial_heart_rate: str = Form(''),
    initial_respiratory_rate: str = Form(''),
    initial_mucous_membranes: str = Form(''),
    initial_crt: str = Form(''),
    initial_hydration: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404, detail='Paciente no encontrado')

    active_hospitalization = (
        db.query(Hospitalization)
        .filter(
            Hospitalization.patient_id == patient.id,
            Hospitalization.status == 'Internado'
        )
        .first()
    )

    if active_hospitalization:
        return RedirectResponse(
            f'/hospitalizations/{active_hospitalization.id}',
            status_code=303
        )

    def to_float(value):
        try:
            return float(str(value).replace(',', '.')) if value and str(value).strip() else None
        except ValueError:
            return None

    def to_int(value):
        try:
            return int(float(str(value).replace(',', '.'))) if value and str(value).strip() else None
        except ValueError:
            return None

    event = ClinicalEvent(
        patient_id=patient.id,
        event_type='Internación',
        title='Ingreso a internación',
        description=reason or '',
        diagnosis=diagnosis or '',
        treatment=treatment_plan or '',
        weight=to_float(initial_weight),
        temperature=to_float(initial_temperature),
        heart_rate=to_int(initial_heart_rate),
        respiratory_rate=to_int(initial_respiratory_rate),
        mucous_membranes=initial_mucous_membranes or '',
        crt=initial_crt or '',
        hydration=initial_hydration or '',
        created_by=user.username,
        event_date=argentina_now()
    )

    db.add(event)
    db.flush()

    # Resultados de laboratorio revisados en la vista previa.
    lab_results_raw = get('lab_results_json')
    if lab_results_raw and str(lab_results_raw).strip():
        try:
            lab_payload = json.loads(lab_results_raw)
        except json.JSONDecodeError:
            lab_payload = {}

        sample_date = parse_date(lab_payload.get('sample_date', '')) or event_datetime.date()
        source_files = ', '.join(lab_payload.get('filenames', []) or [])[:255]

        for row in lab_payload.get('results', []) or []:
            analyte = str(row.get('analyte') or '').strip()
            if not analyte:
                continue
            db.add(LaboratoryResult(
                event_id=event.id,
                patient_id=patient.id,
                sample_date=sample_date,
                panel=str(row.get('panel') or 'Otro')[:120],
                analyte=analyte[:180],
                normalized_analyte=_lab_normalize_analyte(analyte)[:180],
                value_text=str(row.get('value_text') or '')[:120],
                value_numeric=_lab_float(row.get('value_numeric')),
                unit=str(row.get('unit') or '')[:80],
                reference_low=_lab_float(row.get('reference_low')),
                reference_high=_lab_float(row.get('reference_high')),
                reference_text=str(row.get('reference_text') or '')[:120],
                flag=str(row.get('flag') or 'unknown')[:20],
                source_filename=source_files
            ))

    hospitalization = Hospitalization(
        patient_id=patient.id,
        clinical_event_id=event.id,
        cage=cage or '',
        responsible_vet=responsible_vet or '',
        reason=reason or '',
        diagnosis=diagnosis or '',
        treatment_plan=treatment_plan or '',
        notes=notes or '',
        initial_weight=to_float(initial_weight),
        initial_temperature=to_float(initial_temperature),
        initial_heart_rate=to_int(initial_heart_rate),
        initial_respiratory_rate=to_int(initial_respiratory_rate),
        initial_mucous_membranes=initial_mucous_membranes or '',
        initial_crt=initial_crt or '',
        initial_hydration=initial_hydration or '',
        created_by=user.username,
        admission_date=argentina_now(),
        status='Internado'
    )

    db.add(hospitalization)

    if initial_weight and str(initial_weight).strip():
        patient.weight = to_float(initial_weight)

    db.commit()
    db.refresh(hospitalization)

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}',
        status_code=303
    )
def build_hospitalization_ai(hospitalization, patient, related_events, medications, fluids):
    now = argentina_now().replace(tzinfo=None)

    def to_float(value):
        try:
            if value is None:
                return None
            value = str(value).replace(',', '.').strip()
            if not value:
                return None
            return float(value)
        except Exception:
            return None

    def hours_from(dt):
        if not dt:
            return None
        return (now - dt.replace(tzinfo=None)).total_seconds() / 3600

    def clean_text(value):
        return (value or '').strip().lower()

    def med_minutes(time_text):
        try:
            if not time_text:
                return None
            h, m = str(time_text).split(':')[:2]
            return int(h) * 60 + int(m)
        except Exception:
            return None

    alerts = []
    score = 0

    vitals_events = [
        e for e in related_events
        if e.temperature or e.heart_rate or e.respiratory_rate or e.weight or e.mucous_membranes or e.crt or e.hydration
    ]

    last_control = related_events[0] if related_events else None
    last_vital = vitals_events[0] if vitals_events else None
    previous_vital = vitals_events[1] if len(vitals_events) > 1 else None

    temp = to_float(getattr(last_vital, 'temperature', None)) if last_vital else to_float(hospitalization.initial_temperature)
    hr = to_float(getattr(last_vital, 'heart_rate', None)) if last_vital else to_float(hospitalization.initial_heart_rate)
    rr = to_float(getattr(last_vital, 'respiratory_rate', None)) if last_vital else to_float(hospitalization.initial_respiratory_rate)
    weight = to_float(getattr(last_vital, 'weight', None)) if last_vital else to_float(hospitalization.initial_weight or patient.weight)

    prev_temp = to_float(getattr(previous_vital, 'temperature', None)) if previous_vital else None
    prev_hr = to_float(getattr(previous_vital, 'heart_rate', None)) if previous_vital else None
    prev_rr = to_float(getattr(previous_vital, 'respiratory_rate', None)) if previous_vital else None
    prev_weight = to_float(getattr(previous_vital, 'weight', None)) if previous_vital else None

    mucosas = (getattr(last_vital, 'mucous_membranes', '') if last_vital else hospitalization.initial_mucous_membranes) or ''
    crt = (getattr(last_vital, 'crt', '') if last_vital else hospitalization.initial_crt) or ''
    hidratacion = (getattr(last_vital, 'hydration', '') if last_vital else hospitalization.initial_hydration) or ''

    clinical_text = f"{mucosas} {crt} {hidratacion}".lower()
    all_recent_text = " ".join([clean_text(e.description) for e in related_events[:8]])

    danger_words = ['pálida', 'palida', 'cianótica', 'cianotica', 'gris', 'shock', 'colapso', 'débil', 'debil', 'estupor', 'convuls']
    warning_words = ['ictérica', 'icterica', 'congestiva', 'seca', 'deshidrat', 'pegajosa', 'deca', 'dolor', 'hiporexia']

    if temp is not None:
        if temp >= 40 or temp <= 36.5:
            score += 35
            alerts.append({'level': 'danger', 'icon': '🌡', 'title': 'Temperatura crítica', 'detail': f'Última temperatura: {temp:g} °C.'})
        elif temp >= 39.5 or temp <= 37:
            score += 22
            alerts.append({'level': 'warning', 'icon': '🌡', 'title': 'Temperatura fuera de rango', 'detail': f'Última temperatura: {temp:g} °C.'})

    if prev_temp is not None and temp is not None:
        delta = temp - prev_temp
        if abs(delta) >= 1:
            score += 12
            alerts.append({'level': 'warning', 'icon': '📈', 'title': 'Cambio térmico importante', 'detail': f'Variación de temperatura: {delta:+.1f} °C respecto al control previo.'})

    if hr is not None:
        if hr >= 200 or hr <= 50:
            score += 30
            alerts.append({'level': 'danger', 'icon': '❤️', 'title': 'Frecuencia cardíaca crítica', 'detail': f'Última FC: {hr:g} lpm.'})
        elif hr >= 180 or hr <= 60:
            score += 18
            alerts.append({'level': 'warning', 'icon': '❤️', 'title': 'Frecuencia cardíaca alterada', 'detail': f'Última FC: {hr:g} lpm.'})

    if prev_hr is not None and hr is not None and abs(hr - prev_hr) >= 35:
        score += 10
        alerts.append({'level': 'warning', 'icon': '📈', 'title': 'Cambio de FC', 'detail': f'FC varió {hr - prev_hr:+.0f} lpm respecto al control previo.'})

    if rr is not None:
        if rr >= 80 or rr <= 8:
            score += 30
            alerts.append({'level': 'danger', 'icon': '🫁', 'title': 'Frecuencia respiratoria crítica', 'detail': f'Última FR: {rr:g} rpm.'})
        elif rr >= 60 or rr <= 10:
            score += 18
            alerts.append({'level': 'warning', 'icon': '🫁', 'title': 'Frecuencia respiratoria alterada', 'detail': f'Última FR: {rr:g} rpm.'})

    if prev_rr is not None and rr is not None and abs(rr - prev_rr) >= 20:
        score += 10
        alerts.append({'level': 'warning', 'icon': '📈', 'title': 'Cambio de FR', 'detail': f'FR varió {rr - prev_rr:+.0f} rpm respecto al control previo.'})

    if prev_weight is not None and weight is not None and prev_weight > 0:
        weight_delta_percent = ((weight - prev_weight) / prev_weight) * 100
        if abs(weight_delta_percent) >= 5:
            score += 12
            alerts.append({'level': 'warning', 'icon': '⚖', 'title': 'Cambio de peso', 'detail': f'Peso varió {weight_delta_percent:+.1f}% respecto al control previo.'})

    if any(word in clinical_text for word in danger_words):
        score += 30
        alerts.append({'level': 'danger', 'icon': '👄', 'title': 'Perfusión preocupante', 'detail': 'Mucosas/TRC/hidratación con términos de riesgo.'})
    elif any(word in clinical_text for word in warning_words):
        score += 16
        alerts.append({'level': 'warning', 'icon': '👄', 'title': 'Perfusión a revisar', 'detail': 'Mucosas o hidratación requieren control.'})

    if 'vómitos / diarrea: sí' in all_recent_text or 'vomitos / diarrea: sí' in all_recent_text:
        score += 18
        alerts.append({'level': 'warning', 'icon': '🤢', 'title': 'Vómitos o diarrea registrados', 'detail': 'Checklist reciente con vómitos/diarrea positivos.'})

    if 'comió: no' in all_recent_text or 'comio: no' in all_recent_text:
        score += 10
        alerts.append({'level': 'warning', 'icon': '🍽', 'title': 'No comió', 'detail': 'Checklist reciente indica que no comió.'})

    if 'orinó: no' in all_recent_text or 'orino: no' in all_recent_text:
        score += 12
        alerts.append({'level': 'warning', 'icon': '🚽', 'title': 'No orinó', 'detail': 'Checklist reciente indica que no orinó.'})

    hours_since_control = hours_from(last_control.event_date) if last_control else None

    if hours_since_control is None:
        score += 20
        alerts.append({'level': 'warning', 'icon': '🕓', 'title': 'Sin controles registrados', 'detail': 'Todavía no hay controles/evoluciones en esta internación.'})
    elif hours_since_control >= 8:
        score += 25
        alerts.append({'level': 'danger', 'icon': '🕓', 'title': 'Control muy demorado', 'detail': f'Último control hace {int(hours_since_control)} horas.'})
    elif hours_since_control >= 4:
        score += 12
        alerts.append({'level': 'warning', 'icon': '🕓', 'title': 'Control próximo', 'detail': f'Último control hace {int(hours_since_control)} horas.'})

    now_minutes = now.hour * 60 + now.minute
    pending_medications = [m for m in medications if m.status == 'Pendiente']
    overdue_medications = []

    for m in pending_medications:
        mins = med_minutes(m.scheduled_time)
        if mins is not None and mins <= now_minutes:
            overdue_medications.append(m)

    if overdue_medications:
        score += min(35, len(overdue_medications) * 12)
        alerts.append({'level': 'danger', 'icon': '💊', 'title': 'Medicación vencida', 'detail': f'{len(overdue_medications)} medicación/es ya deberían haberse aplicado.'})
    elif pending_medications:
        score += min(20, len(pending_medications) * 7)
        alerts.append({'level': 'warning', 'icon': '💊', 'title': 'Medicación pendiente', 'detail': f'{len(pending_medications)} medicación/es pendiente/s.'})

    active_fluids = [f for f in fluids if f.status == 'Activo']

    if active_fluids:
        alerts.append({'level': 'ok', 'icon': '💧', 'title': 'Fluidoterapia activa', 'detail': f'{len(active_fluids)} plan/es activo/s.'})
    else:
        if hospitalization.status == 'Internado' and ('deshidrat' in clinical_text or 'shock' in clinical_text):
            score += 18
            alerts.append({'level': 'warning', 'icon': '💧', 'title': 'Revisar fluidoterapia', 'detail': 'Hay datos de hidratación/perfusión a revisar y no figura fluido activo.'})

    if hospitalization.status != 'Internado':
        priority_label = 'Alta / Cerrado'
        priority_class = 'closed'
        priority_icon = '⚪'
    elif score >= 70:
        priority_label = 'Crítico'
        priority_class = 'danger'
        priority_icon = '🔴'
    elif score >= 40:
        priority_label = 'Alta prioridad'
        priority_class = 'warning'
        priority_icon = '🟠'
    elif score >= 15:
        priority_label = 'Control cercano'
        priority_class = 'watch'
        priority_icon = '🟡'
    else:
        priority_label = 'Estable'
        priority_class = 'ok'
        priority_icon = '🟢'

    if not alerts:
        alerts.append({'level': 'ok', 'icon': '🟢', 'title': 'Sin alertas activas', 'detail': 'No se detectan alertas automáticas por ahora.'})

    nursing_tasks = []

    for m in overdue_medications:
        nursing_tasks.append({
            'icon': '💊',
            'title': m.medication_name or 'Medicación vencida',
            'detail': f'{m.scheduled_time or "Sin horario"} · {m.dose or "-"} · {m.route or "-"}',
            'priority': 'Urgente',
            'class': 'danger',
            'hospitalization_id': hospitalization.id
        })

    for m in pending_medications:
        if m not in overdue_medications:
            nursing_tasks.append({
                'icon': '💊',
                'title': m.medication_name or 'Medicación pendiente',
                'detail': f'{m.scheduled_time or "Sin horario"} · {m.dose or "-"} · {m.route or "-"}',
                'priority': 'Pendiente',
                'class': 'warning',
                'hospitalization_id': hospitalization.id
            })

    for f in active_fluids:
        nursing_tasks.append({
            'icon': '💧',
            'title': f.fluid_type or 'Fluidoterapia',
            'detail': f'{f.fluid_rate or "-"} ml/h · {f.drip_set or "-"}',
            'priority': 'Activa',
            'class': 'ok',
            'hospitalization_id': hospitalization.id
        })

    nursing_tasks.append({
        'icon': '📋',
        'title': 'Próximo control clínico',
        'detail': 'Registrar T°, FC, FR, mucosas, TRC, hidratación y evolución.',
        'priority': 'Prioritario' if priority_class in ['danger', 'warning'] else 'Rutina',
        'class': priority_class,
        'hospitalization_id': hospitalization.id
    })

    uti_monitor = [
        {'label': 'Temperatura', 'value': temp if temp is not None else '-', 'unit': '°C', 'class': 'danger' if temp is not None and (temp >= 40 or temp <= 36.5) else 'warning' if temp is not None and (temp >= 39.5 or temp <= 37) else 'ok'},
        {'label': 'FC', 'value': hr if hr is not None else '-', 'unit': 'lpm', 'class': 'danger' if hr is not None and (hr >= 200 or hr <= 50) else 'warning' if hr is not None and (hr >= 180 or hr <= 60) else 'ok'},
        {'label': 'FR', 'value': rr if rr is not None else '-', 'unit': 'rpm', 'class': 'danger' if rr is not None and (rr >= 80 or rr <= 8) else 'warning' if rr is not None and (rr >= 60 or rr <= 10) else 'ok'},
        {'label': 'Peso', 'value': weight if weight is not None else '-', 'unit': 'kg', 'class': 'warning' if prev_weight is not None and weight is not None and prev_weight > 0 and abs(((weight - prev_weight) / prev_weight) * 100) >= 5 else 'ok'},
        {'label': 'Mucosas', 'value': mucosas or '-', 'unit': '', 'class': 'danger' if any(word in clinical_text for word in danger_words) else 'warning' if any(word in clinical_text for word in warning_words) else 'ok'},
        {'label': 'Hidratación', 'value': hidratacion or '-', 'unit': '', 'class': 'warning' if any(word in clinical_text for word in warning_words) else 'ok'}
    ]

    clinical_summary = []
    clinical_summary.append(f"{priority_icon} Prioridad: {priority_label}.")
    if temp is not None:
        clinical_summary.append(f"Temperatura actual: {temp:g} °C.")
    if hr is not None:
        clinical_summary.append(f"FC actual: {hr:g} lpm.")
    if rr is not None:
        clinical_summary.append(f"FR actual: {rr:g} rpm.")
    if pending_medications:
        clinical_summary.append(f"Medicación pendiente: {len(pending_medications)}.")
    if active_fluids:
        clinical_summary.append(f"Fluidoterapia activa: {len(active_fluids)} plan/es.")

    return {
        'alerts': alerts,
        'score': min(score, 100),
        'priority_label': priority_label,
        'priority_class': priority_class,
        'priority_icon': priority_icon,
        'uti_monitor': uti_monitor,
        'nursing_tasks': nursing_tasks,
        'last_vital': last_vital,
        'hours_since_control': hours_since_control,
        'clinical_summary': clinical_summary,
        'pending_medications_count': len(pending_medications),
        'overdue_medications_count': len(overdue_medications),
        'active_fluids_count': len(active_fluids)
    }

@app.post('/hospitalizations/{hospitalization_id}/ai')
def hospitalization_ai_real(
    hospitalization_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)

    if not hospitalization:
        raise HTTPException(status_code=404, detail='Internación no encontrada')

    patient = hospitalization.patient

    related_events = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .order_by(ClinicalEvent.event_date.desc())
        .limit(12)
        .all()
    )

    medications = (
        db.query(HospitalizationMedication)
        .filter(HospitalizationMedication.hospitalization_id == hospitalization.id)
        .order_by(HospitalizationMedication.scheduled_time.asc())
        .all()
    )

    fluids = (
        db.query(HospitalizationFluid)
        .filter(HospitalizationFluid.hospitalization_id == hospitalization.id)
        .order_by(HospitalizationFluid.created_at.desc())
        .all()
    )

    fake_event = ClinicalEvent(
        patient_id=patient.id,
        patient=patient,
        event_type='Internación',
        title='Análisis IA de internación',
        description=f"""
Motivo de internación:
{hospitalization.reason or ''}

Diagnóstico inicial:
{hospitalization.diagnosis or ''}

Plan terapéutico:
{hospitalization.treatment_plan or ''}

Notas:
{hospitalization.notes or ''}

Medicación:
{chr(10).join([(m.scheduled_time or '-') + ' · ' + (m.medication_name or '-') + ' · ' + (m.dose or '-') + ' · ' + (m.route or '-') + ' · ' + (m.status or '-') for m in medications])}

Fluidoterapia:
{chr(10).join([(f.fluid_type or '-') + ' · ' + (f.fluid_rate or '-') + ' ml/h · ' + (f.status or '-') for f in fluids])}
""",
        diagnosis=hospitalization.diagnosis or '',
        treatment=hospitalization.treatment_plan or '',
        weight=hospitalization.initial_weight or patient.weight,
        temperature=hospitalization.initial_temperature,
        heart_rate=hospitalization.initial_heart_rate,
        respiratory_rate=hospitalization.initial_respiratory_rate,
        mucous_membranes=hospitalization.initial_mucous_membranes or '',
        crt=hospitalization.initial_crt or '',
        hydration=hospitalization.initial_hydration or '',
        event_date=argentina_now()
    )

    result = ai_clinical_summary(fake_event)

    return JSONResponse(result)
@app.get('/nursing', response_class=HTMLResponse)
def nursing_view(
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    active_hospitalizations = (
        db.query(Hospitalization)
        .filter(Hospitalization.status == 'Internado')
        .order_by(Hospitalization.admission_date.asc())
        .all()
    )

    nursing_patients = []
    global_tasks = []
    stats = {
        'total': 0,
        'critical': 0,
        'warning': 0,
        'watch': 0,
        'stable': 0,
        'overdue_meds': 0,
        'pending_meds': 0,
        'active_fluids': 0
    }

    for hospitalization in active_hospitalizations:
        patient = hospitalization.patient

        related_events = (
            db.query(ClinicalEvent)
            .filter(ClinicalEvent.patient_id == patient.id)
            .filter(ClinicalEvent.event_type.in_(['Internación', 'Control', 'Consulta clínica', 'Alta']))
            .order_by(ClinicalEvent.event_date.desc())
            .limit(30)
            .all()
        )

        medications = (
            db.query(HospitalizationMedication)
            .filter(HospitalizationMedication.hospitalization_id == hospitalization.id)
            .order_by(HospitalizationMedication.scheduled_time.asc(), HospitalizationMedication.created_at.desc())
            .all()
        )

        fluids = (
            db.query(HospitalizationFluid)
            .filter(HospitalizationFluid.hospitalization_id == hospitalization.id)
            .order_by(HospitalizationFluid.created_at.desc())
            .all()
        )

        hospital_ai = build_hospitalization_ai(
            hospitalization,
            patient,
            related_events,
            medications,
            fluids
        )

        stats['total'] += 1
        stats['overdue_meds'] += hospital_ai['overdue_medications_count']
        stats['pending_meds'] += hospital_ai['pending_medications_count']
        stats['active_fluids'] += hospital_ai['active_fluids_count']

        if hospital_ai['priority_class'] == 'danger':
            stats['critical'] += 1
        elif hospital_ai['priority_class'] == 'warning':
            stats['warning'] += 1
        elif hospital_ai['priority_class'] == 'watch':
            stats['watch'] += 1
        else:
            stats['stable'] += 1

        patient_card = {
            'hospitalization': hospitalization,
            'patient': patient,
            'owner': patient.owner if patient else None,
            'ai': hospital_ai,
            'tasks': hospital_ai['nursing_tasks'],
            'alerts': hospital_ai['alerts'],
            'last_event': related_events[0] if related_events else None,
            'medications': medications,
            'fluids': fluids
        }

        nursing_patients.append(patient_card)

        for task in hospital_ai['nursing_tasks']:
            global_tasks.append({
                'patient_name': patient.name if patient else '-',
                'cage': hospitalization.cage or '-',
                'hospitalization_id': hospitalization.id,
                'task': task,
                'priority_class': task.get('class') or hospital_ai['priority_class']
            })

    order_map = {'danger': 0, 'warning': 1, 'watch': 2, 'ok': 3, 'closed': 4}
    nursing_patients = sorted(
        nursing_patients,
        key=lambda item: order_map.get(item['ai']['priority_class'], 9)
    )

    global_tasks = sorted(
        global_tasks,
        key=lambda item: order_map.get(item['priority_class'], 9)
    )

    return templates.TemplateResponse(
        'nursing.html',
        {
            'request': request,
            'nursing_patients': nursing_patients,
            'global_tasks': global_tasks,
            'stats': stats,
            'now': argentina_now().replace(tzinfo=None),
            'today': argentina_now().date()
        }
    )


@app.get('/hospitalizations/nursing', response_class=HTMLResponse)
def hospitalizations_nursing_alias(
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    return RedirectResponse('/nursing', status_code=303)

@app.get('/hospitalizations/new', response_class=HTMLResponse)
def hospitalization_new_redirect(
    patient_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    patient = db.get(Patient, patient_id)

    if not patient:
        raise HTTPException(status_code=404)

    return RedirectResponse(
        f'/patients/{patient.id}/hospitalization',
        status_code=303
    )
@app.get('/hospitalizations/{hospitalization_id}', response_class=HTMLResponse)
def hospitalization_detail(
    hospitalization_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)

    if not hospitalization:
        raise HTTPException(status_code=404, detail='Internación no encontrada')

    patient = hospitalization.patient

    related_events = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .filter(ClinicalEvent.event_type.in_(['Internación', 'Control', 'Consulta clínica', 'Alta']))
        .order_by(ClinicalEvent.event_date.desc())
        .limit(30)
        .all()
    )

    medications = (
        db.query(HospitalizationMedication)
        .filter(
            HospitalizationMedication.hospitalization_id == hospitalization.id
        )
        .order_by(
            HospitalizationMedication.created_at.asc(),
            HospitalizationMedication.id.asc()
        )
        .all()
    )
    
    # ==========================================================
    # Orden cronológico real de las dosis
    # Respeta cambios de día: 23:00 → 07:00 → 15:00
    # ==========================================================
    
    now_argentina = argentina_now()
    treatment_groups = {}
    
    for medication in medications:
        notes_text = medication.notes or ''
    
        treatment_match = re.search(
            r'\[TRATAMIENTO_ACTIVO:([^\]]+)\]',
            notes_text
        )
    
        treatment_key = (
            treatment_match.group(1)
            if treatment_match
            else f'medication-{medication.id}'
        )
    
        treatment_groups.setdefault(
            treatment_key,
            []
        ).append(medication)
    
    for group_medications in treatment_groups.values():
    
        group_medications.sort(
            key=lambda item: item.id
        )
    
        previous_minutes = None
        day_offset = 0
    
        first_created_at = (
            group_medications[0].created_at
            or now_argentina
        )
    
        created_minutes = (
            first_created_at.hour * 60
            + first_created_at.minute
        )
    
        for index, medication in enumerate(group_medications):
    
            try:
                hour_text, minute_text = (
                    medication.scheduled_time or ''
                ).split(':')
    
                scheduled_hour = int(hour_text)
                scheduled_minute = int(minute_text)
    
            except Exception:
                medication.ui_scheduled_at = None
                medication.ui_status = 'Pendiente'
                continue
    
            scheduled_minutes = (
                scheduled_hour * 60
                + scheduled_minute
            )
    
            if index == 0:
                # Si el horario inicial ya pasó ampliamente,
                # se interpreta como perteneciente al día siguiente.
                if scheduled_minutes < created_minutes - 30:
                    day_offset = 1
    
            elif (
                previous_minutes is not None
                and scheduled_minutes < previous_minutes
            ):
                day_offset += 1
    
            scheduled_date = (
                first_created_at.date()
                + timedelta(days=day_offset)
            )
    
            scheduled_datetime = datetime.combine(
                scheduled_date,
                datetime.min.time()
            ).replace(
                hour=scheduled_hour,
                minute=scheduled_minute
            )
    
            medication.ui_scheduled_at = scheduled_datetime
    
            if medication.status == 'Aplicada':
                medication.ui_status = 'Aplicada'
    
            elif medication.status == 'Finalizada':
                medication.ui_status = 'Finalizada'
    
            else:
                minutes_difference = (
                    now_argentina - scheduled_datetime
                ).total_seconds() / 60
    
                if minutes_difference < 0:
                    medication.ui_status = 'Próxima'
    
                elif minutes_difference <= 30:
                    medication.ui_status = 'Pendiente'
    
                else:
                    medication.ui_status = 'Atrasada'
    
            previous_minutes = scheduled_minutes
    
    medications.sort(
        key=lambda medication: (
            medication.ui_scheduled_at
            or datetime.max,
            medication.id
        )
    )

    fluids = (
        db.query(HospitalizationFluid)
        .filter(HospitalizationFluid.hospitalization_id == hospitalization.id)
        .order_by(HospitalizationFluid.created_at.desc())
        .all()
    )
    # ==========================================
    # Últimos registros
    # ==========================================

    latest_event = related_events[0] if related_events else None

    latest_event_time = (
        latest_event.event_date.strftime("%H:%M")
        if latest_event and latest_event.event_date
        else "-"
    )

    latest_checklist = (
        db.query(ClinicalEvent)
        .filter(
            ClinicalEvent.patient_id == patient.id,
            ClinicalEvent.event_type == "Checklist enfermería"
        )
        .order_by(ClinicalEvent.event_date.desc())
        .first()
    )
    latest_checklist_time = (
        latest_checklist.event_date.strftime("%H:%M")
        if latest_checklist and latest_checklist.event_date
        else "-"
    )    
    latest_checklist_time = (
        latest_checklist.event_date.strftime("%H:%M")
        if latest_checklist and latest_checklist.event_date
        else "-"
    )

    hospital_ai = build_hospitalization_ai(
        hospitalization,
        patient,
        related_events,
        medications,
        fluids
    )
    latest_vitals = hospital_ai.get("last_vital")

    current_patient_status = {
        "temperature": (
            getattr(latest_vitals, "temperature", None)
            if latest_vitals and getattr(latest_vitals, "temperature", None) is not None
            else hospitalization.initial_temperature
        ),
        "heart_rate": (
            getattr(latest_vitals, "heart_rate", None)
            if latest_vitals and getattr(latest_vitals, "heart_rate", None) is not None
            else hospitalization.initial_heart_rate
        ),
        "respiratory_rate": (
            getattr(latest_vitals, "respiratory_rate", None)
            if latest_vitals and getattr(latest_vitals, "respiratory_rate", None) is not None
            else hospitalization.initial_respiratory_rate
        ),
        "weight": (
            getattr(latest_vitals, "weight", None)
            if latest_vitals and getattr(latest_vitals, "weight", None) is not None
            else (hospitalization.initial_weight or patient.weight)
        ),
        "mucous_membranes": (
            getattr(latest_vitals, "mucous_membranes", "")
            if latest_vitals and getattr(latest_vitals, "mucous_membranes", "")
            else hospitalization.initial_mucous_membranes
        ),
        "crt": (
            getattr(latest_vitals, "crt", "")
            if latest_vitals and getattr(latest_vitals, "crt", "")
            else hospitalization.initial_crt
        ),
        "hydration": (
            getattr(latest_vitals, "hydration", "")
            if latest_vitals and getattr(latest_vitals, "hydration", "")
            else hospitalization.initial_hydration
        ),
    }
    return templates.TemplateResponse(
        'hospitalization_detail.html',
        {
            'request': request,
            'hospitalization': hospitalization,
            'patient': patient,
            'related_events': related_events,
            'medications': medications,
            'fluids': fluids,
            'hospital_alerts': hospital_ai['alerts'],
            'priority_label': hospital_ai['priority_label'],
            'priority_class': hospital_ai['priority_class'],
            'priority_icon': hospital_ai['priority_icon'],
            'priority_score': hospital_ai['score'],
            'uti_monitor': hospital_ai['uti_monitor'],
            'nursing_tasks': hospital_ai['nursing_tasks'],
            'last_vital': hospital_ai['last_vital'],
            'hours_since_control': hospital_ai['hours_since_control'],
            'clinical_summary': hospital_ai['clinical_summary'],
            'today': argentina_now().date(),
            'current_time_hhmm': argentina_now().strftime('%H:%M'),
            'latest_event_time': latest_event_time,
            'latest_checklist_time': latest_checklist_time,
            'current_patient_status': current_patient_status
        }
    )
@app.get('/hospitalizations/{hospitalization_id}/print', response_class=HTMLResponse)
def hospitalization_print(
    hospitalization_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)

    if not hospitalization:
        raise HTTPException(
            status_code=404,
            detail='Internación no encontrada'
        )

    patient = hospitalization.patient
    owner = patient.owner if patient else None

    related_events = (
        db.query(ClinicalEvent)
        .filter(ClinicalEvent.patient_id == patient.id)
        .filter(
            ClinicalEvent.event_type.in_(
                ['Internación', 'Control', 'Consulta clínica', 'Alta']
            )
        )
        .order_by(ClinicalEvent.event_date.asc())
        .all()
    )

    medications = (
        db.query(HospitalizationMedication)
        .filter(
            HospitalizationMedication.hospitalization_id == hospitalization.id
        )
        .order_by(
            HospitalizationMedication.scheduled_time.asc(),
            HospitalizationMedication.created_at.asc()
        )
        .all()
    )

    fluids = (
        db.query(HospitalizationFluid)
        .filter(
            HospitalizationFluid.hospitalization_id == hospitalization.id
        )
        .order_by(HospitalizationFluid.created_at.asc())
        .all()
    )

    return templates.TemplateResponse(
        "hospitalization_print.html",
        {
            "request": request,
            "hospitalization": hospitalization,
            "patient": patient,
            "owner": owner,
            "related_events": related_events,
            "medications": medications,
            "fluids": fluids,
            "today": argentina_now().date()
        }
    )
@app.post('/hospitalizations/{hospitalization_id}/evolution')
def hospitalization_add_evolution(
    hospitalization_id: int,
    title: str = Form('Evolución de internación'),
    description: str = Form(''),
    temperature: str = Form(''),
    heart_rate: str = Form(''),
    respiratory_rate: str = Form(''),
    weight: str = Form(''),
    mucous_membranes: str = Form(''),
    crt: str = Form(''),
    hydration: str = Form(''),
    treatment: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)

    if not hospitalization:
        raise HTTPException(status_code=404, detail='Internación no encontrada')

    patient = hospitalization.patient

    def to_float(value):
        try:
            return float(str(value).replace(',', '.')) if value and str(value).strip() else None
        except ValueError:
            return None

    def to_int(value):
        try:
            return int(float(str(value).replace(',', '.'))) if value and str(value).strip() else None
        except ValueError:
            return None

    evolution_lines = []

    evolution_lines.append('Constantes / control')
    evolution_lines.append('')

    if temperature and str(temperature).strip():
        evolution_lines.append(f'🌡 Temperatura: {temperature} °C')

    if heart_rate and str(heart_rate).strip():
        evolution_lines.append(f'❤️ FC: {heart_rate} lpm')

    if respiratory_rate and str(respiratory_rate).strip():
        evolution_lines.append(f'🫁 FR: {respiratory_rate} rpm')

    if weight and str(weight).strip():
        evolution_lines.append(f'⚖ Peso: {weight} kg')

    if mucous_membranes and str(mucous_membranes).strip():
        evolution_lines.append(f'👄 Mucosas: {mucous_membranes}')

    if crt and str(crt).strip():
        evolution_lines.append(f'⏱ TRC: {crt}')

    if hydration and str(hydration).strip():
        evolution_lines.append(f'💧 Hidratación: {hydration}')

    if description and description.strip():
        evolution_lines.append('')
        evolution_lines.append('Evolución:')
        evolution_lines.append(description.strip())

    if treatment and treatment.strip():
        evolution_lines.append('')
        evolution_lines.append('Tratamiento / indicaciones:')
        evolution_lines.append(treatment.strip())

    final_description = '\n'.join(evolution_lines)

    event = ClinicalEvent(
        patient_id=patient.id,
        event_type='Control',
        title=title or 'Evolución de internación',
        description=final_description,
        treatment=treatment or '',
        temperature=to_float(temperature),
        heart_rate=to_int(heart_rate),
        respiratory_rate=to_int(respiratory_rate),
        weight=to_float(weight),
        mucous_membranes=mucous_membranes or '',
        crt=crt or '',
        hydration=hydration or '',
        created_by=user.username,
        event_date=argentina_now()
    )

    db.add(event)

    if weight and str(weight).strip():
        patient.weight = to_float(weight)

    db.commit()

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}#evolucion',
        status_code=303
    )
@app.post('/hospitalizations/{hospitalization_id}/checklist')
def hospitalization_checklist(
    hospitalization_id: int,
    ate: str = Form(''),
    drank: str = Form(''),
    urinated: str = Form(''),
    defecated: str = Form(''),
    vomiting_diarrhea: str = Form(''),
    cage_cleaned: str = Form(''),
    checklist_notes: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)

    if not hospitalization:
        raise HTTPException(status_code=404, detail='Internación no encontrada')

    patient = hospitalization.patient

    checklist_lines = []

    checklist_lines.append('Checklist de enfermería')
    checklist_lines.append('')
    checklist_lines.append(f"Comió: {'Sí' if ate else 'No'}")
    checklist_lines.append(f"Tomó agua: {'Sí' if drank else 'No'}")
    checklist_lines.append(f"Orinó: {'Sí' if urinated else 'No'}")
    checklist_lines.append(f"Defecó: {'Sí' if defecated else 'No'}")
    checklist_lines.append(f"Vómitos / diarrea: {'Sí' if vomiting_diarrhea else 'No'}")
    checklist_lines.append(f"Limpieza de jaula: {'Sí' if cage_cleaned else 'No'}")

    if checklist_notes and checklist_notes.strip():
        checklist_lines.append('')
        checklist_lines.append('Observaciones:')
        checklist_lines.append(checklist_notes.strip())

    event = ClinicalEvent(
        patient_id=patient.id,
        event_type='Checklist enfermería',
        title='Checklist de enfermería',
        description='\n'.join(checklist_lines),
        created_by=user.username,
        event_date=argentina_now()
    )

    db.add(event)
    db.commit()

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}#checklist',
        status_code=303
    )
@app.post('/hospitalizations/{hospitalization_id}/fluids')
def hospitalization_fluids(
    hospitalization_id: int,
    fluid_type: str = Form(''),
    fluid_rate: str = Form(''),
    ml_kg_h: str = Form(''),
    drip_set: str = Form(''),
    fluid_notes: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)

    if not hospitalization:
        raise HTTPException(status_code=404, detail='Internación no encontrada')

    patient = hospitalization.patient

    fluid = HospitalizationFluid(
        hospitalization_id=hospitalization.id,
        fluid_type=fluid_type or '',
        fluid_rate=fluid_rate or '',
        ml_kg_h=ml_kg_h or '',
        drip_set=drip_set or '',
        notes=fluid_notes or '',
        status='Activo',
        created_by=user.username
    )

    db.add(fluid)

    lines = [
        'Fluidoterapia',
        '',
        f'Tipo de fluido: {fluid_type or "-"}',
        f'Velocidad: {fluid_rate or "-"} ml/h',
        f'Ml/kg/h: {ml_kg_h or "-"}',
        f'Equipo: {drip_set or "-"}'
    ]

    if fluid_notes and fluid_notes.strip():
        lines.append('')
        lines.append('Observaciones:')
        lines.append(fluid_notes.strip())

    event = ClinicalEvent(
        patient_id=patient.id,
        event_type='Control',
        title='Fluidoterapia',
        description='\n'.join(lines),
        created_by=user.username,
        event_date=argentina_now()
    )

    db.add(event)
    db.commit()

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}#fluidoterapia',
        status_code=303
    )
@app.post('/hospitalizations/{hospitalization_id}/fluids/{fluid_id}/finish')
def hospitalization_fluid_finish(
    hospitalization_id: int,
    fluid_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)
    fluid = db.get(HospitalizationFluid, fluid_id)

    if not hospitalization or not fluid:
        raise HTTPException(status_code=404, detail='Fluidoterapia no encontrada')

    if fluid.hospitalization_id != hospitalization.id:
        raise HTTPException(status_code=404, detail='Fluidoterapia no corresponde a esta internación')

    fluid.status = 'Finalizado'

    if hasattr(fluid, 'finished_at'):
        fluid.finished_at = argentina_now()

    if hasattr(fluid, 'finished_by'):
        fluid.finished_by = user.username
    event = ClinicalEvent(
        patient_id=hospitalization.patient_id,
        event_type='Control',
        title='Fluidoterapia finalizada',
        description=(
            f'Fluido: {fluid.fluid_type or "-"}\n'
            f'Velocidad: {fluid.fluid_rate or "-"} ml/h\n'
            f'Ml/kg/h: {fluid.ml_kg_h or "-"}\n'
            f'Equipo: {fluid.drip_set or "-"}'
        ),
        created_by=user.username,
        event_date=argentina_now()
    )

    db.add(event)
    db.commit()

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}#fluidoterapia',
        status_code=303
    )
@app.post('/hospitalizations/{hospitalization_id}/fluids/{fluid_id}/delete')
def hospitalization_fluid_delete(
    hospitalization_id: int,
    fluid_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)
    fluid = db.get(HospitalizationFluid, fluid_id)

    if not hospitalization or not fluid:
        raise HTTPException(status_code=404, detail='Fluidoterapia no encontrada')

    if fluid.hospitalization_id != hospitalization.id:
        raise HTTPException(status_code=404, detail='Fluidoterapia no corresponde a esta internación')

    db.delete(fluid)
    db.commit()

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}#fluidoterapia',
        status_code=303
    )
def medication_interval_hours(frequency: str):
    freq = (frequency or '').strip().lower()

    freq_map = {
        'sid': 24,
        'cada 24 hs': 24,
        'cada 24h': 24,
        'q24h': 24,
        'bid': 12,
        'cada 12 hs': 12,
        'cada 12h': 12,
        'q12h': 12,
        'tid': 8,
        'cada 8 hs': 8,
        'cada 8h': 8,
        'q8h': 8,
        'qid': 6,
        'cada 6 hs': 6,
        'cada 6h': 6,
        'q6h': 6,
        'q48h': 48,
        'cada 48 hs': 48,
        'cada 48h': 48,
        'q72h': 72,
        'cada 72 hs': 72,
        'cada 72h': 72,
    }

    return freq_map.get(freq)


def build_medication_times(start_time: str, frequency: str):
    interval = medication_interval_hours(frequency)

    if not interval:
        return [start_time] if start_time else []

    if not start_time:
        return []

    try:
        start_dt = datetime.strptime(start_time, '%H:%M')
    except ValueError:
        return [start_time]

    times = []
    current = start_dt

    doses_per_day = max(1, int(24 / interval)) if interval <= 24 else 1

    for _ in range(doses_per_day):
        times.append(current.strftime('%H:%M'))
        current = current + timedelta(hours=interval)

    return times


@app.post('/hospitalizations/{hospitalization_id}/medications')
def hospitalization_medication_create(
    hospitalization_id: int,
    medication_name: str = Form(''),
    dose: str = Form(''),
    route: str = Form(''),
    frequency: str = Form(''),
    scheduled_time: str = Form(''),
    notes: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)

    if not hospitalization:
        raise HTTPException(status_code=404, detail='Internación no encontrada')

    schedule_times = build_medication_times(scheduled_time, frequency)

    if not schedule_times:
        schedule_times = [scheduled_time or '']

    treatment_code = uuid.uuid4().hex[:10]

    for med_time in schedule_times:
        medication = HospitalizationMedication(
            hospitalization_id=hospitalization.id,
            medication_name=medication_name or '',
            dose=dose or '',
            route=route or '',
            frequency=frequency or '',
            scheduled_time=med_time or '',
            notes=(
                f'[TRATAMIENTO_ACTIVO:{treatment_code}]\n'
                f'{notes or ""}'
            ).strip(),
            created_by=user.username,
            status='Pendiente'
        )

        db.add(medication)

    db.commit()

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}#medicacion',
        status_code=303
    )


@app.post('/hospitalizations/{hospitalization_id}/medications/{medication_id}/done')
def hospitalization_medication_done(
    hospitalization_id: int,
    medication_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)
    medication = db.get(HospitalizationMedication, medication_id)

    if not hospitalization or not medication:
        raise HTTPException(status_code=404, detail='Medicación no encontrada')

    if medication.hospitalization_id != hospitalization.id:
        raise HTTPException(status_code=404, detail='Medicación no corresponde a esta internación')

    medication.status = 'Aplicada'
    medication.applied_at = argentina_now()
    medication.applied_by = user.username

    event = ClinicalEvent(
        patient_id=hospitalization.patient_id,
        event_type='Control',
        title='Medicación aplicada',
        description=(
            f'Medicación aplicada: {medication.medication_name or "-"}\n'
            f'Dosis: {medication.dose or "-"}\n'
            f'Vía: {medication.route or "-"}\n'
            f'Frecuencia: {medication.frequency or "-"}\n'
            f'Horario programado: {medication.scheduled_time or "-"}\n'
            f'Aplicada por: {user.username}'
        ),
        created_by=user.username,
        event_date=argentina_now()
    )

    db.add(event)

    notes = medication.notes or ''
    treatment_code = ''

    if '[TRATAMIENTO_ACTIVO:' in notes:
        treatment_code = notes.split('[TRATAMIENTO_ACTIVO:')[1].split(']')[0]

    if treatment_code:
        pending_same_treatment = (
            db.query(HospitalizationMedication)
            .filter(HospitalizationMedication.hospitalization_id == hospitalization.id)
            .filter(HospitalizationMedication.notes.ilike(f'%[TRATAMIENTO_ACTIVO:{treatment_code}]%'))
            .filter(HospitalizationMedication.status == 'Pendiente')
            .filter(HospitalizationMedication.id != medication.id)
            .count()
        )

        if pending_same_treatment == 0:
            schedule_times = build_medication_times(
                medication.scheduled_time or '',
                medication.frequency or ''
            )

            if schedule_times:
                for med_time in schedule_times:
                    new_med = HospitalizationMedication(
                        hospitalization_id=hospitalization.id,
                        medication_name=medication.medication_name or '',
                        dose=medication.dose or '',
                        route=medication.route or '',
                        frequency=medication.frequency or '',
                        scheduled_time=med_time or '',
                        notes=medication.notes or '',
                        created_by=user.username,
                        status='Pendiente'
                    )

                    db.add(new_med)

    db.commit()

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}#medicacion',
        status_code=303
    )


@app.post('/hospitalizations/{hospitalization_id}/medications/{medication_id}/finish')
def hospitalization_medication_finish(
    hospitalization_id: int,
    medication_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)
    medication = db.get(HospitalizationMedication, medication_id)

    if not hospitalization or not medication:
        raise HTTPException(status_code=404, detail='Medicación no encontrada')

    if medication.hospitalization_id != hospitalization.id:
        raise HTTPException(status_code=404, detail='Medicación no corresponde a esta internación')

    notes = medication.notes or ''
    treatment_code = ''

    if '[TRATAMIENTO_ACTIVO:' in notes:
        treatment_code = notes.split('[TRATAMIENTO_ACTIVO:')[1].split(']')[0]

    if treatment_code:
        meds_to_finish = (
            db.query(HospitalizationMedication)
            .filter(HospitalizationMedication.hospitalization_id == hospitalization.id)
            .filter(HospitalizationMedication.notes.ilike(f'%[TRATAMIENTO_ACTIVO:{treatment_code}]%'))
            .filter(HospitalizationMedication.status == 'Pendiente')
            .all()
        )
    else:
        meds_to_finish = [medication]

    for med in meds_to_finish:
        med.status = 'Finalizada'
        med.notes = (med.notes or '').replace('[TRATAMIENTO_ACTIVO:', '[TRATAMIENTO_FINALIZADO:')

    event = ClinicalEvent(
        patient_id=hospitalization.patient_id,
        event_type='Control',
        title='Tratamiento finalizado',
        description=(
            f'Medicación finalizada: {medication.medication_name or "-"}\n'
            f'Dosis: {medication.dose or "-"}\n'
            f'Vía: {medication.route or "-"}\n'
            f'Frecuencia: {medication.frequency or "-"}\n'
            f'Finalizado por: {user.username}'
        ),
        created_by=user.username,
        event_date=argentina_now()
    )

    db.add(event)
    db.commit()

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}#medicacion',
        status_code=303
    )


@app.post('/hospitalizations/{hospitalization_id}/medications/{medication_id}/delete')
def hospitalization_medication_delete(
    hospitalization_id: int,
    medication_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)
    medication = db.get(HospitalizationMedication, medication_id)

    if not hospitalization or not medication:
        raise HTTPException(status_code=404, detail='Medicación no encontrada')

    if medication.hospitalization_id != hospitalization.id:
        raise HTTPException(status_code=404, detail='Medicación no corresponde a esta internación')

    db.delete(medication)
    db.commit()

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}#medicacion',
        status_code=303
    )
@app.post('/hospitalizations/{hospitalization_id}/discharge')
def hospitalization_discharge(
    hospitalization_id: int,
    discharge_summary: str = Form(''),
    discharge_indications: str = Form(''),
    db: Session = Depends(get_db),
    user: User = Depends(require_user)
):
    hospitalization = db.get(Hospitalization, hospitalization_id)

    if not hospitalization:
        raise HTTPException(status_code=404, detail='Internación no encontrada')

    patient = hospitalization.patient

    hospitalization.status = 'Alta'
    hospitalization.discharge_date = argentina_now()
    hospitalization.discharge_summary = discharge_summary or ''
    hospitalization.discharge_indications = discharge_indications or ''

    event = ClinicalEvent(
        patient_id=patient.id,
        event_type='Alta',
        title='Alta de internación',
        description=discharge_summary or '',
        treatment=discharge_indications or '',
        created_by=user.username,
        event_date=argentina_now()
    )

    db.add(event)
    db.commit()

    return RedirectResponse(
        f'/hospitalizations/{hospitalization.id}#alta',
        status_code=303
    )
@app.get('/health')
def health():
    return {'status': 'ok', 'app': 'Los Aromos Cloud'}
    

    
